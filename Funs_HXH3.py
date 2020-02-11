
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from pandas import Series, DataFrame
import math
import scipy.stats as stats
from datetime import datetime, timedelta

import pickle
import pypyodbc
import logging
import sys
import os
import scipy.optimize as sco
from scipy.stats import norm

import smtplib
from email.mime.text import MIMEText

from matplotlib.font_manager import FontProperties
from scipy.stats import norm
from statsmodels import regression

import scipy.optimize as sco
import copy

import time


import matplotlib as mpl
mpl.rcParams['font.sans-serif'] = ['SimHei']
mpl.rcParams['axes.unicode_minus'] = False


# ============================================= LocalSQLServer====================================================== #

# pypyodbc 是pyodbc的Python实现

class Pyodbc:
    """
    对pypyodbc的简单封装

    用法：

    """

    def __init__(self, Server, DataBase, uid='', pwd='', Driver = 'SQL Server', Certificate='SQL Server'):

        self.Certificate = Certificate
        if self.Certificate == 'SQL Server':
            self.Driver = Driver
            self.Server = Server
            self.DataBase = DataBase
            self.uid = uid
            self.pwd = pwd
        elif self.Certificate == 'Windows':
            self.Driver = Driver
            self.Server = Server
            self.DataBase = DataBase
            self.TrustServerCertificate = 'yes'

    def __GetConnect(self):
        """
        得到连接信息
        返回: conn.cursor()
        """
        if not self.DataBase:
            raise(NameError, "No DataBase Specified.")
        if self.Certificate == 'SQL Server':
            self.conn = pypyodbc.connect(
                'Driver={' + self.Driver + '};' +
                'Server=' + self.Server + ';' +
                'Database=' + self.DataBase + ';' +
                'uid=' + self.uid + ';' +
                'pwd=' + self.pwd)

        elif self.Certificate == 'Windows':
            self.conn = pypyodbc.connect(
                'Driver={' + self.Driver + '};' +
                'Server=' + self.Server + ';' +
                'Database=' + self.DataBase + ';' +
                'TrustServerCertificate=' + self.Certificate)

        # connStr = 'Driver={SQL Server};' + 'Server=172.18.1.54;' + 'Database=Wind;' + 'TrustServerCertificate=yes;'

        cur = self.conn.cursor()
        if not cur:
            raise(NameError, 'Failed to connect database.')
        else:
            return cur

    def ExecQuery(self, sql, colnames):
        """
        执行查询语句
        返回的是一个包含tuple的list，list的元素是记录行，tuple的元素是每行记录的字段

        调用示例：
                ms = MSSQL(host="localhost",user="sa",pwd="123456",db="PythonWeiboStatistics")
                resList = ms.ExecQuery("SELECT id,NickName FROM WeiBoUser")
                #返回的是一个包含tuple的list，list的元素是记录行，tuple的元素是每行记录的字段
                for (id,NickName) in resList:
                    print str(id),NickName
        """
        cur = self.__GetConnect()
        try:
            cur.execute(sql)
            res_df = DataFrame(cur.fetchall(), columns=colnames)
        except:
            raise(NameError, 'Fetching Data Error')

        # 查询完毕后必须关闭连接
        self.conn.close()
        return res_df
        # 返回一个dataframe结构，可能为空，None


# Pyodbc2是保留原数据库中表的列名
class Pyodbc2:
    """
    对pypyodbc的简单封装

    用法：

    """

    def __init__(self, Server, DataBase, uid='', pwd='', Driver = 'SQL Server', Certificate='SQL Server'):

        self.Certificate = Certificate
        if self.Certificate == 'SQL Server':
            self.Driver = Driver
            self.Server = Server
            self.DataBase = DataBase
            self.uid = uid
            self.pwd = pwd
        elif self.Certificate == 'Windows':
            self.Driver = Driver
            self.Server = Server
            self.DataBase = DataBase
            self.TrustServerCertificate = 'yes'

    def __GetConnect(self):
        """
        得到连接信息
        返回: conn.cursor()
        """
        if not self.DataBase:
            raise(NameError, "No DataBase Specified.")
        if self.Certificate == 'SQL Server':
            self.conn = pypyodbc.connect(
                'Driver={' + self.Driver + '};' +
                'Server=' + self.Server + ';' +
                'Database=' + self.DataBase + ';' +
                'uid=' + self.uid + ';' +
                'pwd=' + self.pwd)

        elif self.Certificate == 'Windows':
            self.conn = pypyodbc.connect(
                'Driver={' + self.Driver + '};' +
                'Server=' + self.Server + ';' +
                'Database=' + self.DataBase + ';' +
                'TrustServerCertificate=' + self.Certificate)

        # connStr = 'Driver={SQL Server};' + 'Server=172.18.1.54;' + 'Database=Wind;' + 'TrustServerCertificate=yes;'

        cur = self.conn.cursor()
        if not cur:
            raise(NameError, 'Failed to connect database.')
        else:
            return cur

    def ExecQuery(self, sql):
        """
        执行查询语句
        返回的是一个包含tuple的list，list的元素是记录行，tuple的元素是每行记录的字段

        调用示例：
                ms = MSSQL(host="localhost",user="sa",pwd="123456",db="PythonWeiboStatistics")
                resList = ms.ExecQuery("SELECT id,NickName FROM WeiBoUser")
                #返回的是一个包含tuple的list，list的元素是记录行，tuple的元素是每行记录的字段
                for (id,NickName) in resList:
                    print str(id),NickName
        """
        cur = self.__GetConnect()
        try:
            cur.execute(sql)
            col_name_list = [tuple[0].upper() for tuple in cur.description]
            res_df = DataFrame(cur.fetchall(), columns=col_name_list)
        except:
            raise(NameError, 'Fetching Data Error')

        # 查询完毕后必须关闭连接
        self.conn.close()
        return res_df
        # 返回一个dataframe结构，可能为空，None



# ============================================= LogConfig ======================================================= #

def LogConfig(filename):

	user_path = './log/' + datetime.today().strftime('%Y%m%d')
	folder = os.path.exists(user_path)
	if not folder:
		os.makedirs(user_path)

	logging.basicConfig(
		level=logging.INFO,
		format='%(asctime)s %(levelname) -2s: %(message)s',
		filename=user_path + '/' + filename,
		filemode='w'
	)

	console = logging.StreamHandler()
	console.setLevel(logging.INFO)
	formatter = logging.Formatter('%(asctime)s %(levelname) -8s: %(message)s')
	console.setFormatter(formatter)
	logging.getLogger('').addHandler(console)


# ============================================= sendEmail ======================================================= #
def send163email(receivers, subject, content):
    # receivers: list of strings

    # set SMTP Service
    mail_host = "smtp.163.com"  # SMTP Service
    mail_user = "xiaohuhuang6"  # user
    mail_pass = "Huge69064746"  # password

    sender = 'xiaohuhuang6@163.com'  # sender email address

    # email body
    message = MIMEText(content, 'plain', 'utf-8')  # content, format, code
    message['From'] = "{}".format(sender)
    message['To'] = ",".join(receivers)
    message['Subject'] = subject

    try:
        smtpObj = smtplib.SMTP_SSL(mail_host, 465)  # SSL port 465
        smtpObj.login(mail_user, mail_pass)  # log verification
        smtpObj.sendmail(sender, receivers, message.as_string())  # send message
        print("Mail has been send successfully.")
        return 1
    except smtplib.SMTPException as err:
        print(err)
        return 0

# ============================================= quantTools ======================================================= #
# --- .pkl文件读写函数
def pkl_read_write(filename, flag, input_data=None):
	'''
	:param filename: 读写文件名
	:param flag: 读/写
	:param input_data: 写入数据
	'''
	if flag == 'read':
		try:
			pkl_file = open(filename, 'rb')
			data = pickle.load(pkl_file)
			pkl_file.close()
			return data
		except:
			logging.warning(filename + ' does not exist in current path')
			return DataFrame()
	elif flag == 'write':
		pkl_file = open(filename, 'wb')
		pickle.dump(input_data, pkl_file)
		pkl_file.close()
		return True


def dataCheckLoad(dataNames, relativePath):

	if isinstance(dataNames, str):
		dataNames = [dataNames]

	flags = []
	global_vars = np.array(list(globals().keys()))

	for dataName in dataNames:
		if len(np.where(global_vars == dataName)[0]) == 1:
			# 存在该变量，全局
			flags.append(True)
			continue

		filePath = relativePath + dataName + '.pkl'
		dat = pkl_read_write(filePath, 'read')

		if not isinstance(dat, DataFrame):
			globals()[dataName] = dat
			flags.append(True)
			logging.info(dataName + ' is loaded into Global Environment!!')

		else:
			if not dat.empty:
				globals()[dataName] = dat
				flags.append(True)
				logging.info(dataName + ' is loaded into Global Environment!!')
			else:
				flags.append(False)
				logging.warning(dataName + ' load fail!!! Empty or Not Exist')

	if len(flags) == 1:
		flags = flags[0]

	return flags

def dataAssign(dataNames=None, relativePath='./rawData/'):

	if isinstance(dataNames, str):
		dataNames = [dataNames]

	res = []
	for dataName in dataNames:
		if dataCheckLoad(dataName, relativePath):
			res.append(globals()[dataName])
		else:
			res.append([])

	if len(res) == 1:
		res = res[0]

	return res

# 这种处理方式只有在当时存在，因为全局付给了局部变量，在子函数退出时，不再存在!!!



# 在module中，globals赋值不起作用
# 把变量赋值到环境中
def set_EnvData(dataNames, dataList):

	if isinstance(dataNames, str):
		dataNames = [dataNames]
		dataList = [dataList]

	for dataName in dataNames:
		# print(globals()[dataName])
		globals()[dataName] = dataList[dataNames.index(dataName)]

	return True

# 判断环境是否存在变量
def exist_EnvData(dataName):

	global_vars = np.array(list(globals().keys()))
	if len(np.where(global_vars == dataName)[0]) == 1:
		return True
	else:
		return False

# 获取环境中变量的值
def get_EnvData(dataNames, copy=False):

	if isinstance(dataNames, str):
		dataNames = [dataNames]

	res = []
	for dataName in dataNames:
		res.append(globals()[dataName])

	# 返回一个的时候，就不以列表返回
	if len(res) == 1:
		res = res[0]

	if copy:
		return res.copy()
	else:
		return res


def gen_CoreGenData(dataName):
	expr = 'gen_' + dataName + '()'
	tmp = eval(expr)
	return tmp


def get_CoreGenData(dataName, copy=False):

	if exist_EnvData(dataName):
		return get_EnvData(dataName, copy=copy)
	else:
		# 因为中间变量没有保存在本地，因此，需要产生
		gen_CoreGenData(dataName)  # 该函数没有返回真是变量值，只是一个True or False
		return get_EnvData(dataName, copy=copy)

def clearVariables():

	keeps = ['__builtins__', 'sys', 'clearVariables']
	gVariables = list(set(globals().keys()) - set(keeps))
	if len(gVariables) > 0:
		for gVariable in gVariables:
			del globals()[gVariable]

	lVariables = list(set(globals().keys()) - set(keeps))
	if len(lVariables) > 0:
		for lVariable in lVariables:
			del globals()[lVariables]

	return True





# Monte Carlo Simulation of Efficient Frontier
def MCEfficientFrontier(ret, rf):

	def statistics(rets, wts, rf):
		wts = np.array(wts)
		port_return = np.sum(rets.mean() * wts) * 240
		port_std = np.sqrt(np.dot(wts.T, np.dot(rets.cov() * 240, wts)))
		return np.array([port_return, port_std, (port_return - rf)/port_std])

	port_return = []
	port_std = []
	port_sharpe = []
	wtsSet = []

	for i in range(0, 4000):
		wts = np.random.random(3)
		wts /= np.sum(wts)
		temp = statistics(ret, wts, rf)
		port_return.append(temp[0])
		port_std.append(temp[1])
		port_sharpe.append(temp[2])
		wtsSet.append(wts)

	port_return = np.array(port_return)
	port_std = np.array(port_std)
	port_sharpe = np.array(port_sharpe)

	plt.figure()
	plt.scatter(port_std, port_return, c=(port_return -rf)/port_std, marker='o')
	plt.grid(True)
	plt.xlabel('Excepted Volatility')
	plt.ylabel('Expected Return')
	plt.colorbar(label='Sharpe Ratio')
	plt.title('Efficient Frontier')

	plt.show()


# ============================================= Funs_GenPublicData ======================================================= #

# 获取公共数据StockTradeInfo,并存入运行环境中
def gen_StockTradeInfo():

	AShareEODPrices, StockTradeStatus, AShareMoneyFlow,  CapitalizationEOD = \
		dataAssign(['AShareEODPrices', 'StockTradeStatus', 'AShareMoneyFlow', 'CapitalizationEOD'])
	dat = pd.merge(AShareEODPrices.ix[:, ['Symbol', 'Date', 'Close', 'Volume', 'Amount', 'AdjFactor']],
				   StockTradeStatus.ix[:, ['Symbol', 'Date', 'TradeStatus']], on=['Symbol', 'Date'])

	dat = pd.merge(dat, CapitalizationEOD.ix[:, ['Symbol', 'Date', 'FloatAShr', 'TotShr']], on=['Symbol', 'Date'])

	temp_mf = pd.merge(dat.ix[:, ['Symbol', 'Date']], AShareMoneyFlow, how='left', on=['Symbol', 'Date'])
	AShareMoneyFlow = temp_mf.fillna(value=0)

	dat = pd.merge(dat, AShareMoneyFlow.ix[:, ['Symbol', 'Date', 'MFInst', 'MFLarge', 'MFMed', 'MFSmall', 'MFOpen', 'MFClose', 'MFTot']],
				   on=['Symbol', 'Date'] )


	# Add some columns
	dat['AdjClose'] = dat['AdjFactor'] * dat['Close']

	dat['Vwap'] = dat['Close']
	dat['AdjVwap'] = dat['AdjClose']
	tmp = dat['Volume'] > 0
	dat.ix[tmp, 'Vwap'] = dat.ix[tmp, 'Amount'] / dat.ix[tmp, 'Volume']
	dat.ix[tmp, 'AdjVwap'] = dat.ix[tmp, 'AdjFactor'] * dat.ix[tmp, 'Amount'] / dat.ix[tmp, 'Volume']    # uniform the units


	dat['AFloatCap'] = dat['Close'] * dat['FloatAShr']
	dat['MarketValue'] = dat['Close'] * dat['TotShr']
	dat['TurnOver'] = dat['Volume'] / dat['FloatAShr']

	# sort and 存入环境中
	dat = dat.sort_values(by=['Date', 'Symbol'])
	set_EnvData('StockTradeInfo', dat)

	del StockTradeStatus, AShareMoneyFlow,  CapitalizationEOD

	return True

#
# 市值分组
def gen_StockMarketValue():

	if exist_EnvData('StockTradeInfo'):
		dat = get_EnvData('StockTradeInfo')
	else:
		AShareEODPrices,  CapitalizationEOD = \
			dataAssign(['AShareEODPrices', 'CapitalizationEOD'])
		dat = pd.merge(AShareEODPrices.ix[:, ['Symbol', 'Date', 'Close']],
					   CapitalizationEOD.ix[:, ['Symbol', 'Date', 'TotShr']], on=['Symbol', 'Date'])
		dat['MarketValue'] = dat['Close'] * dat['TotShr']

	dat = dat.ix[:, ['Symbol', 'Date', 'MarketValue']]

	dat['MVGroup'] = dat.groupby(by='Date')['MarketValue'].apply(lambda x: pd.qcut(x, 5, labels=range(1, 5 + 1)))

	# sort and 存入环境中
	dat = dat.sort_values(by=['Date', 'Symbol'])
	dat.ix[:, 'MVGroup'] = dat['MVGroup'].astype('float64')
	set_EnvData('StockMarketValue', dat)

	return True


def gen_StockTDayReturnDF():

	AShareEODPrices = dataAssign(['AShareEODPrices'])
	dat = AShareEODPrices.ix[:, ['Symbol', 'Date', 'PctChg']]

	set_EnvData('StockTDayReturnDF', dat)
	return True

def gen_IndexTDayReturnDF():

	AIndexEODPrices = dataAssign(['AIndexEODPrices'])
	dat = AIndexEODPrices.ix[:, ['Symbol', 'Date', 'PctChg']]

	set_EnvData('IndexTDayReturnDF', dat)
	return True



def gen_AllStocks():

	AShareEODPrices = dataAssign(['AShareEODPrices'])
	dat = Series(AShareEODPrices.ix[:, 'Symbol'].unique().tolist())

	set_EnvData('AllStocks', dat)
	return True


def gen_AShareCalendar():

	dat = dataAssign(['AShareCalendar'])
	set_EnvData('AShareCalendar', dat)
	return True


def getStockTDayReturn(stockCode):

	# print(exist_EnvData('StockTDayReturnDF'))
	# 这里StockTDayReturnDF已经在环境中了，没有一直产生
	tmp = get_CoreGenData('StockTDayReturnDF')

	dat = tmp.ix[tmp['Symbol'] == stockCode, ['Date', 'PctChg']].copy()
	dat.index = dat['Date'].tolist()
	dat = dat['PctChg']

	return dat


# 指数收益Series
def getIndexTDayReturn(indexCode):

	tmp = get_CoreGenData('IndexTDayReturnDF')

	dat = tmp.ix[tmp['Symbol'] == indexCode, ['Date', 'PctChg']].copy()
	dat.index = dat['Date'].tolist()
	dat = dat['PctChg']

	return dat


def gen_AShareEODPrices():

	dat = dataAssign('AShareEODPrices')
	set_EnvData('AShareEODPrices', dat)
	return True

def gen_AShareDividend():

	dat = dataAssign('AShareDividend')
	set_EnvData('AShareDividend', dat)
	return True


def gen_REF_DF():

	dat = dataAssign('AShareEODPrices')
	dat = dat.ix[:, ['Symbol', 'Date']]

	set_EnvData('REF_DF', dat)
	return True


def gen_CapitalizationEOD():
	dat = dataAssign('CapitalizationEOD')
	set_EnvData('CapitalizationEOD', dat)
	return True

def gen_IndustryCITICS():
	dat = dataAssign('IndustryCITICS')
	set_EnvData('IndustryCITICS', dat)
	return True

def gen_AIndexHS300FreeWeight():
	dat = dataAssign('AIndexHS300FreeWeight')
	set_EnvData('AIndexHS300FreeWeight', dat)
	return True



def gen_AIndexMembers():
	dat = dataAssign('AIndexMembers')
	set_EnvData('AIndexMembers', dat)
	return True


def gen_refPeriodDates():

	dat = dataAssign('refPeriodDates', './supplementaryData/')
	set_EnvData('refPeriodDates', dat)
	return True

# 底层函数去写路径，不再在外层把路径写出来, 要注意后期和dataAssign整合


# calculate TTMFactor
def gen_TTMFactor():

	AShareEODPrices = get_CoreGenData('AShareEODPrices')
	CapitalizationEOD = get_CoreGenData('CapitalizationEOD')

	tmp = pd.merge(AShareEODPrices, CapitalizationEOD, on=['Symbol', 'Date'])
	tmp = tmp.ix[:, ['Symbol', 'Date', 'Close', 'TotShr']]
	tmp['MarketValue'] = tmp['Close'] * tmp['TotShr']

	dat = dataAssign('TTMHisRevised')

	dat = pd.merge(tmp, dat, on=['Symbol', 'Date'], how='left')
	# dat = pd.merge(tmp, dat, on=['Symbol', 'Date'])
	# dat = dat.groupby('Symbol').apply(lambda c: c.fillna(method='ffill'))

	# calculate
	dat['NetProfitPCTTM'] = dat['NetProfitParent']

	dat['EPSTTM'] = dat['NetProfitParent'] / dat['TotShr']
	dat['InvPETTM'] = dat['NetProfitParent'] / dat['MarketValue']

	dat['TotOpRevPSTTM'] = dat['TotOperRev'] / dat['TotShr']
	dat['TotOpRevPSOverPriceTTM'] = dat['TotOperRev'] / dat['MarketValue']

	dat['CashIncrPSTTM'] = dat['NetIncrCash'] / dat['TotShr']
	dat['CashIncrPSOverPriceTTM'] = dat['NetIncrCash'] / dat['MarketValue']

	dat['CFPSTTM'] = dat['NetCashFlowsOper'] / dat['TotShr']
	dat['CFPSOverPriceTTM'] = dat['NetCashFlowsOper'] / dat['MarketValue']

	dat['OpRatioTTM'] = dat['FaOp'] / dat['FaEBT']
	dat['OpCashFlowTTM'] = dat['NetCashFlowsOper']
	dat['TotOpRevTTM'] = dat['TotOperRev']
	dat['EBITTTM'] = dat['FaEBIT']

	set_EnvData('TTMFactor', dat)

	return True


# calculate BalanceSheetFactor
def gen_BalanceSheetFactor():

	tmp = get_CoreGenData('StockTradeInfo')
	tmp = tmp.ix[:, ['Symbol', 'Date', 'MarketValue']]

	dat = dataAssign('BalanceSheetRevised')

	dat = pd.merge(tmp, dat, on=['Symbol', 'Date'], how='left')
	# dat = dat.groupby('Symbol').apply(lambda c: c.fillna(method='ffill'))

	# calculate
	dat['Equity'] = dat['TotShrHldrExclMin']
	dat['InvPB'] = dat['Equity'] / dat['MarketValue']
	dat['OpCapital'] = dat['TotCurAssets'] - dat['TotCurLiab']

	# dat['TotLiab'] = dat['TotLiab']
	dat['TotAssets'] = dat['TotCurAssets'] + dat['TotNonCurAssets']

	dat['EquityDebtRatio'] = dat['Equity'] / dat['TotLiab']
	# dat['FixAssets'] = dat['FixAssets']

	dat['CurrentRatio'] = dat['TotCurAssets'] / dat['TotCurLiab']

	set_EnvData('BalanceSheetFactor', dat)


	return True



# ============================================= Funs_Common ======================================================= #


# # priceType: Open, Close, Vwap
def get_PFun(stockList, Date, priceType='AdjClose'):

	if isinstance(stockList, str):
		stockList = [stockList]

	dat = get_CoreGenData('StockTradeInfo', copy=False)
	temp = dat.ix[dat['Date'] == Date, ['Symbol', priceType]].copy()
	price = temp.ix[temp['Symbol'].isin(stockList), :]
	price = Series(list(price[priceType]), index=list(price['Symbol']))
	price = price.reindex(stockList)
	return price


# 获取日期函数
def getTradingDays(start_dt, end_dt, flag=0):

	# 完善，是否存在环境中
	tmp = get_CoreGenData('AShareCalendar')
	# 在这个层级还有exist_EnvData, 进去之后就没有了

	if flag == 0:
		dat = tmp.ix[(tmp['Date'] >= start_dt) & (tmp['Date'] <= end_dt), 'Date']
	elif flag == 1:
		dat = tmp.ix[(tmp['Date'] > start_dt) & (tmp['Date'] <= end_dt), 'Date']
	elif flag == 2:
		dat = tmp.ix[(tmp['Date'] >= start_dt) & (tmp['Date'] < end_dt), 'Date']
	elif flag == 3:
		dat = tmp.ix[(tmp['Date'] > start_dt) & (tmp['Date'] < end_dt), 'Date']

	# 如果只有一个数，返回数据类型仍然是Series
	return dat


def getTradeStatus(stockList, Date):

	if isinstance(stockList, str):
		stockList = [stockList]

	dat = get_CoreGenData('StockTradeInfo', copy=False)
	tmp = dat.ix[dat['Date'] == Date, ['Symbol', 'TradeStatus']]
	tmp = tmp.ix[tmp['Symbol'].isin(stockList), :]
	tmp.index = tmp['Symbol'].tolist()
	tmp = tmp['TradeStatus']
	tmp = tmp.reindex(stockList)
	# 保持score的顺序
	return tmp


# def handleRawPortReturn_PeriodReturn(rawPortReturn, portfolioBuiltDates):
# 	# testDates = rawPortReturn.index.tolist()
# 	# periodEndDates = []
# 	#
# 	# for portfolioBuiltDate in portfolioBuiltDates:
# 	# 	# find the day before PortfolioBuiltDate, means the end of a period
# 	# 	if testDates.index(portfolioBuiltDate) != 0:
# 	# 		# 第一天不要
# 	# 		periodEndIndex = testDates.index(portfolioBuiltDate)
# 	# 		periodEndDates = periodEndDates + [testDates[periodEndIndex]]
# 	#
# 	# # Obtain the period return
# 	# portPeriodReturn = rawPortReturn.ix[periodEndDates, :]
#
# 	portPeriodReturn = rawPortReturn.ix[portfolioBuiltDates[1:], :]
# 	return portPeriodReturn


"""
Function: transfer raw port return to normal Trading day netvalue
Input: rawPortReturn, a dataframe (index: date, columns: ports) with accumulative return in a period, PortfolioBuiltDates
Output: port_d_netvalue, normal trading day netvalue
"""

def handleRawPortReturn_TDayNetValue(rawPortReturn, portfolioBuiltDates):

	testDates = rawPortReturn.index.tolist()
	port_d_netvalue = DataFrame(np.nan * np.zeros(rawPortReturn.shape),
						index=rawPortReturn.index, columns=rawPortReturn.columns)

	# the former PortfolioBuiltDate
	periodStartDate = portfolioBuiltDates[0]
	# 初始化1
	periodEndValue = Series(np.ones(rawPortReturn.shape[1]), index=rawPortReturn.columns)

	for portfolioBuiltDate in portfolioBuiltDates:
		if testDates.index(portfolioBuiltDate) != 0:
			periodStartIndex = testDates.index(periodStartDate)
			periodEndIndex = testDates.index(portfolioBuiltDate)
			port_d_netvalue.ix[periodStartIndex:periodEndIndex+1, :] = \
				(rawPortReturn.ix[periodStartIndex:periodEndIndex+1, :] + 1) * periodEndValue

			# Update periodStartDate and periodEndValue
			periodEndValue = port_d_netvalue.ix[periodEndIndex, :]
			if portfolioBuiltDate < testDates[-1]:
				periodStartDate = testDates[testDates.index(portfolioBuiltDate) + 1]

	# Remaining days
	if portfolioBuiltDate < testDates[-1]:
		port_d_netvalue.ix[testDates.index(periodStartDate):, :] = \
					(rawPortReturn.ix[testDates.index(periodStartDate):, :] + 1) * periodEndValue

	return port_d_netvalue



# def handleRawPortReturn_TDayNetValue(rawPortReturn, PortfolioBuiltDates):
#
# 	testDates = rawPortReturn.index.tolist()
# 	port_d_netvalue = DataFrame(np.nan * np.zeros(rawPortReturn.shape), index=rawPortReturn.index, columns=rawPortReturn.columns)
#
# 	# the former PortfolioBuiltDate
# 	oldBenchDate = PortfolioBuiltDates[0]
# 	# 初始化1
# 	oldPeriodEndValue = Series(np.ones(rawPortReturn.shape[1]), index=rawPortReturn.columns)
#
# 	for PortfolioBuiltDate in PortfolioBuiltDates:
# 		if testDates.index(PortfolioBuiltDate) != 0:
# 			oldBenchIndex = testDates.index(oldBenchDate)
# 			periodEndIndex = testDates.index(PortfolioBuiltDate) - 1
# 			port_d_netvalue.ix[oldBenchIndex:periodEndIndex+1, :] = \
# 				(rawPortReturn.ix[oldBenchIndex:periodEndIndex+1, :] + 1) * oldPeriodEndValue
#
# 			# Update oldPeriodValue and oldBenchDate
# 			oldPeriodEndValue = port_d_netvalue.ix[periodEndIndex, :]
# 			oldBenchDate = PortfolioBuiltDate
#
# 	# Remaining days
# 	port_d_netvalue.ix[testDates.index(oldBenchDate):, :] = \
# 				(rawPortReturn.ix[testDates.index(oldBenchDate):, :] + 1) * oldPeriodEndValue
#
# 	return port_d_netvalue


def handleNetvalue_Return(ref_dat, type=1):

	if type == 1:
		# netvalue to return
		if len(ref_dat.shape) > 1:
			# input is a dataframe
			fd_return = np.zeros(ref_dat.shape[1])
			fd_return = fd_return[:, np.newaxis].T
			dat = ref_dat.values[1:, :] / ref_dat.values[:-1, :] - 1
			dat = np.concatenate([fd_return, dat])
			dat = DataFrame(dat, index=ref_dat.index, columns=ref_dat.columns)
		elif len(ref_dat.shape) == 1:
			# input is a Series
			fd_return = np.zeros(1)
			dat = ref_dat.values[1:] / ref_dat.values[:-1] - 1
			dat = np.concatenate([fd_return, dat])
			dat = Series(dat, index=ref_dat.index, name=ref_dat.name)
	elif type == 2:
		# return to netvalue
		dat = (ref_dat + 1).cumprod()

	return dat



# 计算一下相对收益 / 超额收益
def Get_ExcessReturn(ret, refRet):
	excessRet = ret - refRet
	excessNetValue = (1 + excessRet).cumprod()

	IR = excessRet.mean() / excessRet.std()
	res = {'excessRet': excessRet, 'excessNetValue': excessNetValue, 'IR': IR}
	return res


# ============================================ strategyAnalysis ===================================================== #
class strategyAnalysis():

	def __init__(self):
		logging.info('Analyze Strategy Performance')
		self.cycle_fac = 244

	def Get_BasicIndictors(self, ret, rf_rate=0):
		# ret can be Series or dataframe with date index
		self.rf_rate = rf_rate
		self.ret = DataFrame(ret) if isinstance(ret, Series) else ret
		self.numofTD = ret.shape[0]

		# 1. nav
		self.nav = (ret + 1).cumprod()

		# 2. total return
		self.tot_return = self.nav.iloc[-1, :] - 1

		# 3. annulized return
		self.ann_return = (self.tot_return + 1) ** (self.cycle_fac / self.numofTD) - 1

		# 4. annulized volatility
		self.ann_vol = ret.std() * np.sqrt(self.cycle_fac)

		# 5. sharpe ratio
		self.sharpe_ratio = self.ann_return / self.ann_vol

		# 6. downdown
		drawdown = DataFrame(np.zeros(self.nav.shape), index=self.nav.index, columns=self.nav.columns)

		for i in range(0, self.nav.shape[0]):
			drawdown[i:(i + 1)] = self.nav[i:(i + 1)] / self.nav[:(i + 1)].max() - 1
		self.drawdown = drawdown
		self.maxdrawdown = np.abs(self.drawdown.min())

		# 7. calmar ratio
		self.calmar_ratio = self.ann_return / np.abs(self.maxdrawdown)

		res = pd.concat([self.tot_return, self.ann_return, self.ann_vol, self.sharpe_ratio, self.maxdrawdown, self.calmar_ratio], axis=1, sort=True)
		res = round(res.T, 3)
		res.index = ['tot_return', 'ann_return', 'ann_vol', 'sharpe_ratio', 'maxdrawdown', 'calmar_ratio']

		return res

	def Plot(self, title='', sel_ret_plot=0):
		# plot strategy performance

		XTicks = Series(np.linspace(0, self.nav.shape[0] - 1, 8)).astype('int').tolist()
		plt.figure()

		# 1. plot nav
		plt.subplot(311)
		colnames = self.nav.columns.tolist()
		for i in range(0, self.nav.shape[1]):
			plt.plot(range(0, self.numofTD), self.nav[colnames[i]])

		plt.title(title)
		plt.xticks(XTicks, [''] * len(XTicks))
		plt.legend(colnames)
		plt.ylabel('NetValue')

		# 2. plot daily return
		plt.subplot(312)
		plt.bar(range(0, self.numofTD), self.ret.iloc[:, sel_ret_plot])
		plt.xticks(XTicks, [''] * len(XTicks))
		plt.ylabel('Daily Return')

		# 3. plot drawdown
		plt.subplot(313)
		colnames = list(self.drawdown.columns)
		for i in range(0, self.drawdown.shape[1]):
			plt.plot(range(0, self.numofTD), self.drawdown[colnames[i]])

		plt.xticks(XTicks, list(self.drawdown.index[XTicks]))
		plt.legend(colnames)
		plt.ylabel('Drawdown')

		plt.show()


# ============================================ Funs_UpdateRawData ================================================= #
# ------- Configure Local SQLServer ------- #
# Pydb = Pyodbc(Server="10.2.224.104", DataBase="newwind", uid="ax_datacenter", pwd="DataCenter_0727")
Pydb = Pyodbc(Server="172.18.1.54", DataBase="Wind", Certificate='Windows')


#  每个月全量更新一次交易日历
# 1. 交易日历 -- 全量更新
def update_AShareCalendar(config):

	fun_name = 'AShareCalendar'

	tmp_sql = 'SELECT TRADE_DAYS, S_INFO_EXCHMARKET ' \
				'FROM ASHARECALENDAR ' \
				'WHERE S_INFO_EXCHMARKET = \'SSE\' ' \
				'ORDER BY TRADE_DAYS'

	colnames = ['Date', 'StockExchange']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	return dat

### 后期把一些全局变量放进来，不过一定要注意，别被更改了！！！

# 2. EOD PRICE  -- 增量更新
def update_AShareEODPrices(config):

	# fun_name = config['fun_name']
	fun_name = 'AShareEODPrices'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.TRADE_DT, t.S_DQ_OPEN, t.S_DQ_CLOSE, t.S_DQ_AVGPRICE, t.S_DQ_ADJCLOSE, ' \
				   't.S_DQ_PCTCHANGE, t.S_DQ_VOLUME, t.S_DQ_AMOUNT, t.S_DQ_ADJFACTOR, t.S_DQ_TRADESTATUS ' \
				   'FROM dbo.ASHAREEODPRICES t ' \
				   'WHERE t.TRADE_DT <= ' + end_dt + ' AND t.TRADE_DT >= ' + start_dt + ' ' \
				   'ORDER BY t.TRADE_DT,t.S_INFO_WINDCODE'

	colnames = ['Symbol', 'Date', 'Open', 'Close', 'Vwap', 'AdjClose', 'PctChg', 'Volume', 'Amount', 'AdjFactor', 'TradeStatus_Type']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['Open', 'Close', 'Vwap', 'AdjClose', 'PctChg', 'Volume', 'Amount', 'AdjFactor']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	# 调整单位
	dat['Volume'] = dat['Volume'] * 1e2
	dat['Amount'] = dat['Amount'] * 1e3

	if len(dat['Date'].unique()) != Series(periodTDates).shape[0]:
		logging.warning('Trading Dates of Obtained DataFrame cannot match stockCalendar in ' + fun_name + '.' +
						' Latest Date = ' + dat['Date'].unique().max())
		#  这里要晚上更新，不然当天的数据应该还没有来(收盘一段时间后才入库)

	dat = dat.ix[:, ['Symbol', 'Date', 'Open', 'Close', 'Vwap', 'AdjClose', 'PctChg', 'Volume', 'Amount', 'AdjFactor', 'TradeStatus_Type']]
	return dat


# 3. ST  -- 全量更新
def update_AShareST(config):

	fun_name = 'AShareST'

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.S_TYPE_ST, t.ENTRY_DT, t.REMOVE_DT ' \
			'FROM dbo.ASHAREST t ' \
			'ORDER BY t.S_INFO_WINDCODE'
	dat = Pydb.ExecQuery(tmp_sql, ['Symbol', 'TypeST', 'EntryDate', 'RemoveDate'])

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()
	# 要处理一下结束时间

	return dat


# 4. IPO Date -- 全量更新
def update_AShareDescription(config):

	fun_name = 'AShareDesciption'

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.S_INFO_LISTDATE ' \
			'FROM dbo.ASHAREDESCRIPTION t ' \
			'WHERE t.S_INFO_WINDCODE LIKE \'[0-9]%\' ' \
			'ORDER BY t.S_INFO_WINDCODE'
	dat = Pydb.ExecQuery(tmp_sql, ['Symbol', 'IPODate'])

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	return dat





## 资金流向表 ASHAREMONEYFLOW ##
def update_AShareMoneyFlow(config):

	fun_name = 'AShareMoneyFlow'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	tmp_sql = '' \
				'SELECT S_INFO_WINDCODE, TRADE_DT, VALUE_DIFF_INSTITUTE, VALUE_DIFF_LARGE_TRADER, VALUE_DIFF_MED_TRADER, ' \
				'VALUE_DIFF_SMALL_TRADER, S_MFD_INFLOW_OPEN, S_MFD_INFLOW_CLOSE, S_MFD_INFLOW ' \
				'FROM dbo.ASHAREMONEYFLOW ' \
				'WHERE TRADE_DT <= ' + end_dt + ' AND TRADE_DT >= ' + start_dt + ' ' \
				'ORDER BY TRADE_DT, S_INFO_WINDCODE '

	colnames = ['Symbol', 'Date', 'MFInst', 'MFLarge', 'MFMed', 'MFSmall', 'MFOpen', 'MFClose', 'MFTot']
	#          ['机构金额差', '大户金额差','中户金额差' , '散户金额差', '开盘资金流入金额', '收盘资金流入金额', '净流入金额']

	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	dat = dat[dat['Date'].isin(periodTDates)]

	col_trans = ['MFInst', 'MFLarge', 'MFMed', 'MFSmall', 'MFOpen', 'MFClose', 'MFTot']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	# 调整单位
	dat.ix[:, col_trans] = dat.ix[:, col_trans] * 1e4

	if len(dat['Date'].unique()) != Series(periodTDates).shape[0]:
		logging.warning('Trading Dates of Obtained DataFrame cannot match stockCalendar in ' + fun_name + '.' +
						' Latest Date = ' + dat['Date'].unique().max())
	return dat



# MoneyFlow.shape   (916371, 6)
# eodPrice.shape  (985638, 9)
# MoneyFlow比eodPrice记录要少, MoneyFlow['Symbol'].unique().shape --> (3500,); eodPrice['Symbol'].unique().shape --> (3505,)
# 股票数就少一些，要注意，可能有些股票没有该项纪录
# 一般来说，没有的股票是在该阶段退市了，停牌很长时间了，*ST了，这些应该没有影响
# 比如{'600485.SH', '600145.SH', '002075.SZ', '000748.SZ', '000029.SZ'}



## 一致预期数据 ##
# 一直预期的数据太大，把它分成几个元素，后面 #
def update_AShareConsensusRollingData(config):

	fun_name = 'AShareConsensusRollingData'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	# periodTDates = config['periodTDates']

	tmp_sql = '' \
				'SELECT S_INFO_WINDCODE, EST_DT, ROLLING_TYPE, NET_PROFIT, EST_EPS, EST_PEG, EST_ROE ' \
				'FROM dbo.ASHARECONSENSUSROLLINGDATA ' \
				'WHERE EST_DT <= ' + end_dt + ' AND EST_DT >= ' + start_dt + ' ' \
				'ORDER BY EST_DT, S_INFO_WINDCODE '

	colnames = ['Symbol', 'EST_Date', 'Type', 'NetProfit', 'EST_EPS', 'EST_PEG', 'EST_ROE']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['NetProfit', 'EST_EPS', 'EST_PEG', 'EST_ROE']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	return dat


# AshareDividend 分红数据
def update_AShareDividend(config):

	fun_name = 'AShareDividend'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	# periodTDates = config['periodTDates']

	tmp_sql = 'SELECT S_INFO_WINDCODE, EX_DT, STK_DVD_PER_SH, CASH_DVD_PER_SH_AFTER_TAX ' \
				'FROM dbo.ASHAREDIVIDEND ' \
				'WHERE S_DIV_PROGRESS = 3 AND EX_DT <= ' + end_dt + ' AND EX_DT >= ' + start_dt + ' ' \
				'ORDER BY EX_DT, S_INFO_WINDCODE'

	colnames = ['Symbol', 'Date', 'StkRatio', 'CashRatio']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['StkRatio', 'CashRatio']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	return dat



# AIndexMembersCITICS 中信一级行业 - 全量更新
def update_AIndexMembersCITICS(config):

	fun_name = 'AIndexMembersCITICS'

	tmp_sql = 'SELECT S_CON_WINDCODE, S_INFO_WINDCODE, S_CON_INDATE, S_CON_OUTDATE, CUR_SIGN ' \
				'FROM dbo.AINDEXMEMBERSCITICS ' \
				'ORDER BY S_CON_WINDCODE, S_INFO_WINDCODE'

	colnames = ['Symbol', 'IndexCode', 'InDate', 'OutDate', 'Sign']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	return dat


# 2. AIndexIndustriesEODCITICS  -- 增量更新
def update_AIndexIndustriesEODCITICS(config):

	fun_name = 'AIndexIndustriesEODCITICS'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.TRADE_DT, t.S_DQ_PRECLOSE, t.S_DQ_OPEN, t.S_DQ_CLOSE, ' \
				   't.S_DQ_PCTCHANGE, t.S_DQ_VOLUME, t.S_DQ_AMOUNT ' \
				   'FROM dbo.AINDEXINDUSTRIESEODCITICS t ' \
				   'WHERE t.TRADE_DT <= ' + end_dt + ' AND t.TRADE_DT >= ' + start_dt + ' ' \
				   'ORDER BY t.TRADE_DT, t.S_INFO_WINDCODE'

	colnames = ['Symbol', 'Date', 'PreClose', 'Open', 'Close', 'PctChg', 'Volume', 'Amount']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['PreClose', 'Open', 'Close', 'PctChg', 'Volume', 'Amount']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	# 调整单位
	dat['Volume'] = dat['Volume'] * 1e2
	dat['Amount'] = dat['Amount'] * 1e3

	if len(dat['Date'].unique()) != Series(periodTDates).shape[0]:
		logging.warning('Trading Dates of Obtained DataFrame cannot match stockCalendar in ' + fun_name + '.' +
						' Latest Date = ' + dat['Date'].unique().max())

	return dat



# AIndexMembers A股指数成分股 - 全量更新
def update_AIndexMembers(config):

	fun_name = 'AIndexMembers'

	# 上证综指（000001.SH），深证成指（399001.SZ），上证50（000016.SH），沪深300（000300.SH），中证500（000905.SH）
	# 中证800（000906.SH）， 中证1000（000852.SH），中小板指（399005.SZ），创业板（399006.SZ）
	indexSet = ['\'000001.SH\'', '\'399001.SZ\'', '\'000016.SH\'', '\'000300.SH\'', '\'000905.SH\'',
				'\'000906.SH\'', '\'000852.SH\'', '\'399005.SZ\'', '\'399006.SZ\'']

	tmp_sql = 'SELECT S_CON_WINDCODE, S_INFO_WINDCODE, S_CON_INDATE, S_CON_OUTDATE, CUR_SIGN ' \
				'FROM dbo.AINDEXMEMBERS ' \
				'WHERE S_INFO_WINDCODE IN (' + ','.join(indexSet) + ') ' \
				'ORDER BY S_CON_WINDCODE, S_INFO_WINDCODE'

	colnames = ['Symbol', 'IndexCode', 'InDate', 'OutDate', 'Sign']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	return dat


# AIndexEODPrices - 增量更新
def update_AIndexEODPrices(config):

	fun_name = 'AIndexEODPrices'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	# 上证综指（000001.SH），深证成指（399001.SZ），上证50（000016.SH），沪深300（000300.SH），中证500（000905.SH）
	# 中证800（000906.SH）， 中证1000（000852.SH），中小板指（399005.SZ），创业板（399006.SZ）
	indexSet = ['\'000001.SH\'', '\'399001.SZ\'', '\'000016.SH\'', '\'000300.SH\'', '\'000905.SH\'',
				'\'000906.SH\'', '\'000852.SH\'', '\'399005.SZ\'', '\'399006.SZ\'']

	tmp_sql = '' \
				'SELECT t.S_INFO_WINDCODE, t.TRADE_DT, t.S_DQ_PRECLOSE, t.S_DQ_OPEN, ' \
				't.S_DQ_CLOSE, t.S_DQ_PCTCHANGE, t.S_DQ_VOLUME, t.S_DQ_AMOUNT  ' \
				'FROM dbo.AINDEXEODPRICES t ' \
				'WHERE t.TRADE_DT <= ' + end_dt + ' AND t.TRADE_DT >= ' + start_dt + ' AND t.S_INFO_WINDCODE IN (' + ','.join(indexSet) + ') ' \
				'ORDER BY t.TRADE_DT, t.S_INFO_WINDCODE '

	colnames = ['Symbol', 'Date', 'PreClose', 'Open', 'Close', 'PctChg', 'Volume', 'Amount']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	if len(dat['Date'].unique()) != Series(periodTDates).shape[0]:
		logging.warning('Trading Dates of Obtained DataFrame cannot match stockCalendar in ' + fun_name + '.' +
						' Latest Date = ' + dat['Date'].unique().max())

	col_trans = ['PreClose', 'Open', 'Close', 'PctChg', 'Volume', 'Amount']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	# 调整单位
	dat['Volume'] = dat['Volume'] * 1e2
	dat['Amount'] = dat['Amount'] * 1e3

	return dat



# AIndexMembersWIND Wind指数成分股 - 全量更新
def update_AIndexMembersWIND(config):

	fun_name = 'AIndexMembersWIND'

	tmp_sql = 'SELECT S_CON_WINDCODE, F_INFO_WINDCODE, S_CON_INDATE, S_CON_OUTDATE, CUR_SIGN ' \
				'FROM dbo.AINDEXMEMBERSWIND ' \
				'WHERE F_INFO_WINDCODE = \'881001.WI\' ' \
				'ORDER BY S_CON_WINDCODE'

	colnames = ['Symbol', 'IndexCode', 'InDate', 'OutDate', 'Sign']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	return dat




# 2. AIndexWindIndustriesEOD  -- 增量更新
def update_AIndexWindIndustriesEOD(config):

	fun_name = 'AIndexWindIndustriesEOD'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.TRADE_DT, t.S_DQ_PRECLOSE, t.S_DQ_OPEN, t.S_DQ_CLOSE, ' \
				   't.S_DQ_PCTCHANGE, t.S_DQ_VOLUME, t.S_DQ_AMOUNT ' \
				   'FROM dbo.AINDEXWINDINDUSTRIESEOD t ' \
				   'WHERE t.TRADE_DT <= ' + end_dt + ' AND t.TRADE_DT >= ' + start_dt + ' AND t.S_INFO_WINDCODE = \'881001.WI\' ' \
				   'ORDER BY t.TRADE_DT, t.S_INFO_WINDCODE'

	colnames = ['Symbol', 'Date', 'PreClose', 'Open', 'Close', 'PctChg', 'Volume', 'Amount']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['PreClose', 'Open', 'Close', 'PctChg', 'Volume', 'Amount']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	# 调整单位
	dat['Volume'] = dat['Volume'] * 1e2
	dat['Amount'] = dat['Amount'] * 1e3

	if len(dat['Date'].unique()) != Series(periodTDates).shape[0]:
		logging.warning('Trading Dates of Obtained DataFrame cannot match stockCalendar in ' + fun_name + '.' +
						' Latest Date = ' + dat['Date'].unique().max())

	return dat



# 2. AIndexHS300FreeWeight  -- 增量更新
def update_AIndexHS300FreeWeight(config):

	fun_name = 'AIndexHS300FreeWeight'
	start_dt = config['start_dt']
	end_dt = config['end_dt']

	# 上证综指（000001.SH），深证成指（399001.SZ），上证50（000016.SH），沪深300（000300.SH / 399300.SZ），中证500（000905.SH）
	# 中证800（000906.SH）， 中证1000（000852.SH），中小板指（399005.SZ），创业板（399006.SZ）
	indexSet = ['\'000001.SH\'', '\'399001.SZ\'', '\'000016.SH\'', '\'399300.SZ\'', '\'000905.SH\'',
				'\'000906.SH\'', '\'000852.SH\'', '\'399005.SZ\'', '\'399006.SZ\'']

	tmp_sql = '' \
			  	'SELECT t.S_CON_WINDCODE, t.TRADE_DT, t.S_INFO_WINDCODE, t.I_WEIGHT ' \
			  	'FROM dbo.AINDEXHS300FREEWEIGHT t ' \
			  	'WHERE t.TRADE_DT <= ' + end_dt + ' AND t.TRADE_DT >= ' + start_dt + ' AND t.S_INFO_WINDCODE IN (' + ','.join(indexSet) + ') ' \
				'ORDER BY t.TRADE_DT, t.S_INFO_WINDCODE, t.S_CON_WINDCODE '

	# 数据库中，有些指数可能是从2010开始，从20141030开始（中证1000）

	colnames = ['Symbol', 'Date', 'IndexCode', 'Weight']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['Weight']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	# 为了统一，把399300统一更换成000300.SH
	dat.ix[dat['IndexCode'] == '399300.SZ', 'IndexCode'] = '000300.SH'

	return dat


# AShareBalanceSheet 资产负债表 - 增量更新
def update_AShareBalanceSheet(config):

	fun_name = 'AShareBalanceSheet'
	start_dt = config['start_dt']
	end_dt = config['end_dt']

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.ANN_DT, t.REPORT_PERIOD, t.TOT_CUR_ASSETS, t.FIX_ASSETS, t.TOT_NON_CUR_ASSETS, ' \
				   't.TOT_CUR_LIAB, t.TOT_LIAB, t.TOT_SHRHLDR_EQY_EXCL_MIN_INT ' \
				   'FROM dbo.ASHAREBALANCESHEET t ' \
				   'WHERE t.ANN_DT <= ' + end_dt + ' AND t.ANN_DT >= ' + start_dt + ' AND t.STATEMENT_TYPE = 408001000 ' \
				   'ORDER BY t.ANN_DT, t.S_INFO_WINDCODE, t.REPORT_PERIOD'

	colnames = ['Symbol', 'Date', 'ReportPeriod', 'TotCurAssets', 'FixAssets', 'TotNonCurAssets', 'TotCurLiab', 'TotLiab', 'TotShrHldrExclMin']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['TotCurAssets', 'FixAssets', 'TotNonCurAssets', 'TotCurLiab', 'TotLiab', 'TotShrHldrExclMin']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')


	return dat



# AShareCashFlow 现金流量表 - 增量更新
def update_AShareCashFlow(config):

	fun_name = 'AShareCashFlow'
	start_dt = config['start_dt']
	end_dt = config['end_dt']

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.ANN_DT, t.REPORT_PERIOD, t.CASH_CASH_EQU_BEG_PERIOD, t.CASH_CASH_EQU_END_PERIOD ' \
				   'FROM dbo.ASHARECASHFLOW t ' \
				   'WHERE t.ANN_DT <= ' + end_dt + ' AND t.ANN_DT >= ' + start_dt + ' AND t.STATEMENT_TYPE = 408001000 ' \
				   'ORDER BY t.ANN_DT, t.S_INFO_WINDCODE, t.REPORT_PERIOD'

	colnames = ['Symbol', 'Date', 'ReportPeriod', 'CashCashEquBeg', 'CashCashEquEnd']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['CashCashEquBeg', 'CashCashEquEnd']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	return dat



# AShareTTMHis TTM指标历史数据 - 增量更新
def update_AShareTTMHis(config):

	fun_name = 'AShareTTMHis'
	start_dt = config['start_dt']
	end_dt = config['end_dt']

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.ANN_DT, t.REPORT_PERIOD, t.TOT_OPER_REV_TTM, t.OPER_REV_TTM, t.NET_PROFIT_TTM,' \
			  	't.NET_PROFIT_PARENT_COMP_TTM, t.NET_INCR_CASH_CASH_EQU_TTM, t.NET_CASH_FLOWS_OPER_ACT_TTM, t.S_FA_OP_TTM, ' \
			  	't.S_FA_EBT_TTM, t.S_FA_EBIT_TTM  ' \
				'FROM dbo.ASHARETTMHIS t ' \
				'WHERE t.ANN_DT <= ' + end_dt + ' AND t.ANN_DT >= ' + start_dt + ' AND t.STATEMENT_TYPE = \'合并报表\' ' \
				'ORDER BY t.ANN_DT, t.S_INFO_WINDCODE, t.REPORT_PERIOD'

	colnames = ['Symbol', 'Date', 'ReportPeriod', 'TotOperRev', 'OperRev', 'NetProfit', 'NetProfitParent', 'NetIncrCash',
				'NetCashFlowsOper', 'FaOp', 'FaEBT', 'FaEBIT']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['TotOperRev', 'OperRev', 'NetProfit', 'NetProfitParent', 'NetIncrCash',
				'NetCashFlowsOper', 'FaOp', 'FaEBT', 'FaEBIT']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64')

	return dat


# AShareCapitalization  股票股本 - 增量更新
def update_AShareCapitalization(config):

	fun_name = 'AShareCapitalization'

	tmp_sql = 'SELECT t.S_INFO_WINDCODE, t.CHANGE_DT, t.TOT_SHR, t.FLOAT_SHR, t.FLOAT_A_SHR ' \
				   'FROM dbo.ASHARECAPITALIZATION t ' \
				   'ORDER BY t.CHANGE_DT, t.S_INFO_WINDCODE '

	colnames = ['Symbol', 'Date', 'TotShr', 'FloatShr', 'FloatAShr']
	dat = Pydb.ExecQuery(tmp_sql, colnames)

	if dat.empty:
		logging.warning('Obtained DataFrame is empty in ' + fun_name + '.')
		return DataFrame()

	col_trans = ['TotShr', 'FloatShr', 'FloatAShr']
	dat.ix[:, col_trans] = dat.ix[:, col_trans].astype('float64') * 1e4

	return dat




# ============================================ Funs_UpdateDerData ================================================= #


def update_StockST(config):


	fun_name = 'stockST'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	REF_DF = get_CoreGenData('REF_DF')
	dat = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()
	dat['ST'] = 0

	AShareST = dataAssign('AShareST')

	# handle case when REMOVE_DT = None
	tempST = AShareST.copy()   # 保留原始的ST数据
	tempST.ix[tempST['RemoveDate'].isnull(), 'RemoveDate'] = (datetime.today() + timedelta(days=1)).strftime('%Y%m%d')

	for tradDay in list(dat['Date'].unique()):
		temp_stk = dat[dat['Date'] == tradDay]
		temp = pd.merge(temp_stk, tempST, on='Symbol')

		temp.ix[(temp['Date'] >= temp['EntryDate']) & (temp['Date'] < temp['RemoveDate']), 'ST'] = 1
		temp = temp[temp['ST'] == 1]
		dat.ix[(dat['Date'] == tradDay) & dat['Symbol'].isin(list(temp['Symbol'])), 'ST'] = 1

	return dat


def update_StockIPO(config):

	fun_name = 'stockIPO'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	REF_DF = get_CoreGenData('REF_DF')
	dat = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()
	dat['IPO'] = 0

	AShareDescription = dataAssign('AShareDescription')

	for tradDay in list(dat['Date'].unique()):
		temp_IPO = AShareDescription[AShareDescription['IPODate'] == tradDay]
		dat.ix[(dat['Date'] == tradDay) & dat['Symbol'].isin(list(temp_IPO['Symbol'])), 'IPO'] = 1

	return dat


def update_StockSuspension(config):

	fun_name = 'stockSuspension'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']
	# rawData = config['rawData']

	AShareEODPrices = dataAssign('AShareEODPrices')

	REF_DF = get_CoreGenData('REF_DF')
	dat = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()
	dat['Suspension'] = 0

	rawData = AShareEODPrices.ix[(AShareEODPrices['Date'] >= start_dt) & (AShareEODPrices['Date'] <= end_dt), :].copy()
	# 1. handle suspension
	dat.ix[rawData['TradeStatus_Type'] == '停牌', 'Suspension'] = 1

	return dat


def update_StockLimitUpDown(config):

	fun_name = 'stockLimitUpDown'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	AShareEODPrices = dataAssign('AShareEODPrices')
	StockST = dataAssign('StockST')

	REF_DF = get_CoreGenData('REF_DF')
	dat = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()
	dat['LimitUpDown'] = 0

	rawData = AShareEODPrices.ix[(AShareEODPrices['Date'] >= start_dt) & (AShareEODPrices['Date'] <= end_dt), :].copy()
	tmpST = StockST.ix[(StockST['Date'] >= start_dt) & (StockST['Date'] <= end_dt), :].copy()

	# handle limit up and down
	dat.ix[rawData['PctChg'] > 9.7, 'LimitUpDown'] = 1
	dat.ix[rawData['PctChg'] < -9.7, 'LimitUpDown'] = -1

	# ST股的上下限是5%
	dat.ix[(tmpST['ST'] == 1) & (rawData['PctChg'] > 4.8), 'LimitUpDown'] = 1
	dat.ix[(tmpST['ST'] == 1) & (rawData['PctChg'] < -4.8), 'LimitUpDown'] = -1

	return dat



def update_StockTradeStatus(config):

	fun_name = 'stockTradeStatus'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	StockST, StockIPO, StockSuspension, StockLimitUpDown = \
		dataAssign(['StockST', 'StockIPO', 'StockSuspension', 'StockLimitUpDown'])

	dat = pd.merge(StockST, StockIPO, on=['Symbol', 'Date'])
	dat = pd.merge(dat, StockSuspension, on=['Symbol', 'Date'])
	dat = pd.merge(dat, StockLimitUpDown, on=['Symbol', 'Date'])

	# Final trade status
	dat['TradeStatus'] = \
		(1 - dat['Suspension']) * (1 - abs(dat['LimitUpDown'])) * (1 - dat['ST']) * (1 - dat['IPO'])

	dat = dat.ix[(dat['Date'] >= start_dt) & (dat['Date'] <= end_dt), :].copy()

	return dat


def update_CapitalizationEOD(config):

	fun_name = 'CapitalizationEOD'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	AShareCapitalization = dataAssign('AShareCapitalization')

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:
		tmp = AShareCapitalization.ix[AShareCapitalization['Date'] <= td, :]
		tmp = tmp.sort_values(by=['Symbol', 'Date'])
		tmp = tmp.drop_duplicates(['Symbol'], keep='last')

		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp.ix[:, ['Symbol', 'TotShr', 'FloatShr', 'FloatAShr']], on=['Symbol'])
		dat = pd.concat([dat, tmp], sort=True)

	return dat


# 简单清洗BalanceSheet，只在财报节点结束后更新财报
def update_BalanceSheetRevised(config):

	fun_name = 'BalanceSheetRevised'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	AShareDescription = dataAssign('AShareDescription')
	AShareBalanceSheet = dataAssign('AShareBalanceSheet')

	# 把在周末的都往后移到下个交易日
	AShareBalanceSheet.ix[:, 'Date'] = shiftWeekend(AShareBalanceSheet['Date'])

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	reportEndDates = ['0501', '0901', '1101']

	# 能并行吗？
	for td in tradeDates:
		tmp = AShareBalanceSheet.ix[AShareBalanceSheet['Date'] <= td, :]

		# 删除上市不足2年的股票
		tmp_ref = ref.ix[ref['Date'] == td, :]
		tmp_ref = pd.merge(tmp_ref, AShareDescription, on=['Symbol'])
		tmp_ref['2Y'] = tmp_ref['IPODate'].apply(
			lambda x: (datetime.strptime(x, '%Y%m%d') + timedelta(days=730)).strftime('%Y%m%d'))

		tmpStocks = tmp_ref.ix[tmp_ref['Date'] > tmp_ref['2Y'], 'Symbol'].tolist()
		tmp = tmp.ix[tmp['Symbol'].isin(tmpStocks), :]
		tmp = tmp.sort_values(by=['Symbol', 'Date'])

		if td[4:] < reportEndDates[0]:
			# 用去年3季度报

			# 提取所有股票的去年3季度表
			tmp = tmp.ix[tmp['ReportPeriod'] == str(int(td[:4]) - 1) + '0930', :]
			del tmp['Date']
			del tmp['ReportPeriod']

		elif td[4:] >= reportEndDates[2]:

			# 提取所有股票的今年3季度表
			tmp = tmp.ix[tmp['ReportPeriod'] == td[:4] + '0930', :]
			del tmp['Date']
			del tmp['ReportPeriod']


		elif (td[4:] >= reportEndDates[0]) and (td[4:] < reportEndDates[1]):
			# 用去年年报0.8和今年一季度报0.2合成

			# 提取所有股票的去年年报和今年一季度报
			tmp1 = tmp.ix[tmp['ReportPeriod'] == td[:4] + '0331', :]
			tmp1.index = list(tmp1['Symbol'])

			tmp2 = tmp.ix[tmp['ReportPeriod'] == str(int(td[:4]) - 1) + '1231', :]
			tmp2.index = list(tmp2['Symbol'])

			tmp = tmp1.ix[:, 3:] * 0.8 + tmp2.ix[:, 3:] * 0.2
			tmp.insert(0, 'Symbol', tmp.index.tolist())


		elif (td[4:] >= reportEndDates[1]) and (td[4:] < reportEndDates[2]):
			# 用今年中年报

			# 提取所有股票的今年中季度表
			tmp = tmp.ix[tmp['ReportPeriod'] == td[:4] + '0630', :]
			del tmp['Date']
			del tmp['ReportPeriod']

		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp, on=['Symbol'], how='left')

		# 去重
		# tmp = tmp.drop_duplicates()
		dat = pd.concat([dat, tmp], sort=True)

	return dat

# 把财报在周末发的，顺移到最近的下一个交易日
def shiftWeekend(dat):
	# dat is a series
	tradingDates = get_CoreGenData('AShareCalendar')['Date']
	tmp = dat.ix[~dat.isin(tradingDates.tolist())]

	for weekend in list(tmp.unique()):
		replaceDate = tradingDates[tradingDates > weekend].min()  # 顺移到下一个交易日
		dat.ix[dat == weekend] = replaceDate

	return dat



# 简单清洗TTMHis，只在财报节点结束后更新财报
def update_TTMHisRevised(config):

	fun_name = 'TTMHisRevised'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	AShareDescription = dataAssign('AShareDescription')
	AShareTTMHis = dataAssign('AShareTTMHis')

	# 把在周末的都往后移到下个交易日
	AShareTTMHis.ix[:, 'Date'] = shiftWeekend(AShareTTMHis['Date'])

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	reportEndDates = ['0501', '0901', '1101']

	for td in tradeDates:
		tmp = AShareTTMHis.ix[AShareTTMHis['Date'] <= td, :]

		# 删除上市不足2年的股票
		tmp_ref = ref.ix[ref['Date'] == td, :]
		tmp_ref = pd.merge(tmp_ref, AShareDescription, on=['Symbol'])
		tmp_ref['2Y'] = tmp_ref['IPODate'].apply(
			lambda x: (datetime.strptime(x, '%Y%m%d') + timedelta(days=730)).strftime('%Y%m%d'))

		tmpStocks = tmp_ref.ix[tmp_ref['Date'] > tmp_ref['2Y'], 'Symbol'].tolist()
		tmp = tmp.ix[tmp['Symbol'].isin(tmpStocks), :]
		tmp = tmp.sort_values(by=['Symbol', 'Date'])

		if td[4:] < reportEndDates[0]:
			# 用去年3季度报，# 提取所有股票的去年3季度表
			tmp = tmp.ix[tmp['ReportPeriod'] == str(int(td[:4]) - 1) + '0930', :]

		elif td[4:] >= reportEndDates[2]:
			# 当年3季度报
			tmp = tmp.ix[tmp['ReportPeriod'] == td[:4] + '0930', :]

		elif (td[4:] >= reportEndDates[0]) and (td[4:] < reportEndDates[1]):
			# 一季度报
			tmp = tmp.ix[tmp['ReportPeriod'] == td[:4] + '0331', :]

		elif (td[4:] >= reportEndDates[1]) and (td[4:] < reportEndDates[2]):
			# 用今年中年报， # 提取所有股票的今年中季度表
			tmp = tmp.ix[tmp['ReportPeriod'] == td[:4] + '0630', :]

		if tmp.empty:
			continue

		del tmp['Date']
		del tmp['ReportPeriod']

		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp, on=['Symbol'], how='left')
		dat = pd.concat([dat, tmp], sort=True)

	return dat




# 中信一级行业分组
def update_IndustryCITICS(config):

	fun_name = 'IndustryCITICS'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	AIndexMembersCITICS = dataAssign('AIndexMembersCITICS')
	tmp_dat = AIndexMembersCITICS.fillna(value=datetime.today().strftime('%Y%m%d'))

	# 找到指数最早的记录
	firstRecordDate = tmp_dat['InDate'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[(tmp_dat['InDate'] <= td) & (tmp_dat['OutDate'] >= td), :]
		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp, how='left', on=['Symbol'])
		tmp = tmp.fillna(value='Others')
		tmp = tmp.ix[:, ['Symbol', 'Date', 'IndexCode']]
		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.sort_values(by=['Date', 'Symbol'])
	return dat





# 上证50指数成分
def update_SZ50Pool(config):

	fun_name = 'SZ50Pool'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	index_code = '000016.SH'
	AIndexMembers = get_CoreGenData('AIndexMembers')
	IndexMembers = AIndexMembers.ix[AIndexMembers['IndexCode'] == index_code, :]
	tmp_dat = IndexMembers.fillna(value=datetime.today().strftime('%Y%m%d'))


	# 找到指数最早的记录
	firstRecordDate = tmp_dat['InDate'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[(tmp_dat['InDate'] <= td) & (tmp_dat['OutDate'] >= td), :]
		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp, how='left', on=['Symbol'])
		tmp = tmp.fillna(value=0)
		tmp.ix[tmp['IndexCode'] == index_code, 'IndexCode'] = 1

		tmp = tmp.ix[:, ['Symbol', 'Date', 'IndexCode']]
		tmp.columns = ['Symbol', 'Date', 'isMember']

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.sort_values(by=['Date', 'Symbol'])
	return dat


# 上证50指数日权重
def update_SZ50DailyWeight(config):

	fun_name = 'SZ50DailyWeight'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	index_code = '000016.SH'
	AIndexHS300FreeWeight = get_CoreGenData('AIndexHS300FreeWeight')
	tmp_dat = AIndexHS300FreeWeight.ix[AIndexHS300FreeWeight['IndexCode'] == index_code, :]

	if np.any(tmp_dat.isnull()):
		logging.error('Index: ' + index_code + ', Weights Exist NaN ! ')
		return DataFrame()


	# 找到指数最早的记录
	firstRecordDate = tmp_dat['Date'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[tmp_dat['Date'] <= td, :]
		if tmp.empty:
			logging.warning('Index: ' + index_code + ', Weights Missing Before ' + td)
			continue
		tmp = tmp.ix[tmp['Date'] == tmp['Date'].max(), :]
		tmp['Date'] = td

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.ix[:, ['Symbol', 'Date', 'Weight']]
	dat = dat.sort_values(by=['Date', 'Symbol'])

	return dat



# HS300指数成分
def update_HS300Pool(config):

	fun_name = 'HS300Pool'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	index_code = '000300.SH'
	AIndexMembers = get_CoreGenData('AIndexMembers')
	IndexMembers = AIndexMembers.ix[AIndexMembers['IndexCode'] == index_code, :]
	tmp_dat = IndexMembers.fillna(value=datetime.today().strftime('%Y%m%d'))


	# 找到指数最早的记录
	firstRecordDate = tmp_dat['InDate'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[(tmp_dat['InDate'] <= td) & (tmp_dat['OutDate'] >= td), :]
		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp, how='left', on=['Symbol'])
		tmp = tmp.fillna(value=0)
		tmp.ix[tmp['IndexCode'] == index_code, 'IndexCode'] = 1

		tmp = tmp.ix[:, ['Symbol', 'Date', 'IndexCode']]
		tmp.columns = ['Symbol', 'Date', 'isMember']

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.sort_values(by=['Date', 'Symbol'])
	return dat



# HS300指数日权重
def update_HS300DailyWeight(config):

	fun_name = 'HS300DailyWeight'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	index_code = '000300.SH'
	AIndexHS300FreeWeight = get_CoreGenData('AIndexHS300FreeWeight')
	tmp_dat = AIndexHS300FreeWeight.ix[AIndexHS300FreeWeight['IndexCode'] == index_code, :]

	if np.any(tmp_dat.isnull()):
		logging.error('Index: ' + index_code + ', Weights Exist NaN ! ')
		return DataFrame()


	# 找到指数最早的记录
	firstRecordDate = tmp_dat['Date'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[tmp_dat['Date'] <= td, :]
		if tmp.empty:
			logging.warning('Index: ' + index_code + ', Weights Missing Before ' + td)
			continue
		tmp = tmp.ix[tmp['Date'] == tmp['Date'].max(), :]
		tmp['Date'] = td

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.ix[:, ['Symbol', 'Date', 'Weight']]
	dat = dat.sort_values(by=['Date', 'Symbol'])

	return dat



# ZZ500指数成分
def update_ZZ500Pool(config):

	fun_name = 'ZZ500Pool'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	AIndexMembers = get_CoreGenData('AIndexMembers')
	index_code = '000905.SH'
	IndexMembers = AIndexMembers.ix[AIndexMembers['IndexCode'] == index_code, :]
	tmp_dat = IndexMembers.fillna(value=datetime.today().strftime('%Y%m%d'))


	# 找到指数最早的记录
	firstRecordDate = tmp_dat['InDate'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[(tmp_dat['InDate'] <= td) & (tmp_dat['OutDate'] >= td), :]
		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp, how='left', on=['Symbol'])
		tmp = tmp.fillna(value=0)
		tmp.ix[tmp['IndexCode'] == index_code, 'IndexCode'] = 1

		tmp = tmp.ix[:, ['Symbol', 'Date', 'IndexCode']]
		tmp.columns = ['Symbol', 'Date', 'isMember']

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.sort_values(by=['Date', 'Symbol'])
	return dat



# ZZ500指数日权重
def update_ZZ500DailyWeight(config):

	fun_name = 'ZZ500DailyWeight'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	index_code = '000905.SH'
	AIndexHS300FreeWeight = get_CoreGenData('AIndexHS300FreeWeight')
	tmp_dat = AIndexHS300FreeWeight.ix[AIndexHS300FreeWeight['IndexCode'] == index_code, :]

	if np.any(tmp_dat.isnull()):
		logging.error('Index: ' + index_code + ', Weights Exist NaN ! ')
		return DataFrame()

	# 找到指数最早的记录
	firstRecordDate = tmp_dat['Date'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[tmp_dat['Date'] <= td, :]
		if tmp.empty:
			logging.warning('Index: ' + index_code + ', Weights Missing Before ' + td)
			continue
		tmp = tmp.ix[tmp['Date'] == tmp['Date'].max(), :]
		tmp['Date'] = td

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.ix[:, ['Symbol', 'Date', 'Weight']]
	dat = dat.sort_values(by=['Date', 'Symbol'])

	return dat



# ZZ800指数成分
def update_ZZ800Pool(config):

	fun_name = 'ZZ800Pool'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	AIndexMembers = get_CoreGenData('AIndexMembers')
	index_code = '000906.SH'
	IndexMembers = AIndexMembers.ix[AIndexMembers['IndexCode'] == index_code, :]
	tmp_dat = IndexMembers.fillna(value=datetime.today().strftime('%Y%m%d'))


	# 找到指数最早的记录
	firstRecordDate = tmp_dat['InDate'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[(tmp_dat['InDate'] <= td) & (tmp_dat['OutDate'] >= td), :]
		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp, how='left', on=['Symbol'])
		tmp = tmp.fillna(value=0)
		tmp.ix[tmp['IndexCode'] == index_code, 'IndexCode'] = 1

		tmp = tmp.ix[:, ['Symbol', 'Date', 'IndexCode']]
		tmp.columns = ['Symbol', 'Date', 'isMember']

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.sort_values(by=['Date', 'Symbol'])
	return dat



# ZZ800指数日权重
def update_ZZ800DailyWeight(config):

	fun_name = 'ZZ800DailyWeight'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	index_code = '000906.SH'
	AIndexHS300FreeWeight = get_CoreGenData('AIndexHS300FreeWeight')
	tmp_dat = AIndexHS300FreeWeight.ix[AIndexHS300FreeWeight['IndexCode'] == index_code, :]

	if np.any(tmp_dat.isnull()):
		logging.error('Index: ' + index_code + ', Weights Exist NaN ! ')
		return DataFrame()

	# 找到指数最早的记录
	firstRecordDate = tmp_dat['Date'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[tmp_dat['Date'] <= td, :]
		if tmp.empty:
			logging.warning('Index: ' + index_code + ', Weights Missing Before ' + td)
			continue
		tmp = tmp.ix[tmp['Date'] == tmp['Date'].max(), :]
		tmp['Date'] = td

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.ix[:, ['Symbol', 'Date', 'Weight']]
	dat = dat.sort_values(by=['Date', 'Symbol'])

	return dat



# ZZ1000指数成分
def update_ZZ1000Pool(config):

	fun_name = 'ZZ1000Pool'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	AIndexMembers = get_CoreGenData('AIndexMembers')
	index_code = '000852.SH'
	IndexMembers = AIndexMembers.ix[AIndexMembers['IndexCode'] == index_code, :]
	tmp_dat = IndexMembers.fillna(value=datetime.today().strftime('%Y%m%d'))

	# 找到指数最早的记录
	firstRecordDate = tmp_dat['InDate'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[(tmp_dat['InDate'] <= td) & (tmp_dat['OutDate'] >= td), :]
		tmp = pd.merge(ref.ix[ref['Date'] == td, :], tmp, how='left', on=['Symbol'])
		tmp = tmp.fillna(value=0)

		tmp.ix[tmp['IndexCode'] == index_code, 'IndexCode'] = 1

		tmp = tmp.ix[:, ['Symbol', 'Date', 'IndexCode']]
		tmp.columns = ['Symbol', 'Date', 'isMember']

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.sort_values(by=['Date', 'Symbol'])
	return dat




# ZZ1000指数日权重
def update_ZZ1000DailyWeight(config):

	fun_name = 'ZZ1000DailyWeight'
	start_dt = config['start_dt']
	end_dt = config['end_dt']
	periodTDates = config['periodTDates']

	index_code = '000852.SH'
	AIndexHS300FreeWeight = get_CoreGenData('AIndexHS300FreeWeight')
	tmp_dat = AIndexHS300FreeWeight.ix[AIndexHS300FreeWeight['IndexCode'] == index_code, :]

	if np.any(tmp_dat.isnull()):
		logging.error('Index: ' + index_code + ', Weights Exist NaN ! ')
		return DataFrame()

	# 找到指数最早的记录
	firstRecordDate = tmp_dat['Date'].min()
	start_dt = start_dt if start_dt > firstRecordDate else firstRecordDate

	REF_DF = get_CoreGenData('REF_DF')
	ref = REF_DF.ix[(REF_DF['Date'] >= start_dt) & (REF_DF['Date'] <= end_dt), :].copy()

	dat = DataFrame()
	tradeDates = ref['Date'].unique().tolist()

	for td in tradeDates:

		tmp = tmp_dat.ix[tmp_dat['Date'] <= td, :]
		if tmp.empty:
			logging.warning('Index: ' + index_code + ', Weights Missing Before ' + td)
			continue
		tmp = tmp.ix[tmp['Date'] == tmp['Date'].max(), :]
		tmp['Date'] = td

		dat = pd.concat([dat, tmp], sort=True)

	dat = dat.ix[:, ['Symbol', 'Date', 'Weight']]
	dat = dat.sort_values(by=['Date', 'Symbol'])

	return dat




# 更新周期节点数据
def update_refPeriodDates():

	stockCalendar = dataAssign('AShareCalendar')

	# # 返回数字年份
	fun_year = lambda x: datetime.strptime(x, '%Y%m%d').year
	stockCalendar['Year'] = stockCalendar['Date'].apply(fun_year)

	# 返回数字季度
	fun_quarter = lambda x: (int(datetime.strptime(x, '%Y%m%d').strftime('%m')) - 1) // 3 + 1
	stockCalendar['Quarter'] = stockCalendar['Date'].apply(fun_quarter)

	# 返回数字月份
	fun_month = lambda x: datetime.strptime(x, '%Y%m%d').month
	stockCalendar['Month'] = stockCalendar['Date'].apply(fun_month)

	# 一年中的星期数(00-53),星期一为星期的开始
	fun_week = lambda x: datetime.strptime(x, '%Y%m%d').strftime('%W')
	stockCalendar['Week'] = stockCalendar['Date'].apply(fun_week)


	# 提取季度节点
	Years = stockCalendar['Year'].unique().tolist()
	Tickers = ['Quarter', 'Month', 'Week']
	Types = ['head', 'tail']

	periodDatesDict = {
		'Quarter': {},
		'Month': {},
		'Week': {}
	}

	for ticker in Tickers:
		for Type in Types:

			res = []
			for year in Years:
				tmpDat = stockCalendar.ix[stockCalendar['Year'] == year, ['Date', ticker]].copy()
				if Type == 'head':
					tmp = tmpDat.groupby(by=ticker).apply(lambda x: x['Date'].tolist()[0])    # 函数结果最好不带索引，不然不好返回
				else:
					tmp = tmpDat.groupby(by=ticker).apply(lambda x: x['Date'].tolist()[-1])

				res = res + list(tmp)
			periodDatesDict[ticker][Type] = res

	pkl_read_write('./supplementaryData/refPeriodDates.pkl', 'write', periodDatesDict)


# ============================================ Funs_CalFactors ================================================= #
# -------------------------------------------------------
# ClosePrice
def calSubFCT_ClosePrice(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('StockTradeInfo')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'Close']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat


# MarketValue
def calSubFCT_MarketValue(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('StockTradeInfo')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'MarketValue']]
	dat['MarketValue'] = np.log(dat['MarketValue'].astype('float64'))
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat

# AFloatCap
def calSubFCT_AFloatCap(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('StockTradeInfo')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'AFloatCap']]
	dat['AFloatCap'] = np.log(dat['AFloatCap'].astype('float64'))
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat



# ----------------------------------------------------------------------------------------
# weighted volatility
def cwv(vec=None, halflife=60):
	if halflife == 0:
		return vec.std()
	else:
		Lambda = (0.5) ** (1 / halflife)
		tmp = Lambda ** np.arange(len(vec), 0, -1)
		weight = tmp / tmp.mean()
		return (vec * weight).std()

def calWeightedVolatility(stockCode, config):

	# print(stockCode)
	actionDates = config['actionDates']
	LBW = config['Params']['LBW']

	dat = getStockTDayReturn(stockCode)


	tradeDates = dat.index.tolist()

	# 这里的优先级要拿出来
	if len(dat) < LBW:
		logging.warning(stockCode + ': length of returns is smaller than LBW!')
		return DataFrame()

	if not actionDates:
		# 如果actionDates为空，那么取所有可能的
		ep = Series(range(LBW - 1, len(tradeDates)))
	else:
		ep = Series(np.where(Series(tradeDates).isin(actionDates))[0])
		ep = ep[ep >= (LBW - 1)]
		if (len(ep) == 0):
			logging.warning('no matched date with LBW tdays')
			return DataFrame()

	func2 = lambda x: cwv(dat[(x - LBW + 1): (x + 1)])
	res = DataFrame(ep.apply(func2))
	res['Date'] = list(np.array(tradeDates)[ep])
	res['Symbol'] = stockCode

	res.columns = ['Value', 'Date', 'Symbol']
	res = res.ix[:, ['Symbol', 'Date', 'Value']]

	return res



def Series2MergeDF(dat):

	tmp = DataFrame()
	for i in range(len(dat)):
		tmp = pd.concat([tmp, dat[i]], sort=True)

	tmp = tmp.sort_values(by=['Date', 'Symbol'])

	return tmp

def calSubFCT_WeightedVolatility(config):

	allStocks = get_EnvData('AllStocks')

	func = lambda x: calWeightedVolatility(x, config)
	res = allStocks.apply(func)

	res = Series2MergeDF(res)
	return res

# ----------------------------------------------------------------------------------------

def calAvgMoneyFlow(stockCode, config):

	# print(stockCode)
	actionDates = config['actionDates']
	LBW = config['Params']['LBW']
	mf_type = config['Params']['mf_type']

	# print(exist_EnvData('StockTradeInfo'))
	tmp = get_CoreGenData('StockTradeInfo')

	if mf_type not in tmp.columns.tolist():
		return DataFrame()

	dat = tmp.ix[tmp['Symbol'].isin([stockCode]), ['Date', mf_type]]
	dat.index = list(dat['Date'])
	dat = dat[mf_type]
	tradeDates = dat.index.tolist()

	if len(dat) < LBW:
		logging.warning(stockCode + ': length of returns is smaller than LBW!')
		return DataFrame()

	if not actionDates:
		ep = Series(range(LBW - 1, len(tradeDates)))
	else:
		ep = Series(np.where(Series(tradeDates).isin(actionDates))[0])
		ep = ep[ep >= (LBW - 1)]
		if (len(ep) == 0):
			logging.warning('no matched date with LBW tdays')
			return DataFrame()

	func2 = lambda x: dat[(x - LBW + 1): (x + 1)].mean()

	res = DataFrame(ep.apply(func2))
	res['Date'] = list(np.array(tradeDates)[ep])
	res['Symbol'] = stockCode

	res.columns = ['Value', 'Date', 'Symbol']
	res = res.ix[:, ['Symbol', 'Date', 'Value']]

	return res


def calSubFCT_AvgMoneyFlow(config):

	allStocks = get_EnvData('AllStocks')

	func = lambda x: calAvgMoneyFlow(x, config)
	res = allStocks.apply(func)

	res = Series2MergeDF(res)
	return res


# -------------------------------------------------

def calAvgMoneyFlowFloatCapRatio(stockCode, config):

	# print(stockCode)
	actionDates = config['actionDates']
	LBW = config['Params']['LBW']
	mf_type = config['Params']['mf_type']

	# print(exist_EnvData('StockTradeInfo'))
	tmp = get_CoreGenData('StockTradeInfo')

	if mf_type not in tmp.columns.tolist():
		return DataFrame()

	dat = tmp.ix[tmp['Symbol'].isin([stockCode]), ['Date', mf_type, 'AFloatCap']]

	dat.index = list(dat['Date'])
	dat = dat.ix[:, [mf_type, 'AFloatCap']]
	tradeDates = dat.index.tolist()

	if len(dat) < LBW:
		logging.warning(stockCode + ': length of returns is smaller than LBW!')
		return DataFrame()

	if not actionDates:
		ep = Series(range(LBW - 1, len(tradeDates)))
	else:
		ep = Series(np.where(Series(tradeDates).isin(actionDates))[0])
		ep = ep[ep >= (LBW - 1)]
		if (len(ep) == 0):
			logging.warning('no matched date with LBW tdays')
			return DataFrame()

	func = lambda x: dat[mf_type][(x - LBW + 1): (x + 1)].mean() / dat['AFloatCap'][(x - LBW + 1): (x + 1)].mean()

	res = DataFrame(ep.apply(func))
	res['Date'] = list(np.array(tradeDates)[ep])
	res['Symbol'] = stockCode

	res.columns = ['Value', 'Date', 'Symbol']
	res = res.ix[:, ['Symbol', 'Date', 'Value']]

	return res


def calSubFCT_AvgMoneyFlowFloatCapRatio(config):

	allStocks = get_EnvData('AllStocks')

	func = lambda x: calAvgMoneyFlowFloatCapRatio(x, config)
	res = allStocks.apply(func)

	res = Series2MergeDF(res)
	return res


# ------------------------------------------------------

def calAvgMoneyFlowTurnOverRatio(stockCode, config):

	# print(stockCode)
	actionDates = config['actionDates']
	LBW = config['Params']['LBW']
	mf_type = config['Params']['mf_type']

	tmp = get_CoreGenData('StockTradeInfo')

	if mf_type not in tmp.columns.tolist():
		return DataFrame()

	dat = tmp.ix[tmp['Symbol'].isin([stockCode]), ['Date', mf_type, 'TurnOver']]

	dat.index = list(dat['Date'])
	dat = dat.ix[:, [mf_type, 'TurnOver']]
	tradeDates = dat.index.tolist()

	if len(dat) < LBW:
		logging.warning(stockCode + ': length of returns is smaller than LBW!')
		return DataFrame()

	if not actionDates:
		ep = Series(range(LBW - 1, len(tradeDates)))
	else:
		ep = Series(np.where(Series(tradeDates).isin(actionDates))[0])
		ep = ep[ep >= (LBW - 1)]
		if (len(ep) == 0):
			logging.warning('no matched date with LBW tdays')
			return DataFrame()

	func = lambda x: dat[mf_type][(x - LBW + 1): (x + 1)].mean() / dat['TurnOver'][(x - LBW + 1): (x + 1)].mean()

	res = DataFrame(ep.apply(func))
	res['Date'] = list(np.array(tradeDates)[ep])
	res['Symbol'] = stockCode

	res.columns = ['Value', 'Date', 'Symbol']
	res = res.ix[:, ['Symbol', 'Date', 'Value']]

	return res


def calSubFCT_AvgMoneyFlowTurnOverRatio(config):

	allStocks = get_EnvData('AllStocks')

	func = lambda x: calAvgMoneyFlowTurnOverRatio(x, config)
	res = allStocks.apply(func)

	res = Series2MergeDF(res)
	return res



# ----------------------------------------------------------------------------------------
def simplelinReg(dat, half_life = 60):

	# 不想要half_life，设为0
	if half_life > 0:
		Lambda = (0.5) ** (1 / half_life)
		tmp = Lambda ** np.arange(dat.shape[0], 0, -1)
		weight = tmp / tmp.mean()
		dat = dat * DataFrame(np.tile(weight, [dat.shape[1], 1]).T, index=dat.index, columns=dat.columns)

	datCov = dat.cov()
	beta = datCov.ix['stockRet', 'indexRet'] / datCov.ix['indexRet', 'indexRet']
	resid_sd = (dat['stockRet'] - beta * dat['indexRet']).std()

	return {'beta': beta, 'resid_sd': resid_sd}


def calBeta(stockCode, config):

	# print(stockCode)
	actionDates = config['actionDates']
	LBW = config['Params']['LBW']
	indexCode = config['Params']['indexCode']

	dat = DataFrame(getStockTDayReturn(stockCode))
	dat.columns = ['stockRet']

	if not indexCode:
		indexCode = '000300.SH'

	indexDat = DataFrame(getIndexTDayReturn(indexCode))
	indexDat.columns = ['indexRet']

	tmp = indexDat.join(dat)
	tmp = pd.DataFrame({"indexRet": list(tmp.indexRet.values),
						"stockRet": list(tmp.stockRet.values)}, index=tmp.index)

	# remove nan
	tmp = tmp.dropna(axis=0, how='any')

	tradeDates = tmp.index.tolist()

	if tmp.shape[0] < LBW:
		logging.warning('InterSection of ' + stockCode + ' and ' + indexCode + ' is smaller than LBW!')
		return DataFrame()

	if not actionDates:
		ep = Series(range(LBW - 1, len(tradeDates)))
	else:
		ep = Series(np.where(Series(tradeDates).isin(actionDates))[0])
		ep = ep[ep >= (LBW - 1)]
		if (len(ep) == 0):
			logging.warning('no matched date with LBW tdays')
			return DataFrame()

	# 计算Beta
	func2 = lambda x: simplelinReg(dat=tmp[(x - LBW + 1): (x + 1)])['beta']

	res = DataFrame(ep.apply(func2))
	res['Date'] = list(np.array(tradeDates)[ep])
	res['Symbol'] = stockCode

	res.columns = ['Value', 'Date', 'Symbol']
	res = res.ix[:, ['Symbol', 'Date', 'Value']]

	return res


def calSubFCT_Beta(config):

	allStocks = get_EnvData('AllStocks')

	func = lambda x: calBeta(x, config)
	res = allStocks.apply(func)

	res = Series2MergeDF(res)
	return res

# --------------------------------------------------------------------------------------------
def calVolofCAPMResid(stockCode, config):

	# print(stockCode)
	actionDates = config['actionDates']
	LBW = config['Params']['LBW']
	indexCode = config['Params']['indexCode']

	dat = DataFrame(getStockTDayReturn(stockCode))
	dat.columns = ['stockRet']


	if not indexCode:
		indexCode = '000300.SH'
	indexDat = DataFrame(getIndexTDayReturn(indexCode))
	indexDat.columns = ['indexRet']

	tmp = indexDat.join(dat)
	tmp = pd.DataFrame({"indexRet": list(tmp.indexRet.values),
						"stockRet": list(tmp.stockRet.values)}, index=tmp.index)

	# remove nan
	tmp = tmp.dropna(axis=0, how='any')

	tradeDates = tmp.index.tolist()

	if tmp.shape[0] < LBW:
		logging.warning('InterSection of ' + stockCode + ' and ' + indexCode + ' is smaller than LBW!')
		return DataFrame()

	if not actionDates:
		ep = Series(range(LBW - 1, len(tradeDates)))
	else:
		ep = Series(np.where(Series(tradeDates).isin(actionDates))[0])
		ep = ep[ep >= (LBW - 1)]
		if (len(ep) == 0):
			logging.warning('no matched date with LBW tdays')
			return DataFrame()

	# 计算residual
	func2 = lambda x: simplelinReg(dat=tmp[(x - LBW + 1): (x + 1)])['resid_sd']

	res = DataFrame(ep.apply(func2))
	res['Date'] = list(np.array(tradeDates)[ep])
	res['Symbol'] = stockCode

	res.columns = ['Value', 'Date', 'Symbol']
	res = res.ix[:, ['Symbol', 'Date', 'Value']]

	return res



def calSubFCT_VolofCAPMResid(config):

	allStocks = get_EnvData('AllStocks')

	func = lambda x: calVolofCAPMResid(x, config)
	res = allStocks.apply(func)

	res = Series2MergeDF(res)
	return res


# --------------------------------------------------------------------------------------------
# Calculate TurnOverRate

def calTurnOverRate(stockCode, config):

	# print(stockCode)
	actionDates = config['actionDates']
	LBW = config['Params']['LBW']

	dat = get_CoreGenData('StockTradeInfo')
	dat = dat.ix[dat['Symbol'].isin([stockCode]), :].copy()

	dat['TurnOverRate'] = dat['Volume'] / dat['FloatAShr']

	dat.index = list(dat['Date'])
	dat = dat['TurnOverRate']
	tradeDates = dat.index.tolist()

	if len(dat) < LBW:
		logging.warning(stockCode + ': length of returns is smaller than LBW!')
		return DataFrame()

	if not actionDates:
		ep = Series(range(LBW - 1, len(tradeDates)))
	else:
		ep = Series(np.where(Series(tradeDates).isin(actionDates))[0])
		ep = ep[ep >= (LBW - 1)]
		if (len(ep) == 0):
			logging.warning('no matched date with LBW tdays')
			return DataFrame()

	func = lambda x: dat[(x - LBW + 1): (x + 1)].mean()

	res = DataFrame(ep.apply(func))
	res['Date'] = list(np.array(tradeDates)[ep])
	res['Symbol'] = stockCode

	res.columns = ['Value', 'Date', 'Symbol']
	res = res.ix[:, ['Symbol', 'Date', 'Value']]

	return res


def calSubFCT_TurnOverRate(config):

	allStocks = get_EnvData('AllStocks')

	func = lambda x: calTurnOverRate(x, config)
	res = allStocks.apply(func)

	res = Series2MergeDF(res)
	return res


# --------------------------------------------------------------------------------------------
# Calculate TurnOverRateRatio

def calTurnOverRateRatio(stockCode, config):

	# print(stockCode)
	actionDates = config['actionDates']
	LBW_L = config['Params']['LBW_L']
	LBW_S = config['Params']['LBW_S']

	dat = get_CoreGenData('StockTradeInfo')
	dat = dat.ix[dat['Symbol'].isin([stockCode]), :].copy()
	dat['TurnOverRate'] = dat['Volume'] / dat['FloatAShr']

	dat.index = list(dat['Date'])
	dat = dat['TurnOverRate']
	tradeDates = dat.index.tolist()

	if len(dat) < LBW_L:
		logging.warning(stockCode + ': length of returns is smaller than LBW!')
		return DataFrame()

	if not actionDates:
		ep = Series(range(LBW_L - 1, len(tradeDates)))
	else:
		ep = Series(np.where(Series(tradeDates).isin(actionDates))[0])
		ep = ep[ep >= (LBW_L - 1)]
		if (len(ep) == 0):
			logging.warning('no matched date with LBW tdays')
			return DataFrame()

	func = lambda x: dat[(x - LBW_L + 1): (x + 1)].mean() / dat[(x - LBW_S + 1): (x + 1)].mean()

	res = DataFrame(ep.apply(func))
	res['Date'] = list(np.array(tradeDates)[ep])
	res['Symbol'] = stockCode

	res.columns = ['Value', 'Date', 'Symbol']
	res = res.ix[:, ['Symbol', 'Date', 'Value']]

	return res


def calSubFCT_TurnOverRateRatio(config):

	allStocks = get_EnvData('AllStocks')

	func = lambda x: calTurnOverRateRatio(x, config)
	res = allStocks.apply(func)

	res = Series2MergeDF(res)
	return res


# --------------------------- TTMHis 表内容 ------------------------------------------------------------
# 基于TTMFactor数据

# EPSTTM
def calSubFCT_EPSTTM(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('TTMFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'EPSTTM']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat


# CFPSTTM
def calSubFCT_CFPSTTM(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('TTMFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'CFPSTTM']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat


# CFPSOverPriceTTM
def calSubFCT_CFPSOverPriceTTM(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('TTMFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'CFPSOverPriceTTM']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat



# InvPETTM
def calSubFCT_InvPETTM(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('TTMFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'InvPETTM']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat


# OpRatioTTM
def calSubFCT_OpRatio(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('TTMFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'OpRatioTTM']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat


# --------------------------- BalanceSheet 表内容 ------------------------------------------------------------
# 基于BalanceSheetFactor数据

# EquityDebtRatio
def calSubFCT_EquityDebtRatio(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('BalanceSheetFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'EquityDebtRatio']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat

# InvPB
def calSubFCT_InvPB(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('BalanceSheetFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'InvPB']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat

# OpCapital
def calSubFCT_OpCapital(config):
	actionDates = config['actionDates']

	dat = get_CoreGenData('BalanceSheetFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'OpCapital']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat


# CurrentRatio
def calSubFCT_CurrentRatio(config):

	actionDates = config['actionDates']

	dat = get_CoreGenData('BalanceSheetFactor')

	# 保留actionDates的因子值
	dat = dat.ix[dat['Date'].isin(actionDates), ['Symbol', 'Date', 'CurrentRatio']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat

# --------------------------- TTMHis AND BalanceSheet ---------------------------------------
# 基于 TTMFactor AND BalanceSheetFactor

# EBITOpCapitalRatio
def calSubFCT_EBITOpCapitalRatio(config):

	actionDates = config['actionDates']

	tmp = get_CoreGenData('TTMFactor')
	tmp = tmp.ix[tmp['Date'].isin(actionDates), :]

	dat = get_CoreGenData('BalanceSheetFactor')
	dat = dat.ix[dat['Date'].isin(actionDates), :]

	dat = pd.merge(dat.ix[:, ['Symbol', 'Date', 'OpCapital']],
				   tmp.ix[:, ['Symbol', 'Date', 'EBITTTM']], on=['Symbol', 'Date'])

	# OpCapital可能有0项，看看是不是TTMFactor和BalanceSheetFactor处理NaN了
	dat.ix[dat['OpCapital'] == 0, 'OpCapital'] = np.nan

	dat['EBITOpCapitalRatio'] = dat['EBITTTM'] / dat['OpCapital']
	dat = dat.ix[:, ['Symbol', 'Date', 'EBITOpCapitalRatio']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat



# FixAssetTurnover
def calSubFCT_FixAssetsTurnover(config):

	actionDates = config['actionDates']

	tmp = get_CoreGenData('TTMFactor')
	tmp = tmp.ix[tmp['Date'].isin(actionDates), :]

	dat = get_CoreGenData('BalanceSheetFactor')
	dat = dat.ix[dat['Date'].isin(actionDates), :]

	dat = pd.merge(dat.ix[:, ['Symbol', 'Date', 'FixAssets']],
				   tmp.ix[:, ['Symbol', 'Date', 'TotOpRevTTM']], on=['Symbol', 'Date'])

	# FixAssets，看看是不是TTMFactor和BalanceSheetFactor处理NaN了
	dat.ix[dat['FixAssets'] == 0, 'FixAssets'] = np.nan

	dat['FixAssetsTurnover'] = dat['TotOpRevTTM'] / dat['FixAssets']
	dat = dat.ix[:, ['Symbol', 'Date', 'FixAssetsTurnover']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat


# TotAssetTurnover
def calSubFCT_TotAssetsTurnover(config):

	actionDates = config['actionDates']

	tmp = get_CoreGenData('TTMFactor')
	tmp = tmp.ix[tmp['Date'].isin(actionDates), :]

	dat = get_CoreGenData('BalanceSheetFactor')
	dat = dat.ix[dat['Date'].isin(actionDates), :]

	dat = pd.merge(dat.ix[:, ['Symbol', 'Date', 'TotAssets']],
				   tmp.ix[:, ['Symbol', 'Date', 'TotOpRevTTM']], on=['Symbol', 'Date'])

	# FixAssets，看看是不是TTMFactor和BalanceSheetFactor处理NaN了
	dat.ix[dat['TotAssets'] == 0, 'TotAssets'] = np.nan

	dat['TotAssetsTurnover'] = dat['TotOpRevTTM'] / dat['TotAssets']
	dat = dat.ix[:, ['Symbol', 'Date', 'TotAssetsTurnover']]
	dat.columns = ['Symbol', 'Date', 'Value']

	return dat



# ============================================ Funs_NormFactors ================================================= #


def handleRawFactorValue(dat=None):
	# dat = Series(dat.values.astype('float64'), index=list(dat.index), name=dat.name)

	MED = dat.median()

	if math.isnan(MED):
		logging.warning('Normalization: Raw Factor Values are ALL NAN (' + dat.name + ')')
		dat[:] = 0
		return dat

	# 完善细节！！NaN超过一定比例

	MAD = abs(dat - MED).median()
	dat[dat > MED + 3 * 1.4826 * MAD] = MED + 3 * 1.4826 * MAD
	dat[dat < MED - 3 * 1.4826 * MAD] = MED - 3 * 1.4826 * MAD

	std = dat.std()
	if math.isnan(std) or std == 0:
		logging.warning('Standard Deviation is NaN or 0! Be Carefully!')
		std = abs(dat).mean()
	dat = (dat - dat.mean()) / std

	# handle nan
	dat = dat.fillna(value=0)
	return dat

# 这里要分组求，至少按照日期，这里的正则化是相当于对截面股票数据


def normlizeFactor(dat, norm_type='Norm'):

	colnames = list(dat.columns)
	if ('Symbol' not in colnames) or ('Date' not in colnames):
		logging.warning('Raw factor value missing column Symbol or Date')
		return DataFrame()
	tmp = dat.ix[:, ['Symbol', 'Date']].copy()
	len_dat = dat.shape[0]

	colnames.remove('Symbol')
	colnames.remove('Date')
	fct_names = colnames

	if norm_type == 'Norm':
		# 全市场
		for fct_name in fct_names:
			tmp[fct_name] = dat.groupby(by='Date')[fct_name].apply(handleRawFactorValue)

	# 市值中性
	if norm_type == 'MVNorm':
		StockMarketValue = get_CoreGenData('StockMarketValue')

		# StockMarketValue.ix[:, 'Group'] = StockMarketValue['Group'].astype('float64')
		dat = pd.merge(dat, StockMarketValue.ix[:, ['Symbol', 'Date', 'MVGroup']], on=['Symbol', 'Date'])

		if len_dat != dat.shape[0]:
			logging.warning('Merge of StockMarketValue and Raw Factor cannot match Original Raw Factor!')
			return DataFrame()

		for fct_name in fct_names:
			tmp[fct_name] = dat.groupby(by=['Date', 'MVGroup'])[fct_name].apply(handleRawFactorValue)

	# 行业中性
	if norm_type == 'IndNorm':
		IndustryCITICS = get_CoreGenData('IndustryCITICS')
		dat = pd.merge(dat, IndustryCITICS.ix[:, ['Symbol', 'Date', 'IndexCode']], on=['Symbol', 'Date'])

		if len_dat != dat.shape[0]:
			logging.warning('Merge of IndustryCITICS and Raw Factor cannot match Original Raw Factor!')
			return DataFrame()

		for fct_name in fct_names:
			tmp[fct_name] = dat.groupby(by=['Date', 'IndexCode'])[fct_name].apply(handleRawFactorValue)
			tmp.ix[dat['IndexCode'] == 'Others', fct_name] = 0  # Others 用0代替

	# 市值行业都中性
	if norm_type == 'IndMVNorm':
		IndustryCITICS = get_CoreGenData('IndustryCITICS')
		StockMarketValue = get_CoreGenData('StockMarketValue')
		dat = pd.merge(dat, IndustryCITICS.ix[:, ['Symbol', 'Date', 'IndexCode']], on=['Symbol', 'Date'])
		dat = pd.merge(dat, StockMarketValue.ix[:, ['Symbol', 'Date', 'MVGroup']])

		if len_dat != dat.shape[0]:
			logging.warning('Merge of IndustryCITICS, MarketValue and Raw Factor cannot match Original Raw Factor!')
			return DataFrame()

		for fct_name in fct_names:
			tmp[fct_name] = dat.groupby(by=['Date', 'IndexCode', 'MVGroup'])[fct_name].apply(handleRawFactorValue)
			tmp.ix[dat['IndexCode'] == 'Others', fct_name] = 0  # Others 用0代替

	if np.any(tmp.isnull()):
		logging.warning('Normalized Factors Exist NaN ! Be careful !')

	return tmp

# ============================================ Funs_EvalFactors ================================================= #

# dateType = 'head' or 'tail' ，返回月初还是月末数据
# periodDates 数据包含periodDates的首尾信息
def get_PortfolioBuiltDates(actionDates, freqofTrades='Month', dateType='head'):

	refPeriodDates = get_CoreGenData('refPeriodDates')

	tmp = refPeriodDates[freqofTrades][dateType]
	res_ind = [td in actionDates for td in tmp]
	res = list(Series(tmp)[res_ind])

	return res

def get_FactorValues(fct_name, relativePath='./factorData/Norm/'):

	fct_filename = relativePath + fct_name + '.pkl'
	dat = pkl_read_write(fct_filename, 'read')

	return dat

def assignGroupValue(dat, numofGroups):
	# dat is a series with Symbol as index
	dat_len = len(dat)
	interval = int(dat_len / numofGroups)
	dat = dat.rank()

	for i in range(numofGroups):
		if i < (numofGroups - 1):
			dat[(dat >= i * interval) & (dat < (i + 1) * interval)] = int(i + 1)
		else:
			dat[(dat >= i * interval)] = int(i + 1)

	if len(dat.unique()) < numofGroups:
		logging.warning('Group Number < ' + repr(numofGroups) + ', Grouped by StockCode, result is not reliable!')
		# 按股票代码顺序给分组
		for i in range(numofGroups):
			# print(i)
			if i < (numofGroups - 1):
				dat[(i * interval):((i + 1) * interval)] = int(i + 1)
			else:
				dat[(i * interval):] = int(i + 1)

	return dat

def removeAllZeros(factorValue):
	# factorValue:  Symbol | Date | Value

	for td in factorValue['Date'].unique().tolist():

		if not np.any(~ (factorValue.ix[factorValue['Date'] == td, 'Value'] == 0)):
			continue
		else:
			# 遇到不是所有股票因子值为0的日期，退出，从这一天开始算有效日期
			break

	factorValue = factorValue.ix[factorValue['Date'] >= td, :]
	return factorValue


def evalFactorPerformance(evalFactor, config):

	actionDates = config['actionDates']
	freqofTrades = config['freqofTrades']
	# selectPool = config['selectPool']  # 在某个股票池里选股
	weightType = config['weightType']
	dateType = config['dateType']
	numofGroups = config['numofGroups']
	totCost = config['totCost']

	# 把StockTradeInfo LOAD 到环境中
	StockTradeInfo = get_CoreGenData('StockTradeInfo')

	# 获取因子值，这里后期要改善， # 处理前端的NaN,如果对于Normalized Factor, 就是全0
	factorValue = removeAllZeros(get_FactorValues(evalFactor + '_Norm'))

	# 参照stockTradeInfo
	dat = pd.merge(StockTradeInfo, factorValue, on=['Date', 'Symbol']).ix[:,
		  	['Symbol', 'Date', 'TradeStatus', 'Value']]


	# 确认回测日期
	if len(actionDates) == 0:
		actionDates = factorValue['Date'].unique().tolist()
	else:
		ind = [td in factorValue['Date'].unique().tolist() for td in actionDates]
		actionDates = list(Series(actionDates)[ind])
		# 这里耗时太长，非常不划算，这里的for循环还是要想办法替代

	portfolioBuiltDates = get_PortfolioBuiltDates(actionDates, freqofTrades, dateType)

	# 根据调仓时间来调整actionDates
	actionDates = actionDates[actionDates.index(portfolioBuiltDates[0]):]

	#
	dat['periodRet'] = 0
	dat['Weight'] = 0

	portLabels = \
		list(Series(['port'] * numofGroups) + Series(list(np.arange(1, numofGroups + 1).astype('str'))))
	portRets = DataFrame(np.zeros([numofGroups, len(actionDates)]), index=portLabels, columns=actionDates)

	for actionDate in actionDates:

		# 看是不是调仓第一天，用收盘价的回测，当天收益为0，但是后期调仓日也是有收益的！
		if actionDates.index(actionDate):
			print('BackTest Date: ' + actionDate)
			Symbols = candidates.index.tolist()
			candidates.ix[:, 'periodRet'] = \
				(1 - totCost) * get_PFun(Symbols, actionDate) / get_PFun(Symbols, portfolioBuiltDate) - 1

			# calculate port return， 默认为0了
			portRets.ix[:, actionDate] = list(
				candidates.groupby('Group').apply(lambda c: (c['periodRet'] * c['Weight']).sum()))

		if actionDate in portfolioBuiltDates:

			if not actionDates.index(actionDate):
				print('BackTest Date: ' + actionDate)

			portfolioBuiltDate = actionDate
			candidates = dat.ix[dat['Date'] == portfolioBuiltDate, :]
			candidates.index = list(candidates['Symbol'])
			del candidates['Symbol']

			# Delete untradable stocks on portfolio built day
			candidates = candidates.ix[candidates['TradeStatus'] == 1, :].copy()
			try:
				candidates['Group'] = pd.qcut(candidates['Value'], numofGroups, labels=range(1, numofGroups+1))
			except:
				logging.warning('Too Much Extreme on PortfolioBuiltDate: ' + portfolioBuiltDate)
				candidates['Group'] = assignGroupValue(candidates['Value'], numofGroups)
			# 看看这里要不要排序

			if weightType == 1:
				# Equal weights
				candidates['Weight'] = \
					candidates.groupby('Group')['Weight'].apply(lambda c: Series(np.ones(len(c)), index=list(c.index)) / len(c))
					# 财报TTM中有股票代码，日期重复的数据，在index=xx时，会保持
			else:
				# weighted
				candidates['Weight'] = \
					candidates.groupby('Group')['Weight'].apply(lambda c: np.log(2 + c.rank()) / np.log(2 + c.rank()).sum())

	# handle raw port_d_ret
	portRets = portRets.T

	# handle raw results
	# periodRets = handleRawPortReturn_PeriodReturn(portRets, portfolioBuiltDates)
	NetValue = handleRawPortReturn_TDayNetValue(portRets, portfolioBuiltDates)
	rets = handleNetvalue_Return(NetValue, 1)

	return {'rets': rets, 'portfolioBuiltDates': portfolioBuiltDates, 'config': config}


def calPortPeriodReturn(eval_result):

	rets = eval_result['rets']
	portfolioBuiltDates = eval_result['portfolioBuiltDates']

	tradeDates = rets.index.tolist()
	res = DataFrame()

	for portfolioBuiltDate in portfolioBuiltDates:
		if portfolioBuiltDates.index(portfolioBuiltDate) != 0:
			preBuiltDate = portfolioBuiltDates[portfolioBuiltDates.index(portfolioBuiltDate) - 1]
			periodDates = tradeDates[tradeDates.index(preBuiltDate): (tradeDates.index(portfolioBuiltDate) + 1)]
			tmp = rets.ix[periodDates, :]

			# 把第一天设置为0
			tmp.ix[0, :] = 0
			# periodStartDate = tradeDates[tradeDates.index(preBuiltDate) + 1]
			res = pd.concat([res, DataFrame((tmp + 1).cumprod().ix[-1, :] - 1).T], sort=True)

	return res


def calLongShortCombo(port_rets):
	'''
		# (x_mean - mu) / (x_std / sqrt(n -1))
		stat_value =
					(combo_rets.mean() - 0) / combo_rets.std() * (np.sqrt(len(combo_rets) - 1)
	'''

	if port_rets.shape[1] < 2:
		logging.warning('Input Returns with Less than 2 Ports, Long-Short Combo Can not Be Done !')
		return []

	# rets = eval_result['rets']
	rets = port_rets
	combo_rets = rets.ix[:, -1] - rets.ix[:, 0]
	combo_netvalue = handleNetvalue_Return(combo_rets, 2)

	# 是否显著不为0
	stat_value, pvalue = stats.ttest_1samp(combo_rets, 0)

	res = {'rets': combo_rets,
		   'netvalue': combo_netvalue,
		   'stat_value': stat_value,
		   'pvalue': pvalue
		   }
	return res



# ============================================ Funs_CombineFactors =============================================== #

def get_dir_pkl(path):
	file_names = os.listdir(path)
	pkl_files = []
	for file_name in file_names:
		if file_name[-4:] == '.pkl':
			pkl_files.append(file_name)
	return pkl_files


def fetch_files_dir(user_dir, file_type=None):
	'''
	功能：查找指定路径下的某类文件（如.csv）文件名,或所有文件
	————————————————————————————————————————————————————————————
	:param user_dir: 指定路径，可以是绝对路径，可以是相对路径
	:param file_type: 需要查找的某类文件名（如.csv）
	:return: 返回查找的文件名
	'''

	file_names = os.listdir(user_dir)

	if file_type == None:
		pass
	else:
		tmp_file_names = []
		for file_name in file_names:
			if os.path.splitext(file_name)[1] == file_type:
				tmp_file_names.append(file_name)
		file_names = tmp_file_names

	return file_names


# 该函数后面要完善, 对应不同函数名
def mergeFactorValues(path):

	fct_files = get_dir_pkl(path)

	dat = pkl_read_write(path + '/' + fct_files[0], 'read')
	for fct_file in fct_files[1:]:
		dat = pd.merge(dat, pkl_read_write(path + '/' + fct_file, 'read'), on=['Symbol', 'Date'])

	return dat

def selectFactors(portRetDict, factorValues, config):


	sepThreshold = config['sepThreshold']
	minNumofValidObs = config['minNumofValidObs']
	numofFactorsToSelect = config['numofFactorsToSelect']
	corThreshold = config['corThreshold']
	inSampleDates = config['inSampleDates']

	selectedFactorIndex = []
	selectedFactorNames = []
	factorOrderFlag = []
	inSampleRet = []
	inSampleSharpe = []

	numofFactors = len(portRetDict)
	factorNames = list(portRetDict.keys())
	actionDates = factorValues['Date'].unique().tolist()
	calFactorDates = Series(actionDates)[Series(actionDates).isin(inSampleDates)].tolist()

	# 第一步： 筛选因子值对收益确实有影响的因子
	for i in range(0, numofFactors):
		rets = portRetDict[factorNames[i]]['rets'].ix[inSampleDates, :]
		rets = rets.dropna(how='any')

		if rets.shape[0] < minNumofValidObs:
			continue

		# 计算平均日收益
		avg_ret = rets.mean()
		# 计算夏普比率
		avg_sharpe = rets.mean() / rets.std() * np.sqrt(244)

		separation = 0.5 * np.corrcoef(avg_ret.values, np.array(range(1, (len(avg_ret) + 1))))[0, 1] \
					 + 0.5 * np.corrcoef(avg_sharpe.values, np.array(range(1, (len(avg_ret) + 1))))[0, 1]
		# 这里可以考虑用spearman相关性

		if separation == np.nan:
			continue

		if np.abs(separation) >= sepThreshold:
			flag = 1 if separation > 0 else -1
			# 三元判断

			if flag == 1:
				inSampleRet.append(avg_ret[-1])
				inSampleSharpe.append(avg_sharpe[-1])
			else:
				inSampleRet.append(avg_ret[0])
				inSampleSharpe.append(avg_sharpe[0])
			selectedFactorIndex.append(i)
			selectedFactorNames.append(factorNames[i])
			factorOrderFlag.append(flag)


	# 第二步： 筛选一定数量的因子，Top N 的因子
	candidates = DataFrame({'factorIndex': selectedFactorIndex, 'factorOrderFlag': factorOrderFlag,
							'factorRet': inSampleRet, 'factorSharpe': inSampleSharpe, 'factorName': selectedFactorNames})
	candidates = candidates.ix[:, ['factorName', 'factorIndex', 'factorOrderFlag', 'factorRet', 'factorSharpe']]

	# 根据0.5的ret和0.5的夏普排序
	candidates['fct_rank'] = list(0.5 * candidates['factorRet'].rank() + 0.5 * candidates['factorSharpe'].rank())
	# 出现相同rank取第一个

	tmp = candidates.ix[:, ['factorName', 'fct_rank']].copy()
	tmp.index = tmp['factorName']
	tmp = tmp['fct_rank']

	ls = {'candidates': tmp, 'selected': [], 'corMat': []}
	# candidates is a Series with factorName as index and fct_rank as values
	ls['corMat'] = \
		DataFrame(np.eye(len(candidates['factorName'])), index=candidates['factorName'], columns=candidates['factorName'])


	def selectFromCandidates(ls, corThreshold):
		# print('Come into selectFromCandidates function...')
		selected = ls['selected']
		candidates = ls['candidates']
		corMat = ls['corMat']

		if len(selected) == 0:
			factorNameList = list(candidates.index)
			tempk = np.where(candidates[:] == candidates[:].max())[0][0]
			ls['selected'].append(factorNameList[tempk])
			ls['candidates'] = ls['candidates'].drop(factorNameList[tempk])
			return ls

		benchFactorName = selected[-1]

		tempList = []
		for k in range(len(candidates)):
			compareFactorName = list(candidates.index)[k]
			dat = factorValues.ix[factorValues['Date'].isin(calFactorDates), ['Date', benchFactorName, compareFactorName]]
			corMean = (dat.groupby('Date').apply(lambda tmp: tmp.corr(method='spearman').ix[0, 1])).mean()
			# mean along dates

			corMat.ix[benchFactorName, compareFactorName] = corMean
			corMat.ix[compareFactorName, benchFactorName] = corMean

			if np.abs(corMean) < corThreshold:
				tempList.append(compareFactorName)

		# 找到满足条件的股票
		if len(tempList) > 0:
			tempRank = candidates[tempList]
			tempk = np.where(tempRank[:] == tempRank[:].max())[0][0]
			ls['selected'].append(tempRank.index.tolist()[tempk])
			candidates = ls['candidates'].drop(tempRank.index.tolist()[tempk])
		else:
			candidates = Series()

		ls['candidates'] = candidates
		ls['corMat'] = corMat
		return ls


	# Program to call
	while True:
		ls = selectFromCandidates(ls, corThreshold)
		if len(ls['selected']) >= numofFactorsToSelect or len(ls['candidates']) == 0:
			break

	selectedFactors = candidates.ix[candidates['factorName'].isin(ls['selected']), :].copy()  # 这里加copy后面就不错有warning

	return selectedFactors

# 这个要完善
def evalPerformanceOfTopStk(dat, fct_name, tradeList, weightType=1, totCost=0.00246):

	actionDates = list(dat['Date'].unique())
	portfolioBuiltDates = list(tradeList.keys())

	# add two columns: periodReturn, Weight
	dat['periodReturn'] = 0   # period return is accumulative according to portfolio built day!!!
	dat['Weight'] = 0

	# Set first actionDate as first portfolio built day
	if actionDates[0] != portfolioBuiltDates[0]:
		actionDates = actionDates[actionDates.index(portfolioBuiltDates[0]):]

	periodRets = Series(np.zeros(len(actionDates)), index=actionDates, name='periodRets')
	for actionDate in actionDates:

		if actionDate in portfolioBuiltDates:
			print('BackTest Date: ' + actionDate)
			portfolioBuiltDate = actionDate
			stockCandidates = dat.ix[(dat['Date'].isin([portfolioBuiltDate])) & (dat['Symbol'].isin(tradeList[portfolioBuiltDate])), :].copy()

			if weightType == 1:
				stockCandidates['Weight'] = 1 / len(tradeList[portfolioBuiltDate])  #
			else:
				stockCandidates['Weight'] = np.log(2 + stockCandidates[fct_name].rank()) / np.log(2 + stockCandidates[fct_name].rank()).sum()

		else:
			print('BackTest Date: ' + actionDate)
			Symbols = list(stockCandidates['Symbol'])
			stockCandidates.ix[:, 'periodReturn'] = \
				list((1 - totCost) * get_PFun(Symbols, actionDate) / get_PFun(Symbols, portfolioBuiltDate) - 1)
			# the time cost becomes longer, maybe because of get_PFun()

		# calculate port return
		periodRets[actionDate] = (stockCandidates['periodReturn'] * stockCandidates['Weight']).sum()

	NetValue = handleRawPortReturn_TDayNetValue(DataFrame(periodRets), portfolioBuiltDates)
	ret = handleNetvalue_Return(NetValue, 1)

	return ret


def calNegQ(returns, wts, risk_aversion=2):
	wts = np.array(wts)
	Q = np.sum(returns.mean() * wts) - risk_aversion / 2 * np.dot(wts.T, np.dot(returns.cov(), wts))
	NegQ = - Q
	return NegQ

def cal_fctwt(returns, risk_aversion=2):

	if returns.empty:
		return Series(index=list(returns.columns), name='weights')

	func = lambda wts: calNegQ(returns, wts, risk_aversion)

	m = returns.shape[1]

	cons = ({'type': 'eq', 'fun': lambda x: np.sum(x) - 1})
	bnds = tuple((0, 1) for x in range(m))

	opts = sco.minimize(func, m * [1. / m, ], method='SLSQP', bounds=bnds, constraints=cons)
	# print(opts)
	return Series(opts.x, index=list(returns.columns), name='weights')
	# 注意：有可能出现优化不到最优值！！！


# dat is the current factorValues
def getCombinedFactor(dat, selFactors):

	fct_names = selFactors.index.tolist()
	tmp = dat.ix[:, ['Symbol', 'Date'] + fct_names].copy()

	tmp['combinedFactor'] = \
		(tmp.ix[:, 2:] * selFactors['factorOrderFlag'] * selFactors['factorWeight']).sum(axis=1)

	res = tmp.ix[:, ['Symbol', 'Date', 'combinedFactor']]
	return res


def getPoolTradeList(data_org, config):

	# "['MarketValue'] not in index"

	fct_name = config['fct_name']
	flag = config['flag']
	portfolioBuiltDates = config['portfolioBuiltDates']
	numofStocks = config['numofStocks']

	tmpDataOrg = data_org.ix[data_org['Date'].isin(portfolioBuiltDates), ['Symbol', 'Date', 'TradeStatus', fct_name]]
	tmpDataOrg.ix[:, fct_name] = flag * tmpDataOrg.ix[:, fct_name]
	# 统一因子方向

	tradeList = dict()
	for portfolioBuiltDate in portfolioBuiltDates:
		tmp = tmpDataOrg.ix[(tmpDataOrg['Date'].isin([portfolioBuiltDate])) & (tmpDataOrg['TradeStatus'] == 1), :]
		tmp = tmp.sort_values(by=[fct_name, 'Symbol'], ascending=[False, True])
		tradeList[portfolioBuiltDate] = list(tmp['Symbol'][:numofStocks])

	return tradeList


def get_TopStkRet(data_org, config):

	fct_name = config['fct_name']
	inSampleDates = config['inSampleDates']

	tradeList = getPoolTradeList(data_org, config)

	tmpData = data_org.ix[data_org['Date'].isin(inSampleDates), ['Symbol', 'Date', fct_name]].copy()
	tmpRet = evalPerformanceOfTopStk(tmpData, fct_name, tradeList, 2)

	tmpRet.columns = [fct_name]
	return tmpRet



def calFactorWeights(retDict, factorValues, config):

	selFactors = selectFactors(retDict, factorValues, config)

	# calculate factor weights
	selFactors['factorWeight'] = 0

	if selFactors.shape[0] == 1:
		selFactors['factorWeight'] = 1

	else:
		tmpFactorValues = factorValues.ix[:, ['Symbol', 'Date'] + selFactors['factorName'].tolist()]
		data_org = pd.merge(get_CoreGenData('StockTradeInfo'), tmpFactorValues, on=['Symbol', 'Date'])    # 这里Merge应该产生了问题, 有重复列名 Ben @20200116
		inSampleDates = config['inSampleDates']
		portfolioBuiltDates = get_PortfolioBuiltDates(inSampleDates)

		rets = DataFrame()
		numofStocks = config['numofStocks']

		for i in range(0, selFactors.shape[0]):
			# 根据factorOrderFlag更新因子值
			config = dict()
			config['numofStocks'] = numofStocks
			config['fct_name'] = selFactors['factorName'].tolist()[i]
			config['flag'] = selFactors['factorOrderFlag'].tolist()[i]
			config['portfolioBuiltDates'] = portfolioBuiltDates
			config['inSampleDates'] = inSampleDates

			rets = pd.concat([rets, get_TopStkRet(data_org, config)], axis=1, sort=True)

		# 因子合并，权重分配
		# 1. 测试mean-variance optimization on absolute returns
		wts = cal_fctwt(rets, risk_aversion=2)
		selFactors.index = selFactors['factorName']
		del selFactors['factorName']
		selFactors['factorWeight'] = wts

	return selFactors


# ============================================ Funs_BackTest =============================================== #

def assignStockWeights(stockScore, weight_type):
	stockScore = DataFrame(stockScore)
	if weight_type == 1:
		# equal weight
		stockScore['weight'] = 1 / stockScore.shape[0]
	else:
		stockScore['weight'] = np.log(2 + stockScore['score'].rank()) / np.log(
			2 + stockScore['score'].rank()).sum()
	return stockScore


# 可能出现除权除息，分红可能因为银行的进度，导致一定时间内不可用，这里先不考虑
def updatePortfolioBeforeMarketOpen(portfolio, actionDate):

	# load 分红数据
	dat = get_CoreGenData('AShareDividend')

	# 分红当天开盘前就应该处理，万得数据在分红当天就操作了
	tmpDividend = dat.ix[dat['Date'] == actionDate, :]

	if not tmpDividend.empty:

		# 事先最好检查一下dividendData,不能出现nan
		if np.any(dat.isnull()):
			logging.error('Dividend Data Exist NAN / None , Portfolio Update cannot be Done! ')
			return None

		tmpDividend.index = tmpDividend['Symbol'].tolist()
		tmpPosition = portfolio['position'].copy()
		tmpPosition['Symbol'] = tmpPosition.index.tolist()
		# 会多出一个Symbol列

		tmp = pd.merge(tmpPosition, tmpDividend, how='inner', on=['Symbol'])

		if not tmp.empty:
			tmp.index = tmp['Symbol'].tolist()
			# print('Dividend Occurs' + ': \n', tmp)
			cashDiv = (tmp['shares'] * tmp['CashRatio']).sum()
			tmpPosition.ix[tmp['Symbol'].tolist(), 'shares'] = \
				tmpPosition.ix[tmp['Symbol'].tolist(), 'shares'] * (1 + tmp['StkRatio'])

			del tmpPosition['Symbol']
			portfolio['position'] = tmpPosition
			# 注意这里没有去更改close， eodAmt值
			portfolio['cash'] = portfolio['cash'] + cashDiv
	# 分红的股再一次性全部卖掉时，零股可以卖掉

	return portfolio

# def checkPortfolio(portfolio, actionDate):
#
# 	if (portfolio['cash'] < 0) or (portfolio['asset'] < 0):
# 		logging.error('Cash and Asset cannot be < 0 !!!')
# 		raise ValueError
#
# 	if np.any(portfolio['position'].isnull()):
#
# 		# print(portfolio)
# 		# 遇到退市的股票，手动干预，去掉，用昨天收盘价
# 		tmp = portfolio['position'].ix[portfolio['position'].isnull().any(axis=1), :].copy()
# 		logging.warning('Portfolio Position Exists NaN !!!' +
# 						','.join(tmp.index.tolist()) + ' on ' + actionDate + '. ')
#
# 		tradeDates = get_CoreGenData('AShareCalendar')['Date'].tolist()
# 		preActionDate = tradeDates[tradeDates.index(actionDate) - 1]
# 		tmp['preClose'] = get_PFun(tmp.index.tolist(), preActionDate, 'Close')
#
# 		portfolio['cash'] = portfolio['cash'] + (1 - 1e-3 - 2e-4 - 1e-3) * (tmp['preClose'] * tmp['shares']).sum()
# 		portfolio['position'] = portfolio['position'].ix[~(portfolio['position'].isnull().any(axis=1)), :]
# 		portfolio['asset'] = portfolio['position']['eodAmt'].sum() + portfolio['cash']
#
# 	return portfolio

def checkPortfolio(portfolio, tradeDetail, actionDate):

	if (portfolio['cash'] < 0) or (portfolio['asset'] < 0):
		logging.error('Cash and Asset cannot be < 0 !!!')
		raise ValueError

	if np.any(portfolio['position'].isnull()):

		# print(portfolio)
		# 遇到退市的股票，手动干预，去掉，用昨天收盘价
		tmp = portfolio['position'].ix[portfolio['position'].isnull().any(axis=1), :].copy()
		logging.warning('Portfolio Position Exists NaN !!!' +
						','.join(tmp.index.tolist()) + ' on ' + actionDate + '. ')

		tradeDates = get_CoreGenData('AShareCalendar')['Date'].tolist()
		preActionDate = tradeDates[tradeDates.index(actionDate) - 1]
		tmp['preClose'] = get_PFun(tmp.index.tolist(), preActionDate, 'Close')

		portfolio['cash'] = portfolio['cash'] + (1 - 1e-3 - 2e-4 - 1e-3) * (tmp['preClose'] * tmp['shares']).sum()
		portfolio['position'] = portfolio['position'].ix[~(portfolio['position'].isnull().any(axis=1)), :]
		portfolio['asset'] = portfolio['position']['eodAmt'].sum() + portfolio['cash']

		# 把这种特殊情况加入到tradeDetail中
		tmpTradeDetail = tmp.copy()
		tmpTradeDetail['closePrice'] = tmp['preClose']
		tmpTradeDetail['tradePrice'] = (1 - 1e-3 - 2e-4 - 1e-3) * tmpTradeDetail['closePrice']
		tmpTradeDetail['tradeAmt'] = tmpTradeDetail['tradePrice'] * tmpTradeDetail['shares']
		tmpTradeDetail['flag'] = 'sell'
		tmpTradeDetail['shares'] = - tmpTradeDetail['shares']
		tmpTradeDetail['Date'] = preActionDate

		tmpTradeDetail = tmpTradeDetail.ix[:, ['weight', 'tradeAmt', 'closePrice', 'tradePrice', 'shares', 'flag', 'Date']]
		tradeDetail[preActionDate] = tmpTradeDetail

	return portfolio, tradeDetail



# 持仓股中出现ST，当天可以卖掉立即卖掉，可以通过isThisPeriod=0操作,先不处理这个！！！

def BKTCoreEngine(combinedFactors, config):

	logging.info('Starts to BackTest ... ')

	buyImpact = config['buyImpact']
	sellImpact = config['sellImpact']
	commissionFee = config['commissionFee']
	stampTax = config['stampTax']
	minCashRatio = config['minCashRatio']
	sharePerHand = config['sharePerHand']

	initAmount = config['initAmt']
	numofStocks = config['numofStocks']
	bktPriceType = config['bktPriceType']
	weightType = config['weightType']
	freqofTrade = config['freqofTrade']  # 季度，月度，周度调仓

	buyCostFactor = 1 + commissionFee + buyImpact
	sellCostFactor = 1 - sellImpact - commissionFee - stampTax

	bktPriceType = bktPriceType + 'Price'
	priceDict = {
		'openPrice': 'Open',
		'closePrice': 'Close',
		'vwapPrice': 'Vwap'
	}

	# load 必须的数据
	StockTradeInfo = get_CoreGenData('StockTradeInfo')
	refPeriodDates = get_CoreGenData('refPeriodDates')[freqofTrade]['tail']

	portfolio = {'cash': initAmount, 'position': None, 'asset': initAmount}
	# 如果出现卖不掉的情况，要备一些候补股票(一般来说，太迟了，月底就被卖了)
	# 要加一个portfolio_record纪录买卖情况，好为后续的归因分析做准备

	latestNetValue = []
	portfolioRecord = dict()

	tradeDates = StockTradeInfo['Date'].unique().tolist()

	# 确定调仓时间和回测时间段，暂时用combinedFactors的时间节点
	prePortfolioBuiltDates = combinedFactors['Date'].unique().tolist()

	prePortfolioBuiltDates = Series(prePortfolioBuiltDates)[Series(prePortfolioBuiltDates) < tradeDates[-1]].tolist()
	portfolioBuiltDates = [tradeDates[tradeDates.index(preDay) + 1] for preDay in prePortfolioBuiltDates]

	actionDatesStartIndex = tradeDates.index(portfolioBuiltDates[0])
	actionDatesEnd = Series([refPeriodDates[refPeriodDates.index(prePortfolioBuiltDates[-1]) + 1], tradeDates[-1]]).min()
	actionDatesEndIndex = tradeDates.index(actionDatesEnd)

	actionDates = tradeDates[actionDatesStartIndex:(actionDatesEndIndex + 1)]

	# 开始模拟回测
	BKTDates = []
	tradeDetail = dict()

	for actionDate in actionDates:
		# print(actionDate)

		BKTDates.append(actionDate)
		if actionDate in portfolioBuiltDates:

			preActionDate = tradeDates[tradeDates.index(actionDate) - 1]
			stockScore = combinedFactors.ix[combinedFactors['Date'] == preActionDate, :].sort_values(by=
															['combinedFactor', 'Symbol'], ascending=[False, True])
			# 按score排序
			stockScore.index = stockScore['Symbol'].tolist()
			stockScore = stockScore['combinedFactor']
			stockScore.name = 'score'

			stockList = list(stockScore.index)
			# print('Model is trained and portfolio adjusted on: ' + actionDate)

			tempTradeStatus = getTradeStatus(stockList, actionDate)
			selTradableStock = tempTradeStatus[tempTradeStatus == 1][:numofStocks]   # !!!! 确定股票 !!!!
			selTradableStockList = selTradableStock.index.tolist()

			if len(latestNetValue) == 0:
				# 取前numofStocks个可以交易的
				selStockScore = stockScore[selTradableStockList]
				selStock = assignStockWeights(selStockScore, weightType)   # !!!! 确定权重 !!!!

				selStock[bktPriceType] = get_PFun(selStock.index.tolist(), actionDate, priceDict[bktPriceType])
				selStock['tradePrice'] = buyCostFactor * selStock[bktPriceType]
				availCash = (1 - minCashRatio) * portfolio['cash']
				selStock['shares'] = np.floor(availCash * selStock['weight'] / selStock['tradePrice'] / sharePerHand) * sharePerHand
				selStock['tradeAmt'] = selStock['tradePrice'] * selStock['shares']
				selStock['isThisPeriod'] = 1

				# -- record trade history
				tradeDetail[actionDate] = selStock.ix[:, ['weight', 'tradeAmt', bktPriceType, 'tradePrice', 'shares']]
				tradeDetail[actionDate]['flag'] = 'buy'
				tradeDetail[actionDate]['Date'] = actionDate

				# -- record eod position
				selStock['close'] = get_PFun(selTradableStockList, actionDate, 'Close')
				selStock['eodAmt'] = selStock['close'] * selStock['shares']

				portfolio['position'] = selStock.ix[:, ['weight', 'eodAmt', 'close', 'shares', 'isThisPeriod']]
				portfolio['cash'] = portfolio['cash'] - selStock['tradeAmt'].sum()
				portfolio['asset'] = portfolio['cash'] + portfolio['position']['eodAmt'].sum()

				portfolio, tradeDetail = checkPortfolio(portfolio, tradeDetail, actionDate)
				latestNetValue.append(portfolio['asset'] / initAmount)
				portfolioRecord[actionDate] = portfolio.copy()
				print('BackTest Date: ' + actionDate + ' , totalvalue: ' + repr(portfolio['asset'] / initAmount))

			else:
				# ------- 盘前分析 ------ #
				portfolio = updatePortfolioBeforeMarketOpen(portfolio, actionDate)
				holdedPosition = portfolio['position'].copy()

				# 获得持仓股当天的状态,前numofStocks个可交易的，但是也包含不可交易的
				candidates = stockScore[:(stockList.index(selTradableStockList[-1]) + 1)]
				# 把这个和持仓股票比较，如果持仓股票中有停牌，并在candidate中，保留这个股票

				# 获取持仓中股票今日状态
				stockHoldedList = holdedPosition.index.tolist()
				stockHoldedStatus = getTradeStatus(stockHoldedList, actionDate)

				# 是否存在停牌股票
				if any(stockHoldedStatus == 0):
					stockSuspendList = stockHoldedStatus[stockHoldedStatus == 0].index.tolist()

					suspendStockInCandidate = list(set(stockSuspendList) & set(candidates.index.tolist()))
					suspendStockNotInCandidate = list(set(stockSuspendList) - set(candidates.index.tolist()))

					# print('suspendStockInCandidate: ', suspendStockInCandidate)
					# print('suspendStockNotInCandidate: ', suspendStockNotInCandidate)

					# 遇到在candidates中的股票，新加入的个数减去,stockToHoldList包含持仓中的可交易股票
					stockToHoldList = selTradableStock[:(numofStocks - len(suspendStockInCandidate))].index.tolist()
					# 包含了在持仓中可交易的股票
					# print('stock to hold:' + repr(len(stockToHoldList)))

					holdedPosition.ix[suspendStockNotInCandidate, 'isThisPeriod'] = 0
					holdedPosition.ix[suspendStockInCandidate, 'isThisPeriod'] = 1

					# 需要处理上期isThisPeriods=0，这次又重新出现在新的选股中，卖掉还是留着？

				else:
					stockSuspendList = []
					stockToHoldList = selTradableStock.index.tolist()

				stockToHold = stockScore[stockToHoldList]
				# assign weights to stockToHold
				stockToHold = assignStockWeights(stockToHold, weightType)

				# 更新vwap, vwapAmt
				holdedPosition[bktPriceType] = get_PFun(holdedPosition.index.tolist(), actionDate, priceDict[bktPriceType])
				priceAmt = bktPriceType[:-5] + 'Amt'
				holdedPosition[priceAmt] = holdedPosition['shares'] * holdedPosition[bktPriceType]

				# 5. 计算可用资金，包括现在的现金资产和可以卖出的股票价值
				tradableHoldedStockList = stockHoldedStatus[stockHoldedStatus == 1].index.tolist()
				tradablePosition = holdedPosition.ix[tradableHoldedStockList, :]
				availCash = (1 - minCashRatio) * (portfolio['cash'] + tradablePosition[priceAmt].sum() * sellCostFactor)

				# 6. 算资金轧差，不要算股票轧差
				stockToHold['targetAmt'] = (availCash / buyCostFactor) * stockToHold['weight']
				# 注意这里要加buyCostFactor,因为买有成本
				chgAmount = stockToHold.ix[:, ['weight', 'targetAmt']].join(tradablePosition.ix[:, [priceAmt, 'shares']], how='outer')
				chgAmount.columns = ['weight', 'targetAmt', priceAmt, 'orgShares']
				# 应该是在这里引入了weight=0的情况，这里先不管，后面可以不要这个列！！！


				# nan -> 0
				chgAmount = chgAmount.fillna(value=0)
				chgAmount['deltaAmt'] = chgAmount['targetAmt'] - chgAmount[priceAmt]
				chgAmount['shares'] = chgAmount['orgShares']

				# 8. 卖出deltaAmt为负数的股票对应的轧差
				chgAmount[bktPriceType] = get_PFun(chgAmount.index.tolist(), actionDate, priceDict[bktPriceType])
				chgAmount['chgShares'] = np.floor(abs(chgAmount['deltaAmt']) / chgAmount[bktPriceType] / sharePerHand) * sharePerHand

				# I. 负数要处理一下，np.floor(-1.2) = -2，这样卖的多一点,不可实现
				chgAmount.ix[chgAmount['deltaAmt'] < 0, 'chgShares'] = - chgAmount.ix[chgAmount['deltaAmt'] < 0, 'chgShares']
				# II. 有可能由于分红，多出了一些零散股，如果要卖完，是不用一定为100的倍数
				chgAmount.ix[chgAmount['targetAmt'] == 0, 'chgShares'] = \
					chgAmount.ix[chgAmount['targetAmt'] == 0, 'deltaAmt'] / chgAmount.ix[chgAmount['targetAmt'] == 0, bktPriceType]

				chgAmount['actChgShares'] = 0

				# handle sell position
				positionSell = chgAmount.ix[chgAmount['chgShares'] < 0, :].copy()
				positionSell['tradePrice'] = sellCostFactor * positionSell.ix[:, bktPriceType]
				portfolio['cash'] = portfolio['cash'] + (abs(positionSell['chgShares']) * positionSell['tradePrice']).sum()

				positionSell['actChgShares'] = positionSell['chgShares']
				positionSell['shares'] = positionSell['orgShares'] + positionSell['actChgShares']
				# 这种赋值写法会报warning！！！

				chgAmount.ix[positionSell.index.tolist(), 'actChgShares'] = positionSell['chgShares']
				chgAmount.ix[positionSell.index.tolist(), 'shares'] = positionSell['shares']

				# 加入交易信息tradeDetail
				tmpTradeDetailSell = positionSell.ix[:, ['weight', 'actChgShares', bktPriceType, 'tradePrice']]
				tmpTradeDetailSell['tradeAmt'] = abs(tmpTradeDetailSell['actChgShares']) * tmpTradeDetailSell['tradePrice']
				tmpTradeDetailSell = tmpTradeDetailSell.ix[:, ['weight', 'tradeAmt', bktPriceType, 'tradePrice', 'actChgShares']]
				tmpTradeDetailSell.columns = ['weight', 'tradeAmt', bktPriceType, 'tradePrice', 'shares']
				tmpTradeDetailSell['flag'] = 'sell'
				tmpTradeDetailSell['Date'] = actionDate

				# 9. 计算买入仓位并买入
				positionBuy = chgAmount.ix[chgAmount['chgShares'] > 0, :].copy()
				positionBuy['tradePrice'] = buyCostFactor * positionBuy[bktPriceType]
				positionBuy['actChgShares'] = positionBuy['chgShares']
				positionBuy['shares'] = positionBuy['orgShares'] + positionBuy['actChgShares']

				chgAmount.ix[positionBuy.index.tolist(), 'actChgShares'] = positionBuy['chgShares']
				chgAmount.ix[positionBuy.index.tolist(), 'shares'] = positionBuy['shares']
				# 这里买一定能买到chgShares，因为加入了minCashRatio
				portfolio['cash'] = portfolio['cash'] - (positionBuy['tradePrice'] * positionBuy['actChgShares']).sum()

				# 加入交易信息tradeDetail
				tmpTradeDetailBuy = positionBuy.ix[:, ['weight', 'actChgShares', bktPriceType, 'tradePrice']]
				tmpTradeDetailBuy['tradeAmt'] = tmpTradeDetailBuy['actChgShares'] * tmpTradeDetailBuy['tradePrice']
				tmpTradeDetailBuy = tmpTradeDetailBuy.ix[:, ['weight', 'tradeAmt', bktPriceType, 'tradePrice', 'actChgShares']]
				tmpTradeDetailBuy.columns = ['weight', 'tradeAmt', bktPriceType, 'tradePrice', 'shares']
				tmpTradeDetailBuy['flag'] = 'buy'
				tmpTradeDetailBuy['Date'] = actionDate

				tradeDetail[actionDate] = pd.concat([tmpTradeDetailSell, tmpTradeDetailBuy], sort=True)

				# 10. 更新仓位信息
				updatedPosition = chgAmount.ix[chgAmount['shares'] > 1e-1, ['weight', 'shares']]
				# 这里targetAmt = 0 的股票是一定能够卖完的，然后持仓中就不在包括
				updatedPosition['close'] = get_PFun(updatedPosition.index.tolist(), actionDate, 'Close')
				updatedPosition['eodAmt'] = updatedPosition['shares'] * updatedPosition['close']
				updatedPosition['isThisPeriod'] = 1

				# 加上之前持仓中不可交易的股票
				if len(stockSuspendList) > 0:
					stockSuspendPosition = holdedPosition.ix[stockSuspendList, :]
					# 更新eodAmt
					stockSuspendPosition['close'] = get_PFun(stockSuspendPosition.index.tolist(), actionDate, 'Close')
					stockSuspendPosition['eodAmt'] = stockSuspendPosition['shares'] * stockSuspendPosition['close']
					updatedPosition = pd.concat([updatedPosition, stockSuspendPosition], sort=True)

				#更新portfolio其他信息
				portfolio['position'] = updatedPosition.ix[:, ['weight', 'eodAmt', 'close', 'shares', 'isThisPeriod']]
				portfolio['asset'] = portfolio['cash'] + portfolio['position']['eodAmt'].sum()

				portfolio, tradeDetail = checkPortfolio(portfolio, tradeDetail, actionDate)
				latestNetValue.append(portfolio['asset'] / initAmount)
				portfolioRecord[actionDate] = portfolio.copy()
				print('BackTest Date: ' + actionDate + ' , totalvalue: ' + repr(portfolio['asset'] / initAmount))

				# if actionDate == '20100201' or actionDate == '20130902':
				# 	print('hello')

		else:

			# ------- 盘前分析 ------ #
			portfolio = updatePortfolioBeforeMarketOpen(portfolio, actionDate)

			# 更新价格和股票价值
			holdedPosition = portfolio['position'].copy()
			holdedPosition['close'] = get_PFun(holdedPosition.index.tolist(), actionDate, 'Close')
			holdedPosition['eodAmt'] = holdedPosition['shares'] * holdedPosition['close']
			portfolio['position'] = holdedPosition

			# 对持仓检查标签，是否可以买卖
			if any(holdedPosition['isThisPeriod'] == 0):
				temp = holdedPosition.ix[holdedPosition['isThisPeriod'] == 0, :]
				tempTradeStatus = getTradeStatus(temp.index.tolist(), actionDate)

				# 可以卖了
				if any(tempTradeStatus == 1):
					tempToSellList = tempTradeStatus[tempTradeStatus == 1].index.tolist()
					# sell stocks with isThisPeriod = 0 and tempTradeStatus = 1
					tmpBktPrice = get_PFun(tempToSellList, actionDate, priceDict[bktPriceType])
					portfolio['cash'] = portfolio['cash'] + sellCostFactor * (holdedPosition.ix[tempToSellList, 'shares'] * tmpBktPrice).sum()
					remainStockList = [stk for stk in holdedPosition.index.tolist() if stk not in tempToSellList]
					portfolio['position'] = holdedPosition.ix[remainStockList, :]

					# -- trade record history
					tmpTradeDetail = holdedPosition.ix[tempToSellList, :]
					tmpTradeDetail[bktPriceType] = tmpBktPrice
					tmpTradeDetail['tradePrice'] = sellCostFactor * tmpTradeDetail[bktPriceType]
					tmpTradeDetail['tradeAmt'] = tmpTradeDetail['tradePrice'] * tmpTradeDetail['shares']
					# 算换手率,应该用shares * vwap, tradeAmt包含了其他费用

					# 卖出shares设置为负数
					tmpTradeDetail['shares'] = - tmpTradeDetail['shares']

					tmpTradeDetail['flag'] = 'sell'
					tmpTradeDetail['Date'] = actionDate

					tradeDetail[actionDate] = tmpTradeDetail.ix[:, ['weight', 'tradeAmt', bktPriceType, 'tradePrice', 'shares', 'flag', 'Date']]

			portfolio['asset'] = portfolio['cash'] + portfolio['position']['eodAmt'].sum()

			portfolio, tradeDetail = checkPortfolio(portfolio, tradeDetail, actionDate)
			latestNetValue.append(portfolio['asset'] / initAmount)
			portfolioRecord[actionDate] = portfolio.copy()
			print('BackTest Date: ' + actionDate + ' , totalvalue: ' + repr(portfolio['asset'] / initAmount))

	latestNetValue = Series(latestNetValue, index=BKTDates, name='NetValue')

	ls = {
		'NetValue': latestNetValue,
		'portfolioRecord': portfolioRecord,
		'tradeDetail': tradeDetail,
		'portfolioBuiltDates': portfolioBuiltDates,
		'bktPriceType': bktPriceType

	}

	logging.info('BackTest is Done !!!')
	return ls


# portfolioAnalysis
def portfolioAnalysis(bktResult):

	logging.info('Starts to Analyze Portfolio ... ')
	# 对portfolio进行分析
	stockCalendar = get_CoreGenData('AShareCalendar')

	NetValue = bktResult['NetValue']
	portfolioRecord = bktResult['portfolioRecord']
	tradeDetail = bktResult['tradeDetail']
	portfolioBuiltDates = bktResult['portfolioBuiltDates']
	bktPriceType = bktResult['bktPriceType']

	BKTDates = NetValue.index.tolist()
	tradeDates = stockCalendar['Date'].tolist()

	# 1. 仓位信息
	positionStatus = DataFrame(np.zeros([len(BKTDates), 2]), columns=['stockProportion', 'cashProportion'], index=BKTDates)
	for tmpDate in BKTDates:
		positionStatus.ix[tmpDate, :] = [round(1 - portfolioRecord[tmpDate]['cash'] / portfolioRecord[tmpDate]['asset'], 4),
								  round(portfolioRecord[tmpDate]['cash'] / portfolioRecord[tmpDate]['asset'], 4)]

	# 2. 股票持有期
	tmp = DataFrame()
	numofStkPerDay = []
	for tmpDate in BKTDates:
		tmp = pd.concat([tmp, portfolioRecord[tmpDate]['position']], sort=True)
		numofStkPerDay.append(portfolioRecord[tmpDate]['position'].shape[0])

	numofStkPerDay = Series(numofStkPerDay, index=BKTDates, name='numofStkPerDay')


	stockSet = tmp.index.unique().tolist()
	stockLastDates = dict()

	for tmpStock in stockSet:
		stockLastDates[tmpStock] = []

	for tmpDate in BKTDates:
		stocksIntmpDate = portfolioRecord[tmpDate]['position'].index.tolist()
		for tmpStock in stocksIntmpDate:
			stockLastDates[tmpStock].append(tmpDate)

	stockStat = DataFrame(columns=['Symbol', 'EntryDate', 'ExitDate'])
	for i in range(len(stockSet)):
		tmpStockLastDates = stockLastDates[stockSet[i]]

		# 这里后面可能可以改进，用计算因子的方法！！！
		func = lambda x: np.where(np.array(tradeDates) == x)[0]
		ep = Series(tmpStockLastDates).apply(func).astype('int')
		ep.index = tmpStockLastDates


		epPos = Series(ep[1:].values - ep[:-1].values, index=ep[1:].index)
		if all(epPos == 1):
			entryDate = [tmpStockLastDates[0]]
			exitDate = [tmpStockLastDates[-1]]
		else:
			tmp = np.where(epPos != 1)[0]
			entryDate = [' '] * (len(tmp) + 1)
			exitDate = [' '] * (len(tmp) + 1)

			for j in range(0, len(tmp) + 1):
				if j == 0:
					entryDate[j] = ep.index.tolist()[0]
					exitDate[j] = epPos.index.tolist()[tmp[j] - 1]
				elif j < len(tmp):
					entryDate[j] = epPos.index.tolist()[tmp[j-1]]
					exitDate[j] = epPos.index.tolist()[tmp[j] - 1]
				else:
					entryDate[j] = epPos.index.tolist()[tmp[j - 1]]
					exitDate[j] = ep.index.tolist()[-1]

		tmpStockStat = DataFrame(columns=['Symbol', 'EntryDate', 'ExitDate'])
		for k in range(len(entryDate)):
			tmpStockStat = pd.concat([tmpStockStat, DataFrame([stockSet[i], entryDate[k], exitDate[k]],
															  index=['Symbol', 'EntryDate', 'ExitDate']).T])

		stockStat = pd.concat([stockStat, tmpStockStat], sort=True)

	stockStat = stockStat.sort_values(by=['Symbol', 'EntryDate'])
	stockStat = stockStat.reset_index().ix[:, ['Symbol', 'EntryDate', 'ExitDate']]
	# 或者就按每个月计算这些指标！！！！


	# 3. 持有期天数
	func = lambda x: np.where(np.array(tradeDates) == x)[0]
	stockStat['LastTDays'] = stockStat['ExitDate'].apply(func).astype('int') - stockStat['EntryDate'].apply(func).astype('int')

	# 4. 持有期涨跌幅
	stockStat['PctChg'] = 0

	gen_StockTradeInfo()
	for i in range(stockStat.shape[0]):
		stockStat.ix[i, 'PctChg'] = round(get_PFun([stockStat.ix[i, 'Symbol']], stockStat.ix[i, 'ExitDate'], 'Close')[0] / \
									get_PFun([stockStat.ix[i, 'Symbol']], stockStat.ix[i, 'EntryDate'], 'Close')[0] - 1, 4)


	# 5. 持有期胜率
	winProb = round(stockStat.ix[stockStat['PctChg'] > 0, :].shape[0] / stockStat.shape[0], 3)

	# 7. 换手率
	tradeHist = DataFrame()
	for tradeDay in list(tradeDetail.keys()):
		tradeHist = pd.concat([tradeHist, tradeDetail[tradeDay]], sort=True)

	# 月度换手率
	PBD = portfolioBuiltDates
	turnover_M = Series(index=PBD[:-1], name='turnover_M')
	for i in range(len(PBD) - 1):
		PortfolioDateStart = PBD[i]
		PortfolioDateEnd = tradeDates[tradeDates.index(PBD[i+1]) - 1]
		tmp = tradeHist.ix[(tradeHist['Date'] >= PortfolioDateStart) & (tradeHist['Date'] <= PortfolioDateEnd), :]

		avgAsset = (portfolioRecord[PortfolioDateStart]['asset'] + portfolioRecord[PortfolioDateEnd]['asset']) / 2
		turnover_M[PortfolioDateStart] = (tmp['shares'] * tmp[bktPriceType]).sum() / avgAsset
	# 第一次建仓，不算进去


	# 6. 所选个股行业分布
	ls = {
		'positionStatus': positionStatus,
		'numofStkPerDay': numofStkPerDay,
		'stockStat': stockStat,
		'winProb': winProb,
		'turnover_M': turnover_M
	}

	logging.info('Portfolio Analysis is Done! Relative Indicators are returned! ')
	return ls


# ------------------------------------- Simple BKT ------------------------------ #
def calStatistics(dat):

	'''
	:param dat: dataframe or Series
	:return: statistical indicators
	'''

	if isinstance(dat, DataFrame) or isinstance(dat, Series):
		dat = dat.astype('float32')
		mu = dat.mean()
		std = dat.std()
		corr = dat.corr()
		confidenceLevel = norm.ppf(0.05, mu, std)

		return {'mean': mu,
				'std': std,
				'corr': corr,
				'confidenceLevel': confidenceLevel
				}

	else:
		logging.warning('Input is not Pandas DataType!')
		return []


def calRollingStatistics(dat, rollingPeriod=60):

	'''
	:param dat: DataFrame, usually return
	:param rollingPeriod: lookback window
	:return: rolling statistics
	'''

	if isinstance(dat, DataFrame) or isinstance(dat, Series):

		dat = DataFrame(dat).astype('float32')
		rollingMean = dat.rolling(window=rollingPeriod).mean()
		rollingStd = dat.rolling(window=rollingPeriod).std()

		rollingCov = DataFrame()
		rollingCorr = DataFrame()
		columns = dat.columns.tolist()

		if dat.shape[1] > 1:
			for i in range(dat.shape[1] - 1):
				for j in range(i+1, dat.shape[1]):
					tmpCov = dat.ix[:, [i, j]].rolling(window=rollingPeriod).cov()


					tmpCov = DataFrame(tmpCov.ix[range(1, tmpCov.shape[0], 2), 0].values,
									   index=dat.index, columns=[columns[i] + '_' + columns[j]])
					rollingCov = pd.concat([rollingCov, tmpCov], axis=1, sort=True)


					tmpCorr = dat.ix[:, [i, j]].rolling(window=rollingPeriod).corr()
					tmpCorr = DataFrame(tmpCorr.ix[range(1, tmpCorr.shape[0], 2), 0].values,
										index=dat.index, columns=[columns[i] + '_' + columns[j]])
					rollingCorr = pd.concat([rollingCorr, tmpCorr], axis=1, sort=True)

		return {'rollingMean': rollingMean,
				'rollingStd': rollingStd,
				'rollingCov': rollingCov,
				'rollingCorr': rollingCorr
				}

	else:
		logging.warning('Input is not Pandas DataType!')
		return []



def calRollingBeta(s1, s2, rollingPeriod=60):
	'''
	:param s1: 策略收益,Series or DataFrame
	:param s2: 基准收益,Series or DataFrame
	:param rollingPeriod: 回溯窗口长度
	:return: 滚动Beta, 用于检查策略收益与基准收益在各阶段的相关性
	'''

	if len(s1) != len(s2):
		return []

	dat = pd.concat([DataFrame(s1), DataFrame(s2)], axis=1, sort=True)

	rollingCov = dat.rolling(window=rollingPeriod).cov()
	rollingCov = DataFrame(rollingCov.ix[range(1, rollingCov.shape[0], 2), 0].values,
					   index=dat.index, columns=['Cov'])

	rollingVar = DataFrame(s2).rolling(window=rollingPeriod).var()
	rollingBeta = DataFrame(rollingCov.values / rollingVar.values, index=dat.index, columns=['Beta'])

	return rollingBeta



def calRollingSharpe(s1, rollingPeriod=60, rf_rate=0.03, cycle=244):

	'''

	:param s1: 策略收益， Series or Dataframe
	:param rollingPeriod: 回溯窗口长度
	:param rf_rate: 无风险利率
	:return: 滚动Sharpe, 用于查看策略是否在各阶段表现一致
	'''

	dat = DataFrame(s1)
	rollingMean = dat.rolling(window=rollingPeriod).mean()
	rollingStd = dat.rolling(window=rollingPeriod).std()

	rollingSharpe = DataFrame((rollingMean.values * cycle - rf_rate) / (rollingStd.values * (cycle ** (1/2))),
							  index=dat.index, columns=['Sharpe'])
	return rollingSharpe

# old version
def NV2Ret(dat, type=1):

	'''
	:param dat: data needed to be transform, Series or DataFrame
	:param type: 1 (netvalue -> return) or 2 (return -> netvalue)
	:return: transformed data (return or netvalue)
	'''

	if type == 1:
		# netvalue to return
		if len(dat.shape) > 1:
			# input is a dataframe
			fd_return = np.zeros(dat.shape[1])
			fd_return = fd_return[:, np.newaxis].T
			res = dat.values[1:, :] / dat.values[:-1, :] - 1
			res = np.concatenate([fd_return, res])
			res = DataFrame(res, index=dat.index, columns=dat.columns)
		elif len(dat.shape) == 1:
			# input is a Series
			fd_return = np.zeros(1)
			res = dat.values[1:] / dat.values[:-1] - 1
			res = np.concatenate([fd_return, res])
			res = Series(res, index=dat.index, name=dat.name)
	elif type == 2:
		# return to netvalue
		res = (dat + 1).cumprod()

	return res


def nav2return(dat, type=1):

	'''
	:param dat: data needed to be transform, Series or DataFrame
	:param type: 1 (netvalue -> return) or 2 (return -> netvalue)
	:return: transformed data (return or netvalue)
	'''

	if type == 1:
		# nav to return
		if len(dat.shape) > 1:
			# input is a dataframe
			fd_return = np.zeros(dat.shape[1])
			fd_return = fd_return[:, np.newaxis].T
			res = dat.values[1:, :] / dat.values[:-1, :] - 1
			res = np.concatenate([fd_return, res])
			res = DataFrame(res, index=dat.index, columns=dat.columns)
		elif len(dat.shape) == 1:
			# input is a Series
			fd_return = np.zeros(1)
			res = dat.values[1:] / dat.values[:-1] - 1
			res = np.concatenate([fd_return, res])
			res = Series(res, index=dat.index, name=dat.name)
	elif type == 2:
		# return to nav
		res = (dat + 1).cumprod()

	return res





def PeriodRet2TDailyNV(PeriodRet, chgDates):

	'''

	:param PeriodRet: period accumulative return
	:param chgDates: portfolio change dates
	:return: Trading day netvalue
	'''

	PeriodRet = DataFrame(PeriodRet)
	actionDates = PeriodRet.index.tolist()
	TDailyNV = DataFrame(np.nan * np.zeros(PeriodRet.shape),
						index=PeriodRet.index, columns=PeriodRet.columns)

	# the former PortfolioBuiltDate and # Initialize 1
	periodStartDate = chgDates[0]
	periodEndValue = Series(np.ones(PeriodRet.shape[1]), index=PeriodRet.columns)

	for chgDate in chgDates:
		if actionDates.index(chgDate) != 0:
			periodStartIndex = actionDates.index(periodStartDate)
			periodEndIndex = actionDates.index(chgDate)
			TDailyNV.ix[periodStartIndex:periodEndIndex+1, :] = \
				(PeriodRet.ix[periodStartIndex:periodEndIndex+1, :] + 1) * periodEndValue

			# Update periodStartDate and periodEndValue
			periodEndValue = TDailyNV.ix[periodEndIndex, :]
			if chgDate < actionDates[-1]:
				periodStartDate = actionDates[actionDates.index(chgDate) + 1]

	# Remaining days
	if chgDate < actionDates[-1]:
		TDailyNV.ix[actionDates.index(periodStartDate):, :] = \
					(PeriodRet.ix[actionDates.index(periodStartDate):, :] + 1) * periodEndValue

	return TDailyNV


def TDRet2PeriodRet(ret, chgDates):
	'''

	:param ret: daily returns
	:param chgDates: chgDates
	:return: period returns in each period interval
	'''

	periodRet = DataFrame(index=chgDates, columns=ret.columns)
	actionDates = ret.index.tolist()
	for chgDate in chgDates:
		if chgDates.index(chgDate) == 0:
			periodRet.ix[chgDate, :] = 0
		else:
			tmp = ret.ix[chgDates[chgDates.index(chgDate) - 1]:chgDate, :].copy()
			tmp.ix[0, :] = 0
			periodRet.ix[chgDate, :] = (((tmp + 1).cumprod()).ix[-1, :] - 1)

	return periodRet



def simpleBKTCore(dat, config):

	'''
	:param dat: netvalue of different index/stocks/assets
	:param config: includes chgDates and weights on chgDates
	:return: Daily returns
	'''

	chgDates = config['chgDates']
	weights = config['weights']
	totCost = config['totCost']

	if dat.shape[1] != weights.shape[1]:
		logging.error('Dimensions of input netvalue and weights do not Match! ')
		return []

	actionDates = dat.index.tolist()
	# Adjust actionDates based on chgDates
	actionDates = actionDates[actionDates.index(chgDates[0]):]

	columns = dat.columns.tolist()
	PRet = DataFrame(np.zeros((len(actionDates), len(columns))), index=actionDates, columns=columns)
	PRet['PRet'] = 0

	for actionDate in actionDates:

		# 看是不是调仓第一天，用收盘价的回测，当天收益为0，但是后期调仓日也是有收益的！
		if actionDates.index(actionDate):
			# print('BackTest Date: ' + actionDate)
			PRet.ix[actionDate, columns] = (1 - totCost) * dat.ix[actionDate, :] / dat.ix[chgDate, :] - 1
			PRet.ix[actionDate, 'PRet'] = (PRet.ix[actionDate, columns] * wts).sum()

			# 这里可以把各资产都加进去，加账户


		if actionDate in chgDates:

			# if not actionDates.index(actionDate):
			# 	print('BackTest Date: ' + actionDate)

			chgDate = actionDate
			wts = weights.ix[chgDate, :]

	# handle raw results
	NetValue = PeriodRet2TDailyNV(PRet['PRet'], chgDates)
	ret = NV2Ret(NetValue, 1)

	return {'ret': ret, 'config': config}



def dfPlot(dat, title_str=[], num=8):

	'''
	:param dat: a dataframe to draw
	:param title_str: title
	:return: figure
	'''

	dat = DataFrame(dat)
	XTicks = Series(np.linspace(0, dat.shape[0] - 1, num)).astype('int').tolist()

	dat.plot()
	plt.xticks(XTicks, list(dat.index[XTicks]))
	if isinstance(title_str, str):
		plt.title(title_str)
	plt.grid()
	plt.show()


def dfPlotSave(dat, title_str=[], filename='test', show_flag=1):

	'''
	:param dat: a dataframe to draw
	:param title_str: title
	:return: figure
	'''

	dat = DataFrame(dat)
	XTicks = Series(np.linspace(0, dat.shape[0] - 1, 5)).astype('int').tolist()

	dat.plot(figsize=(10, 7))
	plt.xticks(XTicks, list(dat.index[XTicks]))
	if isinstance(title_str, str):
		plt.title(title_str)
	plt.grid()
	plt.savefig(filename + '.png')
	if show_flag == 1:
		plt.show()
	plt.close()


# # 绘制Time Series图
# def tsPlot(y, lags=30):
# 	'''
# 	:param y: time series, series
# 	:param lags: acf, pacf lags
# 	'''
# 	if not isinstance(y, pd.Series):
# 		y = Series(y)
#
# 	XTicks = Series(np.linspace(0, y.shape[0] - 1, 5)).astype('int').tolist()
#
# 	figsize=(10,8)
# 	fig = plt.figure(figsize=figsize)
# 	layout = (3, 2)
#
# 	ts_ax = plt.subplot2grid(layout, (0, 0), colspan=2)
# 	acf_ax = plt.subplot2grid(layout, (1, 0))
# 	pacf_ax = plt.subplot2grid(layout, (1, 1))
# 	qq_ax = plt.subplot2grid(layout, (2, 0))
# 	pp_ax = plt.subplot2grid(layout, (2, 1))
#
# 	y.plot(ax=ts_ax)
# 	ts_ax.grid()
# 	ts_ax.set_xticks(XTicks)
# 	ts_ax.set_xticklabels(list(y.index[XTicks]))
#
# 	ts_ax.set_title('Time Series Analysis Plots')
#
# 	plot_acf(y, lags=lags, ax=acf_ax)
# 	plot_pacf(y, lags=lags, ax=pacf_ax)
#
# 	sm.qqplot(y, line='s', ax=qq_ax)#QQ图检验是否是正太分布
# 	qq_ax.set_title('QQ Plot')
# 	qq_ax.grid()
# 	scs.probplot(y, sparams=(y.mean(), y.std()), plot=pp_ax)
# 	pp_ax.grid()
#
# 	plt.tight_layout()
# 	plt.show()
#
# 	return True

def diffPlot(dat, path):
	'''
	:param dat: dat is nx2 dataframe, the benchmark is at the end
	:param path: path to save image
	:return:
	'''

	colnames = dat.columns.tolist()
	filename = path + colnames[0] + '_' + colnames[1]
	# plot difference
	datDiff = (dat.ix[:, 0] - dat.ix[:, 1]) * 100

	XTicks = Series(np.linspace(0, dat.shape[0] - 1, 8)).astype('int').tolist()

	plt.figure(figsize=(10, 6))

	plt.subplot(211)
	plt.plot(dat)
	plt.xticks(XTicks, list(dat.index[XTicks]))
	plt.grid()
	plt.title(colnames[0] + ', ' + colnames[1])
	plt.legend(colnames)

	plt.subplot(212)
	plt.bar(range(len(datDiff)), datDiff)
	plt.xticks(XTicks, list(datDiff.index[XTicks]))
	plt.legend([colnames[0] + ' - ' + colnames[1]])

	plt.savefig(filename + '.png')
	plt.close()

	return True


def simpleBKT(dat, config):

	'''

	:param dat: 资产净值或者价格序列
	:param config: config['weights']索引只是chgDates
	:return: 返回回测结果，包括ret, 风险指标，权重
	'''

	bktRes = simpleBKTCore(dat, config)
	ret = bktRes['ret']

	ann = strategyAnalysis()
	indicators = ann.Get_BasicIndictors(ret)

	weights = pd.concat([config['weights'], ret.ix[config['weights'].index.tolist()[0]:, :]],
							 axis=1, join='outer', sort=True).ix[:, :config['weights'].shape[1]]
	weights = weights.fillna(method='ffill')

	rollingSharpe = calRollingSharpe(ret)
	return {
		'ret': ret,
		'indicators': indicators,
		'weights': weights,
		'rollingSharpe': rollingSharpe
	}



# def weightStackPlot(weights):
#
# 	x = np.array(range(0, weights.shape[0]))
# 	y = weights.values.astype('float32')
#
# 	XTicks = Series(np.linspace(0, weights.shape[0] - 1, 5)).astype('int').tolist()
#
# 	plt.stackplot(x, y.T)
# 	plt.xticks(XTicks, list(weights.index[XTicks]))
# 	plt.legend(weights.columns.tolist())
# 	plt.show()

# ---------- 计算换手率 ----------- #
def calTurnOverBasedWeights(wts, tickerRet, tickerNV):
	'''
	:param wts: weights, a dataframe with chgDates as index
	:param tickerRet: monthly returns with chgDates as index
	:param tickerNV: netvalue, a dataframe with chgDates as index
	:return:
	'''

	if (wts.shape[0] != tickerNV.shape[0]) or (wts.shape[0] != tickerRet.shape[0]):
		logging.info('Length of Weights, NetValue and Return at chgDates Do NOT Match!')
		return []

	dat = pd.concat([wts, tickerNV], axis=1, sort=True)

	tmp1 = dat.ix[:-1, :-1].values * np.tile((dat.ix[:-1, -1].values)[:, None], (1, wts.shape[1]))
	tmp2 = dat.ix[1:, :-1].values * np.tile((dat.ix[1:, -1].values)[:, None], (1, wts.shape[1]))
	delta_wts = np.abs(tmp1 * (1 + tickerRet.ix[1:, :].values) - tmp2).sum(axis=1)

	turnover = 2 * delta_wts[:, None] / (tickerNV.ix[:-1, :].values + tickerNV.ix[1:, :].values)
	turnover_df = DataFrame(turnover, index=wts.index.tolist()[1:], columns=['TurnOver'])
	return turnover_df





# # --- 获得调仓日期 ----- #
# def getChgDates(actionDates, freqofTrades='Month', dateType='head'):
# 	refPeriodDates = \
# 		pkl_read_write('D:/WorkAtAX/MultiFactors20180402/supplementaryData/refPeriodDates.pkl', 'read')
#
# 	tmp = refPeriodDates[freqofTrades][dateType]
# 	res_ind = [td in actionDates for td in tmp]
# 	res = list(Series(tmp)[res_ind])
#
# 	return res



def write_csv_bktResult(bktResult, pathBKT):

	'''

	:param bktResult: 回测结果集，字典类型，包含持仓，交割单，净值，调仓日期等等
	:param pathBKT: 存放回测结果的路径
	:return: True
	'''
	pathCombineFactors = './factorData/Combine'
	combinedFactors = pkl_read_write(pathCombineFactors + '/' + 'combinedFactors.pkl', 'read')
	selFactorRecord = pkl_read_write(pathCombineFactors + '/' + 'selFactorRecord.pkl', 'read')

	portfolioRecord = bktResult['portfolioRecord']
	tradeDetail = bktResult['tradeDetail']
	portfolioBuiltDates = bktResult['portfolioBuiltDates']

	# 把bktResult转成.csv格式
	td_all = get_CoreGenData('AShareCalendar')['Date'].tolist()
	for portfolioBuiltDate in portfolioBuiltDates:
		portfolioRecord[portfolioBuiltDate]['position'].to_csv(
			pathBKT + '/csv/portfolioRecord/' + portfolioBuiltDate + '_portfolioRecord.csv')

		# selected factors and combined factor
		calDate = td_all[td_all.index(portfolioBuiltDate) - 1]
		selFactorRecord[calDate].to_csv(pathBKT + '/csv/factorSelect/' + portfolioBuiltDate + '_factorSelect.csv')

		tmp_combinedFactors = combinedFactors.ix[combinedFactors['Date'] == calDate, :].sort_values(
			by=['combinedFactor'], ascending=False)
		tmp_combinedFactors = tmp_combinedFactors.reset_index(drop=True)
		tmp_combinedFactors.to_csv(pathBKT + '/csv/combinedFactor/' + portfolioBuiltDate + '_combinedFactor.csv')

	# 交割单
	for tradeDate in list(tradeDetail.keys()):
		tradeDetail[tradeDate].to_csv(pathBKT + '/csv/tradeDetail/' + tradeDate + '_tradeDetail.csv')

	# 交割单转成.csv
	tradeDetail_df = DataFrame()
	for tradeDate in list(tradeDetail.keys()):
		tradeDetail_df = pd.concat([tradeDetail_df, tradeDetail[tradeDate]], axis=0, sort=True)

	tradeDetail_df = tradeDetail_df.sort_values(by=['Date'])
	tradeDetail_df.to_csv('tradeDetail.csv')

	return True


# 创建文件夹路径
def checkdir(pth):
	'''
	功能：查看所传入路径是否存在，如果不存在，创建该路径
	:param pth: 路径（绝对路径和相对路径）
	:return: 返回该路径
	'''
	user_path = pth
	folder = os.path.exists(user_path)
	if not folder:
		os.makedirs(user_path)

	return user_path


# --------------------------------------------------------------------------------------
def nav2yearlyreturn(dat):    # 对应part4_Funs_HXH中函数 NV2YearlyReturn_Revised
	'''
	功能：根据净值或指数计算年度收益，年度收益为本年度最后一个交易日与上年度最后一个交易日的比值减1
	--------------------------
	:param dat: dataframe或Series, index为时间序列且格式为 %Y%m%d
	:return: index为年份的年度收益结果
	'''

	if isinstance(dat, Series):
		dat = DataFrame(dat)

	# 根据dat索引确定计算年度收益起止日
	start_year = int(dat.index[0][:4])
	end_year = int(dat.index[-1][:4])
	year_list = np.arange(start_year, end_year + 1)

	# 计算年度收益
	yearly_ret = DataFrame(0, columns=dat.columns, index=year_list)

	for iYear in year_list:
		if list(year_list).index(iYear) == 0:
			tmp = dat.ix[(repr(iYear) + '0101'): (repr(iYear) + '1231'), :].copy()
		else:
			tmp_shift1 = dat.ix[(repr(iYear-1) + '0101'): (repr(iYear-1) + '1231'), :].copy()
			tmp = dat.ix[tmp_shift1.index[-1]: (repr(iYear) + '1231'), :].copy()

		yearly_ret.loc[iYear, :] = (tmp.iloc[-1, :] / tmp.iloc[0, :] - 1)

	return yearly_ret



def nav2monthlyreturn(dat):     # 可以把index转成时间序列后，再用resample来做
	'''
	功能：根据净值或指数计算年度收益和月度收益，1) 年度收益为本年度最后一个交易日与上年度最后一个交易日的比值-1
	                                        2) 月度收益为上个月最后一个交易日与上年度最后一个交易日的比值-1
	--------------------------
	:param dat: dataframe或Series, index为时间序列且格式为 %Y%m%d
	:return: dict类型, 月度收益结果，index为年份，colums为月度
	'''

	if isinstance(dat, Series):
		dat = DataFrame(dat)
		dat.columns = ['组合']

	start_date = dat.index[0]
	start_year = start_date[:4]
	start_month = start_date[4:6]

	end_date = dat.index[-1]
	end_year = end_date[:4]
	end_month = end_date[4:6]

	year_list = np.arange(int(start_year), int(end_year) + 1)
	month_list = np.arange(1, 13)

	res = dict()

	for column in dat.columns.tolist():

		tmp_dat = dat[column].copy()

		monthly_ret = DataFrame(index=year_list, columns=month_list)
		k = -1
		for iYear in year_list:
			for iMonth in month_list:
				k = k + 1
				if ((iYear == int(start_year)) & (iMonth < int(start_month))) | (
						(iYear == int(end_year)) & (iMonth > int(end_month))):
					# 历史区间之外的月份设置为nan
					monthly_ret.loc[iYear, iMonth] = np.nan
				else:
					if k == 0:
						tmp = tmp_dat[(repr(iYear) + (('0' + repr(iMonth)) if iMonth < 10 else repr(iMonth)) + '01'):].copy()
						tmp = tmp[:(repr(iYear) + (('0' + repr(iMonth)) if iMonth < 10 else repr(iMonth)) + '31')]
					else:
						if iMonth == 1:
							tmp_shift1 = tmp_dat[(repr(iYear - 1) + '1201'): (repr(iYear - 1) + '1231')].copy()
						else:
							s_date_shift = (repr(iYear) + (('0' + repr(iMonth - 1)) if (iMonth - 1) < 10 else repr(iMonth - 1)) + '01')
							e_date_shift = (repr(iYear) + (('0' + repr(iMonth - 1)) if (iMonth - 1) < 10 else repr(iMonth - 1)) + '31')
							tmp_shift1 = tmp_dat[s_date_shift:e_date_shift]

						e_date = (repr(iYear) + (('0' + repr(iMonth)) if (iMonth) < 10 else repr(iMonth)) + '31')
						tmp = tmp_dat[tmp_shift1.index[-1]:e_date].copy()

					monthly_ret.loc[iYear, iMonth] = tmp.iloc[-1] / tmp.iloc[0] - 1

		monthly_ret.columns = list(Series(monthly_ret.columns).astype('str') + Series(['月'] * 12))

		# 计算月度胜率
		num_month = (monthly_ret.shape[0] * monthly_ret.shape[1] - (monthly_ret.isnull().sum()).sum())
		num_month_win = ((monthly_ret > 0).sum()).sum()
		monthly_win_prob = num_month_win / num_month

		# 最低/高月度收益
		mr_max = (monthly_ret.max()).max()
		mr_min = (monthly_ret.min()).min()

		# 加入年度收益
		yearly_ret = nav2yearlyreturn(tmp_dat)
		yearly_ret.columns = ['全年']

		# 合并
		monthly_ret = pd.concat([monthly_ret, yearly_ret], axis=1)

		res[column] = {'monthly_ret': monthly_ret,
					'monthly_win_prob': monthly_win_prob,
					'mr_max': mr_max,
					'mr_min': mr_min}

	return res


def nav2periodreturn(dat, period='Month'):
	'''
    通过净值曲线计算周度/月度/季度/年度收益情况
	-------------------------------
	:param dat:  Series/DataFrame, 净值
	:param period: str, 区间，周度/月度/季度/年度
	:return: dict, 返回区间收益情况
	'''

	if isinstance(dat, Series):
		dat = DataFrame(dat)

	if isinstance(dat.index[0], str):
		try:
			dat.index = dat.index.map(lambda x: datetime.strptime(x, '%Y%m%d'))
		except:
			print('index of dat is not %Y%m%d, please check!')

	period_dict = {'Week': '1W', 'Month': '1M', 'Quarter': '1Q', 'Year': '1Y'}

	sampled_dat = dat.resample(period_dict[period], closed='right', label='right').last()

	# 把第一个数添加到sampled_dat
	sampled_dat = pd.concat([dat.ix[0:1, :], sampled_dat])

	# 计算period收益
	ret = sampled_dat / sampled_dat.shift(1) - 1
	# 删除第一行
	ret = ret.ix[1:, :]
	ret.index = ret.index.map(lambda x: x.strftime('%Y%m%d'))

	# 计算区间胜率
	win_prob = (ret > 0).sum() / ret.shape[0]

	# 最低/高收益
	r_max = ret.max()
	r_min = ret.min()

	# 合并
	summary = pd.concat([DataFrame(win_prob), DataFrame(r_max), DataFrame(r_min)], axis=1)
	summary.columns = ['win_prob', 'max', 'min']

	return {'ret_table': ret,
			'summary': summary}


# def nav2periodreturn(nav, period='Month'):
# 	'''
#     通过净值曲线计算月度/季度/半年度/年度收益情况
# 	-------------------------------
# 	:param nav:  净值
# 	:param period: 区间，月度/季度/半年度
# 	:return: 返回区间收益情况
# 	'''
#
# 	# 确定起止时间
# 	start_date = nav.index[0]
# 	start_year = start_date[:4]
# 	start_month = start_date[4:6]
# 	start_season = math.ceil(int(start_month) / 3)
# 	start_halfYear = math.ceil(int(start_month) / 6)
#
# 	end_date = nav.index[-1]
# 	end_year = end_date[:4]
# 	end_month = end_date[4:6]
# 	end_season = math.ceil(int(end_month) / 3)
# 	end_halfYear = math.ceil(int(end_month) / 6)
#
# 	# 月度，季度，半年度，年度
# 	month_list = np.arange(1, 13)
# 	season_list = np.arange(1, 5)
# 	halfYear_list = np.arange(1, 3)
# 	year_list = np.arange(int(start_year), int(end_year) + 1)
#
#
# 	if period == 'Month':
# 		period_list = month_list
# 		start_period = start_month
# 		end_period = end_month
# 		period_label = '月'
# 		period_num = 12
#
# 		# 计算阶段表现
# 		period_ret = DataFrame(index=year_list, columns=period_list)
# 		for iYear in year_list:
# 			for iPeriod in period_list:
# 				if ((iYear == int(start_year)) & (iPeriod < int(start_period))) | (
# 						(iYear == int(end_year)) & (iPeriod > int(end_period))):
# 					period_ret.loc[iYear, iPeriod] = np.nan
# 				else:
# 					tmp = nav[(repr(iYear) + (('0' + repr(iPeriod)) if iPeriod < 10 else repr(iPeriod)) + '01'):]
# 					tmp = tmp[:(repr(iYear) + (('0' + repr(iPeriod)) if iPeriod < 10 else repr(iPeriod)) + '31')]
# 					tmp_period_ret = tmp.iloc[-1] / tmp.iloc[0] - 1
# 					period_ret.loc[iYear, iPeriod] = tmp_period_ret if isinstance(tmp_period_ret, float) else tmp_period_ret.values[0]
#
# 	elif period == 'Season':
# 		# 季度
# 		period_list = season_list
# 		start_period = start_season
# 		end_period = end_season
# 		period_label = '季度'
# 		period_num = 4
#
# 		period_ret = DataFrame(index=year_list, columns=period_list)
# 		for iYear in year_list:
# 			for iPeriod in period_list:
# 				if ((iYear == int(start_year)) & (iPeriod < int(start_period))) | (
# 						(iYear == int(end_year)) & (iPeriod > int(end_period))):
# 					period_ret.loc[iYear, iPeriod] = np.nan
# 				else:
# 					if iPeriod == 1:
# 						tmp = nav[(repr(iYear)+'0101'):].copy()
# 						tmp = tmp[:(repr(iYear)+'0331')].copy()
#
# 					elif iPeriod == 2:
# 						tmp = nav[(repr(iYear) + '0401'):].copy()
# 						tmp = tmp[:(repr(iYear) + '0630')].copy()
#
# 					elif iPeriod == 3:
# 						tmp = nav[(repr(iYear) + '0701'):].copy()
# 						tmp = tmp[:(repr(iYear) + '0930')].copy()
#
# 					elif iPeriod == 4:
#
# 						tmp = nav[(repr(iYear) + '1001'):].copy()
# 						tmp = tmp[:(repr(iYear) + '1231')].copy()
#
# 					tmp_period_ret = tmp.iloc[-1] / tmp.iloc[0] - 1
# 					period_ret.loc[iYear, iPeriod] = tmp_period_ret if isinstance(tmp_period_ret, float) else tmp_period_ret.values[0]
#
# 	elif period == 'halfYear':
# 		# 半年
# 		period_list = halfYear_list
# 		start_period = start_halfYear
# 		end_period = end_halfYear
# 		period_label = '半年'
# 		period_num = 2
#
# 		period_ret = DataFrame(index=year_list, columns=period_list)
# 		for iYear in year_list:
# 			for iPeriod in period_list:
# 				if ((iYear == int(start_year)) & (iPeriod < int(start_period))) | (
# 						(iYear == int(end_year)) & (iPeriod > int(end_period))):
# 					period_ret.loc[iYear, iPeriod] = np.nan
# 				else:
# 					if iPeriod == 1:
# 						tmp = nav[(repr(iYear) + '0101'):].copy()
# 						tmp = tmp[:(repr(iYear) + '0630')].copy()
#
# 					elif iPeriod == 2:
# 						tmp = nav[(repr(iYear) + '0701'):].copy()
# 						tmp = tmp[:(repr(iYear) + '1231')].copy()
#
# 					tmp_period_ret = tmp.iloc[-1] / tmp.iloc[0] - 1
# 					period_ret.loc[iYear, iPeriod] = tmp_period_ret if isinstance(tmp_period_ret, float) else tmp_period_ret.values[0]
#
# 	period_ret.columns = list(Series(period_ret.columns).astype('str') + Series([period_label] * period_num))
#
#
# 	# 计算区间胜率
# 	num_period = (period_ret.shape[0] * period_ret.shape[1] - (period_ret.isnull().sum()).sum())
# 	num_period_win = ((period_ret > 0).sum()).sum()
# 	period_win_prob = num_period_win / num_period
#
# 	# 最低/高收益
# 	period_r_max = (period_ret.max()).max()
# 	period_r_min = (period_ret.min()).min()
#
# 	# 加入年度收益
# 	yearly_ret = NV2YearlyReturnV2(nav, int(start_year), int(end_year))
# 	yearly_ret.columns = ['全年']
#
# 	# 合并
# 	period_ret = pd.concat([period_ret, yearly_ret], axis=1)
#
# 	return {'period_ret_table': period_ret,
# 			'period_win_prob': period_win_prob,
# 			'period_r_max': period_r_max,
# 			'period_r_min': period_r_min}



# ----------------------------------------------------------------------------------
# 计算相对基准回撤
def cal_relative_drawback(dat):
	'''
	功能：计算策略/基金相对与基准的最大回撤及回撤发生的时间
	计算公式：MaxLoss = min(NAV_s,t2 / NAV_s,t1 - NAV_BM,t2 / NAV_BM,t1), t1 < t2
	    其中：NAV_s 为策略净值, NAV_BM为基准净值
	--------------------------------------------------------------
	:param dat: DataFrame,nx2，第一列为策略净值，第二列为基准净值
	:return: Dataframe, 相对回撤及其起止时间
	'''

	columns = ['startDate', 'endDate', 'drawback']
	res = DataFrame(columns=columns)
	actionDates = dat.index.tolist()
	for i in range(dat.shape[0]):
		if i == 0:
			res = DataFrame([actionDates[i], actionDates[i], 0], index=columns).T
		else:
			tmp = dat.iloc[i, 0] / dat.iloc[:i, 0] - dat.iloc[i, 1] / dat.iloc[:i, 1]
			# 寻找最小值和对应日期
			res = pd.concat([res, DataFrame(Series([tmp.idxmin(), actionDates[i], tmp.min()],
												   index=['startDate', 'endDate', 'drawback'])).T])

		res = res.reset_index(drop=True)
	return res


# 计算相对基准最大回撤
def cal_relative_maxdrawback(dat):
	'''
	功能：计算策略/基金相对与基准的最大回撤及回撤发生的时间
	计算公式：MaxLoss = min(NAV_s,t2 / NAV_s,t1 - NAV_BM,t2 / NAV_BM,t1), t1 < t2
	    其中：NAV_s 为策略净值, NAV_BM为基准净值
	--------------------------------------------------------------
	:param dat: DataFrame,nx2，第一列为策略净值，第二列为基准净值
	:return: Series, 最大回撤及其起止时间
	'''

	res = cal_relative_drawback(dat)
	res = res.sort_values(by='drawback')
	maxLoss = res.iloc[0, :].copy()

	return maxLoss



def cal_yearly_relative_drawback(dat):

	'''
	功能：计算年度相对基准的最大回撤
	--------------------------------------------
	:param dat: DataFrame,nx2，第一列为策略净值，第二列为基准净值
	:return: DataFrame，每年的相对基准最大回撤，包括年度最大回撤及其起止时间
	'''

	start_year = int(dat.index[0][:4])
	end_year = int(dat.index[-1][:4])
	year_list = np.arange(start_year, end_year + 1)

	columns = ['startDate', 'endDate', 'drawback']
	yearly_maxdb = DataFrame(0, columns=columns, index=year_list)
	for iYear in year_list:
		if list(year_list).index(iYear) == 0:
			tmp = dat.ix[(repr(iYear) + '0101'): (repr(iYear) + '1231'), :].copy()
		else:
			tmp_shift1 = dat.ix[(repr(iYear-1) + '0101'): (repr(iYear-1) + '1231'), :].copy()
			tmp = dat.ix[tmp_shift1.index[-1]: (repr(iYear) + '1231'), :].copy()

		yearly_maxdb.loc[iYear, :] = cal_relative_maxdrawback(tmp)

	# yearly_maxdb = yearly_maxdb.rename(columns={'startDate': '回撤开始日期', 'endDate': '回撤结束日期', 'drawback': '年度最大回撤'})
	return yearly_maxdb


# ------------------------ 资产配置风险模型 ---------------------------- #


# --------------------- weights update ---------------------------------------#
def calSharpe(returns, wts, rf_ret=0):
    wts = np.array(wts)
    port_return = np.sum(returns.mean() * wts) * 244
    port_std = np.sqrt(np.dot(wts.T, np.dot(returns.cov() * 244, wts)))
    return (port_return - rf_ret) / port_std

def get_maxSharpeWeights(returns, rf_ret, calSharpe):

	calNegSharpe = lambda wts: - calSharpe(returns, wts, rf_ret)

	m = returns.shape[1]
	cons = ({'type': 'eq', 'fun': lambda x: np.sum(x) - 1})
	bnds = tuple((0, 1) for x in range(m))
	opts = sco.minimize(calNegSharpe, m * [1. / m, ], method='SLSQP', bounds=bnds, constraints=cons)
	wts = opts.x

	optSharpe = calSharpe(returns, opts.x)
	if optSharpe < 0:
		print('Attention: the optimized sharpe ratio is negative, return the asset with smallest fall!!!')
		# return the asset with smallest fall
		wts = np.where(np.array(returns.mean()) == returns.mean().max(), 1, 0)

	return Series(wts, index=list(returns.columns), name='weights')

def calTotSEofTRC(x, SIGMA):
	product = np.dot(SIGMA, x)
	TRC = x * product
	totSE = 0
	for i in range(0, SIGMA.shape[0]):
		for j in range(0, SIGMA.shape[0]):
			totSE = totSE + 1e12 * (TRC[i] - TRC[j]) ** 2
			# The original TRC is too small, which easily makes the optimization terminate. So multiple 1e12
	return totSE

# optimize risk parity
def optRiskPrityByTotSE(SIGMA, calTotSEofTRC):
	fun = lambda x: calTotSEofTRC(x, SIGMA)
	x0 = (1 / np.sqrt(np.diag(SIGMA))) / (1 / np.sqrt(np.diag(SIGMA))).sum()   # 1 / np.sqrt(np.diag(SIGMA))  # m * [1. / m, ]
	m = SIGMA.shape[0]
	cons = ({'type': 'eq', 'fun': lambda x: np.sum(x) - 1})
	bnds = tuple((0, 1) for x in range(m))
	opts = sco.minimize(fun, x0, method='SLSQP', bounds=bnds, constraints=cons)
	wts = opts.x
	# TRC = wts * np.dot(SIGMA, wts)
	return Series(wts, index=list(SIGMA.columns), name='weights')

def optMinVariance(SIGMA):
	minSigma = lambda x: np.sqrt(np.dot(x.T, np.dot(SIGMA, x)))
	m = SIGMA.shape[0]
	x0 = (1 / np.sqrt(np.diag(SIGMA))) / (1 / np.sqrt(np.diag(SIGMA))).sum()
	cons = ({'type': 'eq', 'fun': lambda x: np.sum(x) - 1})
	bnds = tuple((0, 1) for x in range(m))
	opts = sco.minimize(minSigma, x0, method='SLSQP', bounds=bnds, constraints=cons)
	wts = opts.x
	return Series(wts, index=list(SIGMA.columns), name='weights')


# ---------------- 最新资产配置代码 --------------- #

def downside_vol(dat):
	'''
	功能：计算区间下行波动率：1）先计算收益率均值；2）去掉大于均值的数值；3）计算小于均值的数值与均值的均方根
	-----------------------------------------
	:param ret: 日收益率序列
	:return: 区间下行波动率
	'''
	ret_mean = dat.mean()
	dat[dat > ret_mean] = np.nan
	dvol = np.sqrt(((dat - ret_mean) ** 2).mean())

	return dvol

def downside_vol_m2(ret):
	'''
	功能：计算区间下行波动率：1）去掉大于0的数值；3）计算小于0的数值的均方根
	----------------------------------------------
	:param ret: 日收益率序列
	:return: 区间下行波动率
	'''

	ret[ret > 0] = 0
	dvol = np.sqrt((ret ** 2).mean())
	return dvol


def drawback(dat):
	'''
	功能：计算截至到终止日期的回撤，drawback = abs(NaV_t / NaV_t_max - 1), NaV_t_max为t前的最大值
	----------------------------------------------------------------------
	:param dat: Series/DataFrame， 资产净值或指数行情
	:return: Series/单个数据
	'''

	if isinstance(dat, Series):
		dat = DataFrame(dat)

	# 一段时间内的回撤
	dat_max = dat.max()
	drawback = dat.ix[-1, :] / dat_max - 1

	return abs(drawback)

def self_defined_risk(ret, risk_type='normal', fct=244):

	'''
	功能：根据日收益率序列和日净值序列，计算风险值
	----------------------------------------------------------
	:param ret: Series/DataFrame，资产日收益率
	:param nav: Series/DataFrame，资产净值，计算drawback风险需要用到
	:param risk_type: 'normal', 'downside', 'drawback'
	:return: float32，risk结果
	'''

	if risk_type == 'normal':
		res = ret.std() * np.sqrt(fct)
	elif risk_type == 'downside':
		res = downside_vol(ret) * np.sqrt(fct)
	else:
		res = drawback(NV2Ret(ret, 2))

	return res

# 自定义计算资产协方差矩阵
def self_defined_cov(ret, LBW_STD=120, LBW_CORR=240):
	'''
	功能：自定义计算资产协方差矩阵，主要考虑标准差和相关系数的稳定性，设置不同的选择长度
	---------------------------------------------------------------------------
	:param ret: Series/DataFrame,日收益率序列
	:param LBW_STD: int, 计算标准差的窗口长度
	:param LBW_CORR: int, 计算相关系数的窗口长度
	:return: DataFrame, 协方差矩阵
	'''

	if isinstance(ret, Series):
		ret = DataFrame(ret)

	if ret.shape[0] < LBW_STD:
		print('length of ret < LBW_STD, SIGMA cannot be Calculated!')
		return []

	est_std = ret.ix[-LBW_STD:, :].std()
	est_corr = ret.ix[-LBW_CORR:, :].corr()

	# construct fct matrix
	fct_matrix = np.diag(est_std ** 2)
	for i in range(fct_matrix.shape[0]):
		for j in range(fct_matrix.shape[1]):
			fct_matrix[i, j] = est_std.values[i] * est_std.values[j]

	est_cov = DataFrame(fct_matrix * est_corr.values, index=ret.columns, columns=ret.columns)

	return est_cov   # 注意ret要用调仓日之前的数据   # Ben @20190709



def risk_target_func(x, risk, risk_target):
	'''
	功能：根据组合成分的波动率，计算某个权重x下，组合的波动率数值
	----------------------------------------------------------
	:param x: list/numpy.array，资产权重
	:param risk: list/numpy.array，资产风险
	:param risk_target: 目标风险预期值
	:return: 默认资产相关性为0的情形下，组合波动率与目标风险预期值的误差平方
	'''

	num = len(risk)
	tot_risk = 0
	for i in range(0, num):
		tot_risk = tot_risk + ((x[i]) ** 2) * (risk[i] ** 2)   # 默认资产相关性为0
	err = ((np.sqrt(tot_risk) - risk_target)) ** 2 * 1e6   # multiple 1e6，提高精度，也可在迭代中设置精度
	return err


# 目标风险优化资产权重值
def opt_risk_target(risk, risk_target, constraints=[]):
	'''
	功能：根据目标波动率优化组合资产权重，注意资产个数超过2以后可能需要更多约束才能优化出结果
	----------------------------------------------
	:param risk: numpy.array, 资产波动率
	:param risk_target: float32, 组合目标风险
	:param constraints: dict / list of dict, 优化约束条件
	:return: dict, 优化结果包括权重value，是否优化达到误差小的程度flag，提示信息message
	'''

	res = {
		'value': None,
		'flag': None,
		'message': ''
	}

	if (len(risk[risk > risk_target]) == len(risk)) | (len(risk[risk > risk_target]) == 0):
		# 所有资产波动率均大于或小于risk_target,无法优化
		res['message'] = 'Asset volatilities are all larger than target or smaller than target, cannot optimize!'
	else:

		fun = lambda x: risk_target_func(x, risk, risk_target)

		m = len(risk)
		x0 = np.array([0.1] * m)  # m * [1. / m]   # 初值很重要[0.5, 0.5]不好，加入hess是否会更好？

		constraints = [constraints] if isinstance(constraints, dict) else constraints
		cons = [{'type': 'eq', 'fun': lambda x: np.sum(x) - 1}] + constraints   # 如果没有更多限制条件，多个资产权重无穷个解
		bnds = tuple((0, 1) for j in range(m))
		opts = sco.minimize(fun, x0, method='SLSQP', bounds=bnds, constraints=cons)
		wts = opts.x

		res['value'] = wts
		res['flag'] = 1 if opts.fun < 1e-1 else 0
		res['message'] = 'Optimization Completed!'

	return res


def risk_budget_func(x, SIGMA, risk_budget):
	'''
	功能：计算资产对组合的风险贡献度与风险预算比例的误差平方和，用于资产权重优化
	--------------------------------------------------------------------
	:param x: numpy.array, 资产权重
	:param SIGMA: DataFrame, 资产协方差
	:param risk_budget: list/numpy.array, 资产风险预算值
	:return: float32, 误差项
	'''

	product = np.dot(SIGMA, x)
	TRC = x * product
	totSE = 0

	for i in range(0, SIGMA.shape[0]):
		for j in range(0, SIGMA.shape[0]):
			totSE = totSE + 1e12 * (TRC[i] / risk_budget[i] - TRC[j] / risk_budget[j]) ** 2

	return totSE


def opt_risk_budget(SIGMA, risk_budget, constraints=[]):
	'''
	功能：根据风险预算优化资产权重
	-------------------------------------------
	:param SIGMA: DataFrame, 资产协方差
	:param risk_budget: list/numpy.array, 资产风险预算值
	:param constraints: dict / list of dict, 优化约束条件
	:return: 优化结果包括权重value，是否优化达到误差小的程度flag，提示信息message
	'''

	res = {
		'value': None,
		'flag': None,
		'message': ''
	}

	fun = lambda x: risk_budget_func(x, SIGMA, risk_budget)
	x0 = (1 / np.sqrt(np.diag(SIGMA))) / (1 / np.sqrt(np.diag(SIGMA))).sum()
	m = SIGMA.shape[0]

	constraints = [constraints] if isinstance(constraints, dict) else constraints
	cons = [{'type': 'eq', 'fun': lambda x: np.sum(x) - 1}] + constraints
	bnds = tuple((0, 1) for j in range(m))
	opts = sco.minimize(fun, x0, method='SLSQP', bounds=bnds, constraints=cons)
	wts = opts.x

	res['value'] = wts
	res['flag'] = 1 if opts.fun < 1e-1 else 0
	res['message'] = 'Optimization Completed!'

	return res


def cal_fix_weights(ret, chgDates, fixed_ratio=[0.2, 0.8]):
	'''
	功能：固定股债配比
	----------------------------------------
	:param ret: DataFrame, 股债日收益率序列
	:param chgDates: list, 调仓日列表
	:param fixed_ratio: list/numpy.array, 股债固定比例
	:return: DataFrame, 权重序列
	'''

	wts = DataFrame(0, index=chgDates, columns=ret.columns)
	wts.ix[:, ret.columns[0]] = fixed_ratio[0]
	wts.ix[:, ret.columns[1]] = fixed_ratio[1]

	return wts


# 这种函数放在Project中，不要放在公共函数中
# 把这个函数拆开s

def cal_asset_risk(ret, chgDates, risk_type='normal'):

	risk = DataFrame(index=chgDates, columns=ret.columns)
	actionDates = ret.index.tolist()

	for chgDate in chgDates:

		# 定义风险--> 1) 波动率; 2) 下行波动率; 3) 回撤
		ind = actionDates.index(chgDate)
		LBW = 240 if risk_type == 'drawback' else 120
		if ret.shape[0] > LBW:
			used_ret = ret.ix[(ind - LBW):ind, :].copy()
		else:
			used_ret = ret.ix[:ind, :].copy()

		risk.ix[chgDate, :] = self_defined_risk(used_ret, risk_type=risk_type)

	return risk



def adjust_risk(asset_risk, adjust_fct=None):

	if asset_risk.shape[0] == adjust_fct.shape[0]:
		risk_adjusted = asset_risk.copy()
		if adjust_fct is not None:
			adj_columns = list(set(asset_risk.columns).intersection(set(adjust_fct.columns)))
			risk_adjusted.ix[:, adj_columns] = adjust_fct.ix[:, adj_columns] * risk_adjusted.ix[:, adj_columns]
	else:
		print('The length of asset_risk and adjust_fct do not match, Please Check!')
		risk_adjusted = []

	return risk_adjusted




def get_risk_adjust_fct(w, varCodes, varNames, indicator, start_dt, end_dt):

	if isinstance(indicator, str):
		indicator = [indicator]

	dat = get_wsd(w, varCodes, indicator,  varNames, start_dt, end_dt)['Data']

	# 取中位数/均值
	s = Series(range(dat.shape[0]))
	dat_median = s.apply(lambda x: dat.iloc[:(x+1), :].median())

	# 根据估值指标生成调整因子（看这个循环怎么修改下，节约运行时间）
	adjust_fct = DataFrame(dat.values / dat_median.values, index=dat.index, columns=dat.columns)

	# 前面没有估值的数据调整因子设为1
	adjust_fct = adjust_fct.fillna(value=1)

	return adjust_fct




def cal_risk_target_wts(asset_risk, risk_target=0.04, stock_limit=0.2):

	actionDates = asset_risk.index.tolist()
	wts = DataFrame(0, index=actionDates, columns=asset_risk.columns)
	constraints = {'type': 'ineq', 'fun': lambda x: stock_limit - x[0]}   # 股票排在前面，股票权重<=stock_limit

	for actionDate in actionDates:
		tmp = opt_risk_target(asset_risk.ix[actionDate, :].values, risk_target, constraints=constraints)  # 可以自定义加约束
		wts.ix[actionDate, :] = tmp['value']
		if tmp['flag'] == 0:
			print(actionDate + ' optimization may not be reached, Be Careful！')

	return wts


# 包含货币基金
def cal_risk_target_wts_wmf(asset_risk, risk_target=0.04, stock_limit=0.2):
	# 不用传入货币基金信息？

	actionDates = asset_risk.index.tolist()
	wts = DataFrame(0, index=actionDates, columns=asset_risk.columns.tolist() + ['货币基金'])
	constraints = {'type': 'ineq', 'fun': lambda x: stock_limit - x[0]}   # 股票权重<=stock_limit
	MF_vol = 0.02

	for actionDate in actionDates:
		tmp = opt_risk_target(asset_risk.ix[actionDate, :].values, risk_target, constraints=constraints)  # 可以自定义加约束
		tmp_wts = tmp['value']
		if tmp['flag'] == 0:
			print(actionDate + ' optimization may not be reached, Be Careful！')

		Bond_vol = asset_risk.ix[actionDate, 1]   # 计算资产风险,提取债券波动率
		tmp_wts2 = tmp_wts[1] * np.array([MF_vol / (Bond_vol + MF_vol), Bond_vol / (Bond_vol + MF_vol)])

		wts.ix[actionDate, :] = [tmp_wts[0], tmp_wts2[0], tmp_wts2[1]]

	return wts



#
# def getRiskTargetWeights_V2(ret, nav, chgDates, risk_target=0.04, risk_type='normal', stock_limit=0.2, adjust_fct=None):
#
# 	wts = DataFrame(0, index=chgDates, columns=ret.columns)
# 	actionDates = ret.index.tolist()
#
# 	risk = DataFrame(index=chgDates, columns=ret.columns)
#
# 	for chgDate in chgDates:
#
# 		# 1. 定义风险--> 1) 波动率; 2) 下行波动率; 3) 回撤。  加权或不加权，调整或不调整
# 		ind = actionDates.index(chgDate)
#
# 		used_ret = ret.ix[(ind - 120):ind, :].copy()   # 注意不要引入当日收盘价    Ben @20190710
#
# 		if risk_type == 'normal':
# 			# 1) 波动率
# 			risk.ix[chgDate, :] = self_defined_risk(used_ret)
# 		elif risk_type == 'downside':
# 			# 2）下行波动率
# 			risk.ix[chgDate, :] = self_defined_risk(used_ret, risk_type='downside')
# 		elif risk_type == 'drawback':
# 			# 3）回撤
# 			if nav.shape[0] > 240:
# 				used_nav = nav.ix[(ind - 240):ind, :].copy()   # 回撤是2年历史数据
# 			else:
# 				used_nav = nav.ix[:ind, :].copy()  # 回撤是2年历史数据
# 			risk.ix[chgDate, :] = self_defined_risk(used_ret, nav=used_nav, risk_type='drawback')
#
# 		# 如果债券指数或货币基金回撤为0，设置为1%
#
# 		# 股票估值调整，债券不用
# 		if adjust_fct is not None:
# 			adj_columns = list(set(risk.columns).intersection(set(adjust_fct.columns)))
# 			risk.ix[chgDate, adj_columns] = adjust_fct.ix[adjust_fct.index.tolist().index(chgDate)-1, :] * risk.ix[chgDate, adj_columns]
#
# 		# if time_weighted:
#
#
# 		# 目标波动率
# 		# 1）不考虑货币基金
# 		# cons = []    # {'type': 'ineq', 'fun': lambda x: stock_limit - x[0]}
#
# 		if risk_type == 'drawback':
#
# 			# 如果股债回撤都小于1，股债各50%, “回撤时一阶的，可以直接比例求”
# 			if (risk.ix[chgDate, :] > risk_target).sum() == 1:
# 				wts.ix[chgDate, 0] = (risk_target - risk.ix[chgDate, 1]) / (risk.ix[chgDate, 0] - risk.ix[chgDate, 1])
# 				wts.ix[chgDate, 1] = 1 - wts.ix[chgDate, 0]
# 			elif (risk.ix[chgDate, :] > risk_target).sum() == 2:
# 				wts.ix[chgDate, :] = [0, 1]   # 都大于risk_target全部配债
# 			else:
# 				wts.ix[chgDate, :] = [0.5, 0.5]
#
# 		else:
# 			wts.ix[chgDate, :], flag = optRiskTarget(risk.ix[chgDate, :].values, risk_target)   # 可以自定义加约束
# 			if flag == 0:
# 				print(chgDate + ' optimization may not be reached, Be Careful！')
#
#
# 	return risk, wts
#




def cal_risk_parity_wts(ret, chgDates, risk_budget=None, stock_limit=0.2):   # old fun name: getRiskParityWeights_V1

	wts = DataFrame(0, index=chgDates, columns=ret.columns)
	actionDates = ret.index.tolist()

	constraints = {'type': 'ineq', 'fun': lambda x: stock_limit - x[0]}  # 股票权重<=stock_limit

	# 如果不提供risk_budget，等分风险预算
	if risk_budget is None:
		risk_budget = np.array([1 / ret.shape[1]] * ret.shape[1])

	for chgDate in chgDates:

		# 计算协方差矩阵
		ind = actionDates.index(chgDate)
		if ind >=240:
			used_ret = ret.ix[:ind, :].copy()
			est_cov = self_defined_cov(used_ret, LBW_STD=120, LBW_CORR=240)

			wts.ix[chgDate, :] = opt_risk_budget(est_cov, risk_budget, constraints=constraints)['value']
		else:
			print('Check length of trading dates before 1st chgDate!')
			pass

	return wts




def cal_risk_parity_wts_wmf(ret, chgDates, risk_budget=None, stock_limit=0.2):    # old fun name: getRiskParityWMFWeights_V1

	# --- 赋权逻辑
	# 1. 先把股票收益率合起来，与债券用风险平价/预算，算出大类资产比例
	# 2. 股票内部用风险平价（波动率分之一加权）
	# 3. 加入货币基金，或者把货币基金当作第三种配置资产，设置波动率为4% (可以调整)。

	wts = DataFrame(0, index=chgDates, columns=ret.columns.tolist() + ['货币基金'])
	actionDates = ret.index.tolist()
	MF_vol = 0.04

	# 如果不提供risk_budget，等分风险预算
	if risk_budget is None:
		risk_budget = np.array([1 / ret.shape[1]] * ret.shape[1])

	constraints = {'type': 'ineq', 'fun': lambda x: stock_limit - x[0]}  # 股票权重<=stock_limit

	for chgDate in chgDates:

		# 计算协方差矩阵
		ind = actionDates.index(chgDate)
		if ind >= 240:
			used_ret = ret.ix[:ind, :].copy()
			est_cov = self_defined_cov(used_ret, LBW_STD=120, LBW_CORR=240)

			tmp = opt_risk_budget(est_cov, risk_budget, constraints=constraints)
			tmp_wts = tmp['value']
			if tmp['flag'] == 0:
				print(chgDate + ' optimization may not be reached, Be Careful！')

			Bond_vol = np.sqrt(est_cov.ix[1, 1]) * np.sqrt(244)  # 计算资产风险,提取债券波动率
			tmp_wts2 = tmp_wts[1] * np.array([MF_vol / (Bond_vol + MF_vol), Bond_vol / (Bond_vol + MF_vol)])

			wts.ix[chgDate, :] = [tmp_wts[0], tmp_wts2[0], tmp_wts2[1]]

		else:
			print('Check length of trading dates before 1st chgDate!')
			pass

	return wts


















# ------------------- 研究所资产配置模型 -----------------------------  #

def calRVCWeights(ret, RVCwts):

	wts = Series(np.zeros(ret.shape[1]), index=ret.columns.tolist())

	# 计算均值排序
	rank_mean = ret.mean().rank(ascending=False)
	# 计算波动率排序
	rank_std = ret.std().rank()
	# 计算相关性排序
	rank_corr = (ret.corr().sum() - 1).rank()

	rank = pd.concat([rank_mean, rank_std, rank_corr], axis=1, sort=True)
	rank.columns = ['R', 'V', 'C']

	score = (rank * RVCwts).sum(axis=1)
	# 去掉平均收益率小于0的资产
	score = score[ret.mean() > 0]
	scoreRank = score.rank()

	tmp = (len(scoreRank) - scoreRank + 1) / scoreRank.sum()
	wts[tmp.index.tolist()] = tmp

	if abs(wts.sum() - 1.0) > 1e-3:
		logging.error('Sum of RVC weights is not 1.')

	return wts


def calESWeights(ret, ES_T, wts_init, cash_flag):

	# 计算VaR和ES
	ret_mean = ret.mean()
	ret_std = ret.std()
	# ret_VaR = 20 * ret_mean - (20 ** (1 / 2)) * ret_std * norm.ppf(0.95)  # 20 为月化参数
	ret_ES = - 20 * ret_mean + (20 ** (1 / 2)) * ret_std * norm.pdf(norm.ppf(0.95)) / 0.05
	# ret_ES 在行情好的时候有可能出现负数（特别是Bond, 波动小)

	ES_weights = ES_T / ret_ES
	ES_weights[(ES_weights > 1) | (ret_ES < 0)] = 1

	wts = wts_init * ES_weights
	wts_cash = (wts_init * (1 - ES_weights)).sum()

	ind = (ES_weights >= 1) & (wts_init > 0)
	if cash_flag == 0 and len(ES_weights[ind]) > 0:
		wts[ind] = wts[ind] + wts_cash / len(ES_weights[ind])
		wts['Currency'] = 0
	else:
		wts['Currency'] = wts_cash

	if abs(wts.sum() - 1.0) > 1e-3:
		logging.error('Sum of weights is not 1.')

	return wts





# ------------------------ Wind api 使用函数 -------------------------- #

def get_tradingdate_api(w, curDate):
	'''
	功能：通过Wind API获取交易日数据，包括是否一周/月/季度/年度末尾
	-----------------------------------------------------------------
	:param w: wind api
	:param curDate: str, 当前日期，可以%Y%m%d或%Y-%m-%d 格式
	:return: dict, Data为DataFrame，交易日期信息
	'''

	res = {
		'Data': None,
		'ErrorCode': 0,
		'message': ''
	}

	if len(curDate) == 8:
		dt = datetime.strptime(curDate, '%Y%m%d')
	elif len(curDate) == 10:
		dt = datetime.strptime(curDate, '%Y-%m-%d')
	else:
		logging.error('Input Date Error!')
		return res

	next_year = str(dt.year + 1)

	if w.isconnected():
		rawData_D = w.tdays('1991-01-01', next_year + '-12-31', 'Period=D')
		rawData_W = w.tdays('1991-01-01', next_year + '-12-31', 'Period=W')
		rawData_M = w.tdays('1991-01-01', next_year + '-12-31', 'Period=M')
		rawData_Q = w.tdays('1991-01-01', next_year + '-12-31', 'Period=Q')

		res['ErrorCode'] = rawData_D.ErrorCode | rawData_W.ErrorCode | rawData_M.ErrorCode | rawData_Q.ErrorCode

		if ~res['ErrorCode']:
			actionDates = list(Series(rawData_D.Times).apply(lambda c: c.strftime('%Y%m%d')))
			dat = DataFrame(rawData_D.Data, columns=actionDates, index=['Date']).T
			dat['Date'] = dat['Date'].apply(lambda x: x.strftime("%Y%m%d"))

			Weeks = list((DataFrame(rawData_W.Data, index=['Date']).T)['Date'].apply(lambda x: x.strftime("%Y%m%d")))
			Months = list((DataFrame(rawData_M.Data, index=['Date']).T)['Date'].apply(lambda x: x.strftime("%Y%m%d")))
			Quarters = list((DataFrame(rawData_Q.Data, index=['Date']).T)['Date'].apply(lambda x: x.strftime("%Y%m%d")))

			# 判断是否是月末   # 最后把整个系统这里，应该这个数据是已经存在的     ！！！

			dat['isWeek'] = 0
			dat['isMonth'] = 0
			dat['isQuarter'] = 0

			dat.ix[dat['Date'].isin(Weeks), 'isWeek'] = 1
			dat.ix[dat['Date'].isin(Months), 'isMonth'] = 1
			dat.ix[dat['Date'].isin(Quarters), 'isQuarter'] = 1

		else:
			res['message'] = 'Fetching Data Error '
			print(res['message'])
			dat = DataFrame()
	else:
		res['message'] = 'Wind is not connected!'
		print(res['message'])
		dat = DataFrame()

	res['Data'] = dat

	return res


def getPeriodFisrtLastDates(trading_date, period='Month', flag='first'):
	'''
	功能：根据trading_date，计算月度/季度/年度等的起止日期
	:param trading_date: DataFrame, 包含季节信息, isWeek, isMonth, isQuarter
	:param period: str, 月度或季度
	:param flag: str, first/last
	:return: list, 返回period的起始或结尾日期
	'''

	column = 'is' + period
	trading_date = trading_date.ix[:, ['Date', column]].copy()

	# 避免后面出现out of index range
	if trading_date.iloc[-1, -1] == 1:
		trading_date = trading_date.iloc[:-1, :]

	actionDates = trading_date['Date'].tolist()
	period_end = trading_date.ix[trading_date[column] == 1, :].index.tolist()
	period_start = [actionDates[actionDates.index(td)+1] for td in period_end]
	period_start = [actionDates[0]] + period_start

	if flag == 'first':
		res = period_start
	elif flag == 'last':
		res = period_end

	return res






# ----------------- 获取数据函数

# 提取数据
# -------------------------建议在使用前，用CG确认参数信息
# 1. get_wsd用于提取二维数据，varCodes, indicators不能同时多维
# 2. get_wss用于提取属性数据或三维数据
# 3. get_edb用于提取宏观数据
# 4. get_sector 用于提取板块成分信息

# 将输入参数统一改成others  Ben @20191031
def get_wsd(w, varCodes, indicators, varNames, start_dt, end_dt, others=''):
	'''
	功能：通过提取一段时间内每日行情数据（行情数据等）
	----------------------------------------
	:param w: wind api
	:param varCodes: list, wind代码
	:param indicators: list, 指标数据
	:param varNames: list, wind代码对应的名称
	:param start_dt: 起始日期，%Y%m%d 或 %Y-%m-%d
	:param end_dt: 到期日期，%Y%m%d 或 %Y-%m-%d
	:param others: 其他数据信息，如复权等
	:return: DataFrame, 返回提取的数据信息
	'''

	res = {
		'Data': None,
		'ErrorCode': 0,
		'Message': ''
	}

	if isinstance(indicators, str):        # Ben @20190510  把字符型Indicators转成list
		indicators = [indicators]
	if isinstance(varNames, str):
		varNames = [varNames]

	if (len(varCodes) > 1) & (len(indicators) > 1):
		res['Message'] = 'varCodes and indicators are both Mulit-dimensional, Please use Loop!'
		print(res['Message'])
		dat = DataFrame()

	else:
		if w.isconnected():
			rawData = w.wsd(varCodes, indicators, start_dt, end_dt, others)
			res['ErrorCode'] = rawData.ErrorCode

			if not rawData.ErrorCode:   # 将~改为not  Ben @20191031
				actionDates = list(Series(rawData.Times).apply(lambda c: c.strftime('%Y%m%d')))
				dat = DataFrame(rawData.Data, columns=actionDates, index=varNames).T
			else:
				res['Message'] = 'Wind API Fetching Data Error '
				print(res['Message'])
				dat = DataFrame()
		else:
			res['Message'] = 'Wind is not Connected!'
			print(res['Message'])
			dat = DataFrame()

	res['Data'] = dat

	return res


def get_edb(w, varCodes, varNames, start_dt, end_dt, others=''):
	'''
	功能：提取宏观变量的历史数据
	-----------------------------------------
	:param w: wind api
	:param varCodes: list, wind代码
	:param varNames: list, wind代码对应的名称
	:param start_dt: 起始日期，%Y%m%d 或 %Y-%m-%d
	:param end_dt: 到期日期，%Y%m%d 或 %Y-%m-%d
	:param others: 其他数据信息，如复权等
	:return: DataFrame, 返回提取的数据信息
	'''

	res = {
		'Data': None,
		'ErrorCode': 0,
		'Message': ''
	}

	if w.isconnected():
		rawData = w.edb(varCodes, start_dt, end_dt, others)    # 去掉用前值填充 "Fill=Previous"   Ben @20190510
		res['ErrorCode'] = rawData.ErrorCode

		if ~rawData.ErrorCode:
			actionDates = list(Series(rawData.Times).apply(lambda c: c.strftime('%Y%m%d')))
			if (len(actionDates) > 1):
				dat = DataFrame(rawData.Data, columns=actionDates, index=varNames).T
			else:
				dat = DataFrame(rawData.Data, index=actionDates, columns=varNames)
				# 注意只有一个数据的时候，数据结构，可能不能转置，还需要check   Ben @20190123
		else:
			res['Message'] = 'Wind API Fetching Data Error '
			print(res['Message'])
			dat = DataFrame()

	else:
		res['Message'] = 'Wind is not Connected!'
		print(res['Message'])
		dat = DataFrame()

	res['Data'] = dat

	return res


# 注意收益率others内容 'tradeDate=20191030;credibility=1',比较丰富
def get_wss(w, varCodes, indicators, others=''):
	'''
	功能：通过Wind api提取每日行情数据,按每一天取,注意不同指标可能others是不同的
	:param w: wind api
	:param varCodes: list, wind代码
	:param indicators: list, 指标数据
	:param others: 其他数据信息，如复权等
	:return: 返回提取的数据信息
	'''

	res = {
		'Data': None,
		'ErrorCode': 0,
		'message': ''
	}

	if isinstance(varCodes, str):
		varCodes = [varCodes]
	if isinstance(indicators, str):
		indicators = [indicators]

	if w.isconnected():

		N0 = len(varCodes)
		n = 1000            # 1000个以上用循环
		m = int(N0 / n) + 1
		df_list = list()

		for i in range(m):

			if i < m - 1:
				tmp_codes = varCodes[(i * n): ((i + 1) * n)]
			else:
				tmp_codes = varCodes[(i * n):]

			if len(tmp_codes) > 0:

				rawData = w.wss(tmp_codes, indicators, others)
				tmp = rawData.Data
				ErrorCode = rawData.ErrorCode

				if not ErrorCode:
					tmp = DataFrame(tmp, index=indicators, columns=rawData.Codes).T
					df_list.append(tmp)
				else:
					break

		# 汇总
		res['ErrorCode'] = ErrorCode
		if not res['ErrorCode']:
			dat = DataFrame().append(df_list)
		else:
			res['message'] = 'Wind API Fetching Data Error '
			print(res['message'])
			dat = DataFrame()
	else:
		dat = DataFrame()
		print('Wind is not connected!')

	res['Data'] = dat

	return res


def get_sector(w, sectorids, dt):
	'''
	功能：提取板块在某一天的成分信息
	:param w: wind api
	:param sectorids: 板块id，注意债券板块有些是到期的，也要在回溯中考虑
	:param dt: 当期日期，%Y%m%d 或 %Y-%m-%d
	:return: 返回提取的数据信息
	'''

	res = {
		'Data': None,
		'ErrorCode': 0,
		'message': ''
	}

	if isinstance(sectorids, str):
		sectorids = [sectorids]


	if w.isconnected():

		df_list = list()

		for iid in sectorids:

			rawData = w.wset("sectorconstituent", "date=" + dt + ";sectorid=" + iid)
			ErrorCode = rawData.ErrorCode

			if not ErrorCode:
				tmp = DataFrame(rawData.Data).T
				tmp.columns = ['Date', 'WindCode', 'Name']
				tmp['Date'] = tmp['Date'].apply(lambda x: x.strftime('%Y%m%d'))
				tmp['Sector'] = iid
				df_list.append(tmp)
			else:
				break

		# 汇总
		res['ErrorCode'] = ErrorCode
		if not res['ErrorCode']:
			dat = DataFrame().append(df_list)
		else:
			res['message'] = 'Wind API Fetching Data Error '
			print(res['message'])
			dat = DataFrame()
	else:
		dat = DataFrame()
		print('Wind is not connected!')

	res['Data'] = dat

	return res




# ----------------- 写入word的函数 ------------------------- #
import pickle
import pandas as pd
import numpy as np
from pandas import DataFrame, Series
import logging
import matplotlib.pyplot as plt
import os
from datetime import datetime, timedelta

from docx import Document
from docx.oxml.ns import qn  # 字体
from docx.shared import Pt  # 字号
from docx.shared import RGBColor  # 颜色

from docx.enum.text import WD_ALIGN_PARAGRAPH  # 段落居中
from docx.enum.table import WD_TABLE_ALIGNMENT  # 表格水平居中
from docx.enum.table import WD_ALIGN_VERTICAL  # 表格垂直居中
from docx.enum.text import WD_COLOR  # 高亮 ex. run.font.highlight_color = WD_COLOR.YELLOW

from docx.enum.text import *  # 段落居中
from docx.enum.table import *  # 表格水平居中

from docx.shared import Inches  # 图片尺寸
from docx.oxml.ns import nsdecls  # 填充颜色
from docx.oxml import parse_xml


def any2str(value, decimals=4):
	"""
	任意类型的格式 --> str
	---------------------------------
		decimals: 小数点后的位数
	"""
	if type(value) == str:
		return value
	# 浮点数保留一定位数
	elif type(value) in [float, np.float64, np.float32, np.float16, np.float]:
		return str(np.round(value, decimals=decimals))
	else:
		return str(value)


class user_docx(object):
	"""
	基于python_docx包自定义类
	"""

	# 初始化: 新建Document
	def __init__(self, EnglishFont=u'Times New Roman', ChineseFont=u'微软雅黑', Size=12):
		self.document = Document()
		self.document.styles['Normal'].font.name = EnglishFont  # 英文字体
		self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), ChineseFont)  # 中文字体
		self.document.styles['Normal'].font.size = Pt(Size)  # 字体大小
		# 其他参数 mengzl@20190507
		self.EnglishFont = EnglishFont
		self.ChineseFont = ChineseFont
		self.Size = Size

	# 新建标题
	def user_add_heading(self, text, level,
						 EnglishFont=u"", ChineseFont=u"", Size=0,
						 R=0, G=0, B=0):
		#		"""
		#		设置标题及相应格式
		#		---------------------------------
		#			text: 标题内容
		#			level: 标题等级, 一般不用0
		#			EnglishFont / ChineseFont: 英文/中文字体
		#			Size: 字体大小
		#			R/G/B: rgb三原色
		#		"""
		run = self.document.add_heading('', level=level).add_run(text)
		run.font.color.rgb = RGBColor(R, G, B)

		# 英文字体
		if len(EnglishFont) == 0:
			run.font.name = self.EnglishFont  # 默认填充 mengzl@20190507
		else:
			run.font.name = EnglishFont
		# 中文字体
		if len(ChineseFont) == 0:
			run._element.rPr.rFonts.set(qn('w:eastAsia'), self.ChineseFont)
		else:
			run._element.rPr.rFonts.set(qn('w:eastAsia'), ChineseFont)
		# 字体大小
		if Size == 0:
			run.font.size = Pt(self.Size)
		else:
			run.font.size = Pt(Size)

	def user_add_paragraph(self, text,
						   EnglishFont=u"", ChineseFont=u"",
						   Size=0, bold=False, center=False,
						   R=0, G=0, B=0):
		#	"""
		#	设置段落及相应格式
		#	---------------------------------
		#		text: 标题内容
		#		EnglishFont / ChineseFont: 英文/中文字体, 默认Document本身取值
		#		size: 字体大小
		#		bold: 加粗
		#		center: 居中
		#		R/G/B: rgb三原色
		#	"""

		run = self.document.add_paragraph()
		if center:
			run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		run = run.add_run(text)
		# 加粗与颜色
		run.bold = bold
		run.font.color.rgb = RGBColor(R, G, B)
		# 是否自定义字体/字号
		if len(EnglishFont) != 0:
			run.font.name = EnglishFont  # 英文字体
		if len(ChineseFont) != 0:
			run._element.rPr.rFonts.set(qn('w:eastAsia'), ChineseFont)  # 中文字体
		if Size != 0:
			run.font.size = Pt(Size)  # 字体大小

	@staticmethod
	def __RGB2HEX(value):
		#	"""
		#	RGB --> HEX, 用于填充底色
		#	调整为私有方法, 外部不调用 mengzl@20190416
		#	"""
		digit = list(map(str, range(10))) + list("ABCDEF")
		string = str()
		for i in value:
			a1 = i // 16
			a2 = i % 16
			string += digit[a1] + digit[a2]
		return string

	def __fill_cell(self, cell, text,
					center=True, size=12, bold=True,
					fill=False, R=0, G=0, B=0):
		#	"""
		#	统一处理单元格内 内容/字号/居中/加粗/背景颜色 mengzl@20190416
		#	设置为私有方法, 外部不调用
		#	"""
		if cell.text != "":
			return
		run = cell.paragraphs[0].add_run(any2str(text))
		run.font.size = Pt(size)
		if center == True:
			cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		if bold:  # 加粗
			run.font.bold = True
		if fill:  # 填充底色, 必须每个格子定义一个颜色!
			shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), self.__RGB2HEX((R, G, B))))
			cell._tc.get_or_add_tcPr().append(shading_elm)

	@staticmethod
	def __zscore2RGB(zscore):
		#	"""
		#	zscore --> RGB, 用于填充底色, zscore参考范围: [-2,2] mengzl@20190416
		#	设置为私有方法, 外部不调用
		#	"""
		R = int(min(max(0, np.round(255 * (2 - zscore) / 4)), 255.0))  # 修正了大于255的情况 mengzl@20190507
		G = int(min(max(0, np.round(255 * (2 + zscore) / 4)), 255.0))
		B = int(max(0, np.round(1 - np.abs(1 - np.abs(zscore)))))
		return (R, G, B)

	def user_add_table(self, df,
					   column_insert=True, index_insert=True,
					   center=True,
					   column_size=12, size=12,
					   column_bold=False, index_bold=False,
					   text_fill_column=False, fill_column_R=218, fill_column_G=165, fill_column_B=32,
					   text_fill_row=[], fill_row_R=135, fill_row_G=206, fill_row_B=250,
					   row_merge_list=[], col_merge_list=[], merge_R=0, merge_G=191, merge_B=255,
					   zscore_column=[],
					   table_style='Table Grid'):
		#	"""
		#	将df(dataframe)作为表格, 插入document
		#	---------------------------------
		#		column_insert / index_insert: 是否插入df.columns / df.index
		#		center: 表格内元素是否居中
		#		column_size / size: 列名/内容字号
		#		culumn_bold / index_bold: 列名/行名加粗
		#		text_fill_column: 列名填充底色
		#		text_fill_row: 按行序号填充底色
		#		table_style: 表格样式
		#
		#	mengzl@20190416 修改
		#	根据取值自动染色: 红绿渐变, [-3,3] --> [(255,0,0),(0,255,0)] (R,G,B)
		#	根据row_merge和col_merge合并单元格 (注意不能先输入“”再合并, 会变成换行)
		#	---------------------------------
		#		ex. row_merge_list = [1,1,3], col_merge_list = [[1,3],[4,5],[3,8]]
		#	"""

		# column/index 名称
		column_name = df.columns.tolist()
		index_name = df.index.tolist()
		# 新增表格
		table = self.document.add_table(rows=len(df) + int(column_insert),
										cols=len(column_name) + int(index_insert))
		# 表格样式
		table.style = table_style

		# 插入数据
		for i_row, row in enumerate(table.rows):
			# 需要合并单元格的列
			if i_row in row_merge_list:
				row_merge = True
				col_merge_ind = [ii for ii, x in enumerate(row_merge_list) if i_row == x]
				col_merge_row = [col_merge_list[ind] for ind in col_merge_ind]
			else:
				row_merge = False
			for i_cell, cell in enumerate(row.cells):
				# 1. 插入合并单元格内容
				if row_merge:
					for col_merge in col_merge_row:
						if i_cell in col_merge:
							# 一次插入
							if i_cell == min(col_merge):
								self.__fill_cell(cell,
												 text=df.iloc[i_row - int(column_insert), i_cell - int(index_insert)],
												 size=size, bold=False, center=center,
												 fill=True, R=merge_R, G=merge_G, B=merge_B)
								cell_max = row.cells[max(col_merge)]
								cell.merge(cell_max)
							# 其他跳过
							else:
								continue
				# 2. 插入column_name(第1行)
				if i_row == 0 and column_insert == True:
					if i_cell == 0 and index_insert == True:
						self.__fill_cell(cell, text="",
										 size=column_size, bold=False, center=center,
										 fill=text_fill_column, R=fill_column_R, G=fill_column_G, B=fill_column_B)
					else:
						self.__fill_cell(cell, text=column_name[i_cell - int(index_insert)],
										 size=column_size, bold=column_bold, center=center,
										 fill=text_fill_column, R=fill_column_R, G=fill_column_G, B=fill_column_B)
				# 3. 插入其他行
				else:
					# 按行填充颜色
					text_fill = (True if i_row in text_fill_row else False)
					# 插入index
					if i_cell == 0 and index_insert == True:
						self.__fill_cell(cell, text=index_name[i_row - int(column_insert)],
										 size=size, bold=index_bold, center=center,
										 fill=text_fill, R=fill_row_R, G=fill_row_G, B=fill_row_B)
					# 按zscore值自动插入颜色
					elif i_cell in zscore_column:
						if df.iloc[i_row - int(column_insert), i_cell - int(index_insert)] == "":
							pass
						else:
							R, G, B = self.__zscore2RGB(df.iloc[i_row - int(column_insert), i_cell - int(index_insert)])
							self.__fill_cell(cell, text=df.iloc[i_row - int(column_insert), i_cell - int(index_insert)],
											 size=size, bold=False, center=center,
											 fill=True, R=R, G=G, B=B)
					# 插入内容
					else:
						self.__fill_cell(cell, text=df.iloc[i_row - int(column_insert), i_cell - int(index_insert)],
										 size=size, bold=False, center=center,
										 fill=text_fill, R=fill_row_R, G=fill_row_G, B=fill_row_B)

	def user_add_picture(self, picture_path, width=6):
		#	"""
		#	插入图片
		#	---------------------------------
		#		picture_path: 图片地址
		#		width: 图片尺寸
		#	"""
		self.document.add_picture(picture_path, width=Inches(width))

	def user_add_page_break(self):
		#	"""
		#	增加分页符
		#	"""
		self.document.add_page_break()

	def user_save(self, save_path):
		#	"""
		#	储存document
		#	---------------------------------
		#		save_path: 储存地址
		#	"""
		self.document.save(save_path)


# ------------------- 绘图函数 ------------------- #
def weightStackPlot(weights, num=8):
	'''
	功能：绘制各个资产或股票的历史权重面积图
	---------------------------------------------
	:param weights: DataFrame， index为日期%Y%m%d, columns为每个资产/股票的权重
	:param num: 显示横轴刻度数量
	:return: 简单显示结果
	'''

	x = np.array(range(0, weights.shape[0]))
	y = weights.values.astype('float32')

	XTicks = Series(np.linspace(0, weights.shape[0] - 1, num)).astype('int').tolist()

	fig = plt.figure(figsize=(10, 6))
	ax = fig.subplots()
	ax.stackplot(x, y.T)     #
	ax.legend(weights.columns.tolist())
	plt.xticks(XTicks, list(weights.index[XTicks]))
	plt.show()

	return True


def dfBarPlot(dat, figsize=(10, 6), title_str='', ylabel='', ygrid=False, show=True, save_path_name=''):
	'''
	功能：绘制bar图，可以多列同时绘制
	---------------------------------------------
	:param dat: DataFrame, 需要绘制bar图的数据
	:param figsize: 绘图尺寸
	:param title_str: 图片title
	:param ylabel: y轴名称
	:param ygrid: 横向网格线
	:param show: 是否显示
	:param save_path_name: 保存路径和名称
	:return: True
	'''

	if dat.shape[1] > 6:
		logging.warning('data columns > 6, please separate.')
		return None

	dat_len = dat.shape[0]
	colors = ['lightskyblue', 'yellowgreen', 'b', 'm', 'g', 'k']  	# 最多6个，太多也不好看
	columns = dat.columns.tolist()

	XTicks = Series(np.arange(0, dat_len)).astype('int').tolist()

	plt.figure(figsize=figsize)
	for i in range(len(columns)):
		plt.bar(np.array(range(dat_len)) + (1-0.2) / len(columns) * (i - int(len(columns) / 2)), dat.iloc[:, i],
				width=(1-0.2) / len(columns) + 0.01, facecolor=colors[i], edgecolor='white', label=columns[i], lw=1)

	plt.xticks(XTicks, list(dat.index[XTicks]))
	plt.legend(loc="upper left")   # label的位置在左上，没有这句会找不到label去哪了

	if ygrid:
		plt.grid(axis='y', linestyle='-.')
	plt.title(title_str)
	plt.ylabel(ylabel)

	if show:
		plt.show()
	else:
		plt.savefig(save_path_name + '.png')
		plt.close()

	return True


def calRollingSharpe2(dat, rollingPeriod=60, fct=244):

	dat = DataFrame(dat)
	rolling_ann_mean = dat.rolling(window=rollingPeriod).mean() * fct
	rolling_ann_std = dat.rolling(window=rollingPeriod).std() * (fct ** (1/2))

	rolling_sharpe = rolling_ann_mean / rolling_ann_std
	return rolling_sharpe



def cal_statistic_alongtime(dat):
	'''
	功能：计算时间轴上当前点位的中位数/平均数
	----------------------------------------------
	:param dat: DataFrame, 时间序列指标数据
	:return: DataFrame, 返回每个点之前的中位数或均值
	'''

	if isinstance(dat, Series):
		dat = DataFrame(dat)

	s = Series(range(dat.shape[0]))

	# 计算历史中位数
	dat_median = s.apply(lambda x: dat.iloc[:(x + 1), :].median())

	# 计算历史均值
	dat_mean = s.apply(lambda x: dat.iloc[:(x + 1), :].mean())

	return {'median': dat_median, 'mean': dat_mean}

def cal_quantile_alongtime(dat):
	'''
	功能：计算时间轴上当前点位的百分位
	--------------------------------------------
	:param dat: DataFrame, 时间序列指标数据
	:return: DataFrame, 返回每个点之前的百分位
	'''

	if isinstance(dat, Series):
		dat = DataFrame(dat)

	s = Series(range(dat.shape[0]))

	# 计算历史百分位数
	def func(dat, x):
		tmp = dat.iloc[:(x + 1), :].fillna(method='ffill')
		res = (tmp > tmp.iloc[-1, :]).sum() / (tmp.shape[0] - tmp.isnull().sum())    # np.nan > 1 -> False
		res = 1 - res
		return res

	res = s.apply(lambda x: func(dat, x))
	res.index = dat.index

	return res



def cal_adjust_fct(dat):
	'''
	功能：根据当前点位与历史中位数的比值计算估值调整因子
	---------------------------------------------------------
	:param dat: DataFrame, 时间序列指标数据
	:return: DataFrame, 调整因子，在1上行波动，空值调整因子设为1
	'''

	if isinstance(dat, Series):
		dat = DataFrame(dat)

	dat_median = cal_statistic_alongtime(dat)['median']

	adjust_fct = DataFrame(dat.values / dat_median.values, index=dat.index, columns=dat.columns)

	# 前面没有估值的数据调整因子设为1
	adjust_fct = adjust_fct.fillna(value=1)

	return adjust_fct



# -------------------------------- Class 回测模块 ------------------------------- #
class userPortfolio():

	def __init__(self, cash, position, asset):
		self.cash = cash
		self.position = position
		self.asset = asset


class SimpleBKTCore_V1():

	def __init__(self, initAmount=1e7, minCashRatio=0, buyFee=0, sellFee=0, cashRet=0.03, sharePerHand=1, tradingDatesHist = []):
		print('BackTest Configuration ... ')
		self.initAmount = initAmount
		self.minCashRatio = minCashRatio
		self.buyFee = buyFee
		self.sellFee = sellFee
		self.cashRet = cashRet
		self.sharePerHand = sharePerHand

		# 计算交易成本
		self.buyCostFactor = 1 + self.buyFee
		self.sellCostFactor = 1 - self.sellFee

		# 持仓和交易初始化
		cash = self.initAmount
		position = DataFrame()
		asset = self.initAmount
		self.portfolio = userPortfolio(cash, position, asset)
		self.tradeDetail = DataFrame()

		# 交易和持仓记录
		self.navRecord = dict()  # 初始空字典
		self.tradeRecord = dict()
		self.portfolioRecord = dict()

		# 回测过程的时间变量
		self.curDate = []
		self.tradingDatesHist = []

	def set_switch_variables(self, temporary_adjust_flag=False, stop_loss_flag=False):

		# 临时调整
		self.temporary_adjust_flag = temporary_adjust_flag

		# 止损
		self.stop_loss_flag = stop_loss_flag


	def print_info(self):

		# 回测参数配置
		print('BackTest Params Configuration: ')
		print('-- initAmount: ', self.initAmount)
		print('-- minCashRatio: ', self.minCashRatio)
		print('-- buyFee: ', self.buyFee)
		print('-- sellFee: ', self.sellFee)

		# 初始仓位设置
		print('Portfolio Initialization: ')
		print('-- cash: ', self.portfolio.cash)
		print('-- postion: ', self.portfolio.position)
		print('-- asset: ', self.portfolio.asset)

	def data_bkt_configuration(self, trade_info, weights, chgDates=[]):

		# 处理wts, trade_info, chgDates三个参数
		# 如果没有chgDates, 安装wts的日期index作为chgDates
		self.chgDates = chgDates
		self.weights = weights
		self.trade_info = trade_info

		if self.chgDates == []:
			self.chgDates = self.weights.index.tolist()


	def daily_update_portfolio(self, dt):
		'''
		功能：非调仓日，用持仓资产的收盘价更新持仓信息，此处是否需要传trade_info，还是仅仅传当天价格数据？  Ben@20200109
		-------------------------------------------------
		:param dt: str, 调仓日期%Y%m%d
		:return: dict, 返回更新后的持仓信息portfolio
		'''

		# 更新组合持仓收盘价
		position = self.portfolio.position.copy()
		position['close'] = self.trade_info.ix[dt, position.index.tolist()]
		position['eodAmt'] = position['shares'] * position['close']
		position['weight'] = position['eodAmt'] / position['eodAmt'].sum()

		self.portfolio.position = position.copy()
		self.portfolio.asset = self.portfolio.cash + self.portfolio.position['eodAmt'].sum()

	def daily_update_trade(self, tradeDetail):
		'''
		功能：更新当天的交易信息
		-------------------------------------------
		:param tradeDetail: 调仓产生的交易信息
		'''

		# 1) tradeDetail如果不为空,写入; 2) 如果self.tradeDetail为空，直接赋值; 3) 日期相同concat

		if not tradeDetail.empty:
			if not self.tradeDetail.empty:
				if tradeDetail['Date'].unique()[0] == self.tradeDetail['Date'].unique()[0]:
					self.tradeDetail = pd.concat([self.tradeDetail, tradeDetail])
				else:
					self.tradeDetail = tradeDetail
			else:
				self.tradeDetail = tradeDetail

	def allocate_cash2asset(self, dt):
		'''
		功能：把持仓中超过最低要求的现金分配到当前持仓中
		------------------------------------------------------
		:return: 现金分配后更新的持仓信息
		'''

		# 把多余的现金按比例分配到资产上
		if self.portfolio.cash / self.portfolio.asset - self.minCashRatio > 0.01:
			position = self.portfolio.position.copy()
			remainAvailCash = (self.portfolio.cash - self.minCashRatio * self.portfolio.asset) / (
					1 - self.minCashRatio + self.minCashRatio / self.buyCostFactor)

			self.portfolio.cash = self.portfolio.cash - remainAvailCash

			# 买入
			position['shares'] = position['shares'] + \
								 remainAvailCash / self.buyCostFactor * position['weight'] / position['close']
			position['eodAmt'] = position['shares'] * position['close']

			# 更新portfolio其他信息
			self.portfolio.position = position.copy()
			self.portfolio.asset = self.portfolio.cash + self.portfolio.position['eodAmt'].sum()

			# 补交易日记录
			tradeDetail = DataFrame(index=position.index.tolist(), columns=['tradePrice', 'shares', 'tradeAmt', 'flag', 'Date'])
			tradeDetail['shares'] = remainAvailCash / self.buyCostFactor * position['weight'] / position['close']
			tradeDetail['tradePrice'] = self.buyCostFactor * position['close']
			tradeDetail['tradeAmt'] = tradeDetail['tradePrice'] * tradeDetail['shares']
			tradeDetail['flag'] = 'buy'
			tradeDetail['Date'] = dt

			print('Allocate remaining cash to current position：', tradeDetail)

			self.daily_update_trade(tradeDetail)

	def update_portfolio(self, target_wts, dt):
		'''
		功能：portfolio为昨日持仓在当日尾盘更新后的持仓信息，在当日尾盘换成新的目标权重，返回更新后的持仓信息和当日交易信息
		---------------------------------------------------------------
		:param target_wts: Series, 当日调仓的目标资产权重
		:param dt: str, 调仓日期%Y%m%d
		:return: dict, 更新后的持仓信息portfolio和当日交易信息tradeDetail
		'''

		tradeDetail = DataFrame()

		if self.portfolio.position.empty:

			# 1. 记录交易信息
			tradeDetail = DataFrame(target_wts).rename(columns={target_wts.name: 'target_weight'}).copy()
			tradeDetail['tradePrice'] = self.buyCostFactor * self.trade_info.ix[
				dt, tradeDetail.index.tolist()]  # trade_info格式后面和多因子一致   Ben @20200110
			availCash = (1 - self.minCashRatio) * self.portfolio.cash
			tradeDetail['shares'] = availCash * tradeDetail['target_weight'] / tradeDetail['tradePrice']
			tradeDetail['tradeAmt'] = tradeDetail['tradePrice'] * tradeDetail['shares']
			self.portfolio.cash = self.portfolio.cash - tradeDetail['tradeAmt'].sum()

			tradeDetail['flag'] = 'buy'
			tradeDetail['Date'] = dt
			tradeDetail = tradeDetail.drop(columns=['target_weight'])

			# 2. 记录收盘信息
			updatedPosition = tradeDetail[['shares']].copy()
			updatedPosition['close'] = self.trade_info.ix[dt, updatedPosition.index.tolist()]
			updatedPosition['eodAmt'] = updatedPosition['close'] * updatedPosition['shares']
			updatedPosition['weight'] = updatedPosition['eodAmt'] / updatedPosition['eodAmt'].sum()

		else:

			# 计算可用资金, 假设按当日净值计算 (注意是当日净值，该函数之前有daily_update_portfolio函数)  Ben @20190630
			#  理论上要按照交易价格算
			availCash = (1 - self.minCashRatio) * (self.portfolio.cash +
											  self.portfolio.position['eodAmt'].sum() * self.sellCostFactor)
			# 这里availCash假设把所有现在资产都卖了，所有资产扣除赎回费，实际不会卖掉所有的,

			# 目标权重
			target_position = DataFrame(target_wts).rename(columns={target_wts.name: 'target_weight'})
			target_position['targetAmt'] = (availCash / self.buyCostFactor) * target_position['target_weight']
			# 会导致卖的多买的少  Ben @20190630

			# 计算轧差
			chgAmt = target_position.join(self.portfolio.position, how='outer')  # target_position应该也是需要price信息
			chgAmt = chgAmt.rename(columns={'shares': 'orgShares'})

			# 填充所有资产价格，这里很关键；之前多因子模型有   Ben @20190731
			chgAmt['close'] = self.trade_info.ix[dt, chgAmt.index.tolist()]

			# nan -> 0
			chgAmt = chgAmt.fillna(value=0)
			chgAmt['deltaAmt'] = chgAmt['targetAmt'] - chgAmt['eodAmt']

			# 计算大致买卖份额
			chgAmt['chgShares'] = abs(chgAmt['deltaAmt']) / chgAmt['close']
			# 1. 这个close是当日的，理论上无法实现，但是指数是可以做到 Ben@20190630

			chgAmt.ix[chgAmt['deltaAmt'] < 0, 'chgShares'] = - chgAmt.ix[chgAmt['deltaAmt'] < 0, 'chgShares']

			# 其实已经确定了  Ben @20190731
			chgAmt['shares'] = chgAmt['orgShares'] + chgAmt['chgShares']

			# 卖出持仓
			tradeDetailSell = DataFrame()
			positionSell = chgAmt.ix[chgAmt['chgShares'] < 0, :].copy()

			if not positionSell.empty:  # 这里不能用~来取反，要用not   Ben @20190731
				positionSell['tradePrice'] = self.sellCostFactor * positionSell.ix[:, 'close']
				self.portfolio.cash = self.portfolio.cash + (
						abs(positionSell['chgShares']) * positionSell['tradePrice']).sum()
				positionSell['shares'] = positionSell['orgShares'] + positionSell['chgShares']
				chgAmt.ix[positionSell.index.tolist(), 'shares'] = positionSell['shares']

				# 加入交易信息tradeDetail
				tradeDetailSell = positionSell[['tradePrice', 'chgShares']].copy()
				tradeDetailSell['tradeAmt'] = abs(tradeDetailSell['chgShares']) * tradeDetailSell['tradePrice']
				tradeDetailSell['flag'] = 'sell'
				tradeDetailSell['Date'] = dt

			# 9. 计算买入仓位并买入
			tradeDetailBuy = DataFrame()
			positionBuy = chgAmt.ix[chgAmt['chgShares'] > 0, :].copy()

			if not positionBuy.empty:
				positionBuy['tradePrice'] = self.buyCostFactor * positionBuy['close']
				positionBuy['shares'] = positionBuy['orgShares'] + positionBuy['chgShares']
				chgAmt.ix[positionBuy.index.tolist(), 'shares'] = positionBuy['shares']

				self.portfolio.cash = self.portfolio.cash - (
						positionBuy['tradePrice'] * positionBuy['chgShares']).sum()

				# 加入交易信息tradeDetail
				tradeDetailBuy = positionBuy.ix[:, ['tradePrice', 'chgShares']].copy()
				tradeDetailBuy['tradeAmt'] = tradeDetailBuy['chgShares'] * tradeDetailBuy['tradePrice']
				tradeDetailBuy['flag'] = 'buy'
				tradeDetailBuy['Date'] = dt

			if (not tradeDetailSell.empty) | (not tradeDetailBuy.empty):
				tradeDetail = pd.concat([tradeDetailSell, tradeDetailBuy], sort=False)

			tradeDetail = tradeDetail.rename(columns={'chgShares': 'shares'})

			# 10. 更新仓位信息
			updatedPosition = chgAmt.ix[chgAmt['shares'] > 1e-6, ['shares']]  # 删除卖掉了的仓位

			updatedPosition['close'] = self.trade_info.ix[dt, updatedPosition.index.tolist()]  # 这里有全局变量  Ben @20190703
			updatedPosition['eodAmt'] = updatedPosition['shares'] * updatedPosition['close']
			updatedPosition['weight'] = updatedPosition['eodAmt'] / updatedPosition['eodAmt'].sum()

		# 汇总和调整顺序
		tradeDetail = tradeDetail[['tradePrice', 'shares', 'tradeAmt', 'flag', 'Date']]
		updatedPosition = updatedPosition[['shares', 'close', 'eodAmt', 'weight']]

		# 仓位更新后，更新portfolio其他信息
		self.portfolio.position = updatedPosition.copy()
		self.portfolio.asset = self.portfolio.cash + self.portfolio.position['eodAmt'].sum()

		self.daily_update_trade(tradeDetail)

		# 把多余的现金按比例分配到资产上
		self.allocate_cash2asset(dt)

	def check_portfolio(self):
		'''
		功能：检查当日持仓数据是否有问题
		-----------------------------------------
		:return: 如果剩余现金或资产总值小于0，空值或无穷值，回测过程报错
		'''

		# portfolio.cash 出现问题
		rule_cash_1 = self.portfolio.cash < -1e-2
		rule_cash_2 = self.portfolio.cash is np.nan
		rule_cash_3 = self.portfolio.cash is np.inf

		# portfolio.asset 出现问题
		rule_asset_1 = self.portfolio.asset < -1e-2
		rule_asset_2 = self.portfolio.asset is np.nan
		rule_asset_3 = self.portfolio.asset is np.inf

		if rule_cash_1 or rule_cash_2 or rule_cash_3 or rule_asset_1 or rule_asset_2 or rule_asset_3:

			logging.error('Cash and Asset cannot be < 0 or NaN or Inf !!!')
			logging.warning('Cash: ' + repr(self.portfolio.cash) + ', Asset: ' + repr(self.portfolio.asset))

			raise ValueError

	def record_trade_portfolio(self, dt):

		if (not self.tradeDetail.empty) & (self.tradeDetail['Date'].unique()[0] == dt):
			self.tradeRecord[dt] = self.tradeDetail.copy()

		self.portfolioRecord[dt] = copy.deepcopy(self.portfolio)   # 应该是赋值问题
		# print(self.portfolioRecord[dt].position)
		# 开始模拟回测
		self.navRecord[dt] = self.portfolio.asset / self.initAmount

	# def temporary_adjust(self, ref_wts, dt):
	#
	# 	if self.temporary_adjust_flag:
	#
	# 		# 如果某一天资产比例偏离超过5%，用新的比例换仓
	# 		cur_wts = self.portfolio.positition['weight']
	# 		target_wts = ref_wts.ix[dt, cur_wts.index.tolist()].copy()
	#
	# 		if abs(target_wts - cur_wts).sum() > 1e-1:
	# 			print('当天资产比例与目标比例偏差超过5%，按目标权重调整仓位！')
	# 			print(cur_wts)
	# 			print(target_wts)
	#
	# 			self.update_portfolio(target_wts, dt)

	# # 止损可以作为一个模块，后续再更新，止损的多样性
	# def stop_loss(self, trade_info, dt, stopFlag=True, stopLevel=1):
	# 	'''
	# 	功能：止损模块，
	# 	:param dt: str, 调仓日期%Y%m%d
	# 	:param trade_info: DataFrame, 资产历史行情数据,复权日行情
	# 	:param portfolio: dict, 当日持仓数据
	# 	:param stopFlag: True/False, 是否止损
	# 	:param stopLevel: 1/2，止损级别，半止损和全止损
	# 	:return: dict, 当日止损后的持仓数据
	# 	'''
	#
	# 	s_b = - 0.05
	# 	b_b = 0.03
	# 	lbw = 60
	# 	stopLossDone = 0
	#
	# 	# 计算最近60个交易日涨跌幅
	# 	trading_dates = trade_info.index.tolist()
	# 	ind = trading_dates.index(dt)
	# 	stockDownRatio = trade_info.ix[ind, 0] / trade_info.ix[ind - lbw, 0] - 1
	# 	bondUpRatio = trade_info.ix[ind, 1] / trade_info.ix[ind - lbw, 1] - 1
	#
	# 	stopSignal = (stockDownRatio < s_b) & (bondUpRatio < b_b)
	#
	# 	if stopSignal:
	# 		# 全止损，半仓止损
	# 		print('Stock changes in ' + repr(lbw) + ' trading dates :' + repr(stockDownRatio * 100)[:4] + '% < ' + repr(
	# 			s_b * 100) + '%' +
	# 			  ', and Bond changes: ' + repr(bondUpRatio * 100)[:4] + '% < ' + repr(b_b * 100) + '%')
	#
	# 		if stopFlag:
	# 			if stopLevel == 1:
	# 				# 半止损
	# 				print('half stop loss ...')
	# 				self.portfolio.cash = self.portfolio.cash + self.sellCostFactor * self.portfolio.position['eodAmt'].sum() / 2
	# 				self.portfolio.position[['shares', 'eodAmt']] = self.portfolio.position[['shares', 'eodAmt']] / 2
	# 				self.portfolio.asset = self.portfolio.cash + self.portfolio.position['eodAmt'].sum()
	# 				stopLossDone = 1
	#
	# 			elif stopLevel == 2:
	# 				# 全止损
	# 				print('total stop loss...')
	# 				self.portfolio.cash = self.portfolio.cash + self.sellCostFactor * self.portfolio.position['eodAmt'].sum()
	# 				self.portfolio.position = DataFrame()
	# 				self.portfolio.asset = self.portfolio.cash
	# 				stopLossDone = 1
	#
	# 	# return portfolio, stopLossDone



	def strategy_bkt(self):

		# 如果有chgDates, 要以chgDates为准,作为对照的wts尽量是外部输入的参数
		self.trade_info = self.trade_info.ix[self.chgDates[0]:, :].copy()

		chgDates = self.chgDates
		actionDates = self.trade_info.index.tolist()
		columns = self.weights.columns.tolist()

		# 回测流程
		start_bkt_flag = 1

		for actionDate in actionDates:
			# print(actionDate)

			if actionDate in chgDates:

				if start_bkt_flag == 1:
					start_bkt_flag = 0
				else:
					self.daily_update_portfolio(actionDate)

				# 计算预期调整权重
				target_wts = self.weights.ix[actionDate, columns]

				# 更新组合
				self.update_portfolio(target_wts, actionDate)  # 默认更新portfolio和tradeDetail，tradeDetail，一天内concat

				# 保留组合信息和交易信息
				self.check_portfolio()
				self.record_trade_portfolio(actionDate)

				print('BackTest Date: ' + actionDate + ' , totalvalue: ' + repr(self.navRecord[actionDate]) + ', 今日为调仓时点')

			else:
				# daily更新portfolio
				self.daily_update_portfolio(actionDate)

				# 临时调整
				# self.temporary_adjust(ref_wts=)

				# 保留组合信息和交易信息
				self.check_portfolio()
				self.record_trade_portfolio(actionDate)

				print('BackTest Date: ' + actionDate + ' , totalvalue: ' + repr(self.navRecord[actionDate]))

				# 这里可能加入非常规调整

			# 这里加入止损信息


def daily_ret_characteristic(ret):
	'''
	功能：计算策略日收益率特征，包括 日收益均值/偏度/峰度/最大/小值，发生日期，正收益概率/VaR/CVaR
	-----------------------------------------------------------------------------------
	:param ret: DataFrame, 日收益率序列
	:return: DataFrame, 日收益率特征
	'''

	res = DataFrame(index=ret.columns, columns=['mean', 'skewness', 'kurtosis', 'max', 'max_date', 'min', 'min_date', 'win_prob',
												'VaR(0.05)', 'CVaR(0.05)', 'VaR(0.01)', 'CVaR(0.01)'])

	res.loc[:, 'mean'] = ret.mean()
	res.loc[:, 'skewness'] = ret.skew()
	res.loc[:, 'kurtosis'] = ret.kurtosis()

	res.loc[:, 'max'] = ret.max()
	res.loc[:, 'max_date'] = ret.idxmax()

	res.loc[:, 'min'] = ret.min()
	res.loc[:, 'min_date'] = ret.idxmin()

	res.loc[:, 'win_prob'] = (ret > 0).sum() / ret.shape[0]

	for column in ret.columns.tolist():

		sorted_tmp = ret[column].sort_values()
		if ret.shape[0] > 20:
			res.loc[column, 'VaR(0.05)'] = sorted_tmp[int(sorted_tmp.shape[0] * 0.05)]
			res.loc[column, 'CVaR(0.05)'] = sorted_tmp[:(int(sorted_tmp.shape[0] * 0.05) + 1)].mean()

		if ret.shape[0] > 100:
			res.loc[column, 'VaR(0.01)'] = sorted_tmp[int(sorted_tmp.shape[0] * 0.01)]
			res.loc[column, 'CVaR(0.01)'] = sorted_tmp[:(int(sorted_tmp.shape[0] * 0.01) + 1)].mean()

	return res


def cal_holding_return(nav):
	'''
	功能：计算持有期收益，1M（20个交易日）, 3M（60个交易日）, 6M（120个交易日）, 1Y（240个交易日）
	注意：可以优化，用日历日来做!!!
	--------------------------------------------------------------------------------------
	:param nav: Series/DataFrame, 策略净值
	:return: DataFrame, 持有其收益信息
	'''
	if isinstance(nav, DataFrame):
		nav = nav.iloc[:, 0]

	holdDaysList = [20, 60, 120, 240]
	res = DataFrame(index=holdDaysList, columns=['mean', 'max', 'max_period', 'min', 'min_period', 'prob'])

	for holdDays in holdDaysList:
		holdDays_ret = []
		for i in range(nav.shape[0] - holdDays):
			holdDays_ret.append(nav[i + holdDays] / nav[i] - 1)

		s1 = Series(holdDays_ret)
		s1_mean = s1.mean()

		s1_max = s1.max()
		s1_max_period = nav.index[np.where(s1 == s1_max)[0]][0] + ' - ' + nav.index[np.where(s1 == s1_max)[0] + holdDays][0]

		s1_min = s1.min()
		s1_min_period = nav.index[np.where(s1 == s1_min)[0]][0] + ' - ' + nav.index[np.where(s1 == s1_min)[0] + holdDays][0]

		s1_prob = s1[s1 > 0].shape[0] / s1.shape[0]

		res.loc[holdDays, :] = [s1_mean, s1_max, s1_max_period, s1_min, s1_min_period, s1_prob]

	res.index = ['1M', '3M', '6M', '1Y']

	return res


def nav2yearlystats(dat):
	'''
	功能：根据净值或指数计算年度收益，年度收益为本年度最后一个交易日与上年度最后一个交易日的比值减1
	--------------------------
	:param dat: dataframe或Series, index为时间序列且格式为 %Y%m%d
	:return: index为年份的年度收益结果
	'''

	dat = DataFrame(dat) if isinstance(dat, Series) else dat

	# 根据dat索引确定计算年度收益起止日
	start_year = int(dat.index[0][:4])
	end_year = int(dat.index[-1][:4])
	year_list = np.arange(start_year, end_year + 1)

	# 计算年度收益/波动率/最大回撤
	yearly_ret = DataFrame(0, columns=dat.columns, index=year_list)

	ret = nav2return(dat)
	yearly_vol = DataFrame(0, columns=dat.columns, index=year_list)
	yearly_maxdrawback = DataFrame(0, columns=dat.columns, index=year_list)


	for iYear in year_list:
		if list(year_list).index(iYear) == 0:
			tmp = dat.loc[(repr(iYear) + '0101'): (repr(iYear) + '1231'), :].copy()
		else:
			tmp_shift1 = dat.loc[(repr(iYear-1) + '0101'): (repr(iYear-1) + '1231'), :].copy()
			tmp = dat.loc[tmp_shift1.index[-1]: (repr(iYear) + '1231'), :].copy()

		yearly_ret.loc[iYear, :] = (tmp.iloc[-1, :] / tmp.iloc[0, :] - 1)
		yearly_maxdrawback.loc[iYear, :] = cal_drawback(tmp)['maxdrawback_info']['maxdrawback']

		tmp_ret = ret.loc[(repr(iYear) + '0101'): (repr(iYear) + '1231'), :].copy()
		yearly_vol.loc[iYear, :] = tmp_ret.std() * np.sqrt(244)


	return {'yearly_ret': yearly_ret,
			'yearly_vol': yearly_vol,
			'yearly_maxdrawback': yearly_maxdrawback}


def cal_drawback(nav):
	'''
	功能：计算净值区间的回撤，最大回撤及发生的日期
	---------------------------------------------------------
	:param nav: Series/DataFrame, 策略净值
	:return: dict, 回撤和最大回撤相关信息
	'''

	# 计算drawback
	nav = DataFrame(nav) if isinstance(nav, Series) else nav
	drawback = DataFrame(np.zeros(nav.shape), index=nav.index, columns=nav.columns)

	for i in range(0, nav.shape[0]):
		drawback[i:(i + 1)] = nav[i:(i + 1)] / nav[:(i + 1)].max() - 1

	maxdrawback = drawback.min()
	maxdrawback_date = drawback.idxmin()

	# 计算之前发生的最高点
	maxdrawback_start_date = Series(index=maxdrawback_date.index)
	for item in maxdrawback_date.index.tolist():
		maxdrawback_start_date[item] = nav.loc[:maxdrawback_date[item], item].idxmax()
	maxdrawback_start_date = maxdrawback_start_date.astype('int').astype('str')

	# 最大回撤信息
	maxdrawback_info = pd.concat([maxdrawback, maxdrawback_start_date, maxdrawback_date], axis=1, sort=True)
	maxdrawback_info.columns = ['maxdrawback', 'maxdrawback_start_date', 'maxdrawback_date']

	return {'drawback': drawback,
			'maxdrawback_info': maxdrawback_info}



# 寻找最大回撤的位置
def period_maxdrawdown(nav, start_dt='', end_dt=''):

	# 设置参数
	if len(start_dt) == 0:
		start_dt = nav.index.tolist()[0]
	if len(end_dt) == 0:
		end_dt = nav.index.tolist()[-1]

	# 1. 计算drawdown
	nav = DataFrame(nav) if isinstance(nav, Series) else nav
	drawdown = DataFrame(np.zeros(nav.shape), index=nav.index, columns=nav.columns)

	for i in range(0, nav.shape[0]):
		drawdown[i:(i + 1)] = 1 - nav[i:(i + 1)] / nav[:(i + 1)].max()

	drawdown = drawdown
	maxdrawdown = drawdown.max()

	# 2. 计算区间回撤信息
	drawdown = drawdown.ix[start_dt:, :]
	drawdown = drawdown.ix[:end_dt, :]

	nav = nav.ix[start_dt:, :]
	nav = nav.ix[:end_dt, :]

	maxdrawdown_info = DataFrame(index=drawdown.columns, columns=['maxdrawdownDate', 'navPeakDate', 'chg'])

	for column in drawdown.columns.tolist():
		tmp = drawdown[column]
		tmp_nav = nav[column]

		ind = np.where(tmp == tmp.max())[0][0]
		maxdrawdownDate = tmp.index[ind]

		indmax = np.where(tmp_nav[:(ind+1)] == tmp_nav[:(ind+1)].max())[0][0]
		navPeakDate = tmp_nav[:(ind+1)].index[indmax]

		maxdrawdown_info.ix[column, :] = [maxdrawdownDate, navPeakDate, tmp.max()]

	return {'drawdown': drawdown,
			'maxdrawdown': maxdrawdown,
			'maxdrawdown_info': maxdrawdown_info
			}








